#!/usr/bin/env python3
"""
Scan PDF and Word documents (.docx) for Danish CPR numbers and dates.
Handles text-based and image-based (scanned) PDFs automatically via OCR.
Supports masking, full anonymisation, dry-run preview, and JSON logging.

Supported formats: .pdf, .docx, .xlsx, .xlsm, .csv, .jpg, .jpeg, .png, .bmp, .tiff, .webp
  (.doc requires conversion: soffice --headless --convert-to docx file.doc)

Usage:
    python document_scanner.py file.pdf
    python document_scanner.py file.docx
    python document_scanner.py file1.pdf file2.docx spreadsheet.xlsx /path/to/folder/

Options:
    --mask            Redact CPR numbers only    -> <n>_masked.pdf/.docx
    --anonymise       Redact all personal data   -> <n>_anonymised.pdf/.docx
                      (CPR, names, addresses, phone numbers, emails)
    --dry-run         Scan and report without writing any output files
    --log FILE        Write a structured JSON log of all findings to FILE
    --older-than DAYS List files with CPR numbers AND dates older than DAYS
    --ocr             Force OCR on every page (even if text is extractable)
    --lang LANG       Tesseract language(s), default: dan+eng
    --dpi DPI         DPI for OCR image rendering, default: 300
    --poppler PATH    Path to Poppler bin folder (Windows only)

Dependencies:
    pip install pdfplumber pdf2image pytesseract pypdf reportlab spacy python-docx openpyxl opencv-python
    python -m spacy download da_core_news_lg   # Danish NER model (~500 MB)

    System packages:
        macOS:  brew install tesseract tesseract-lang poppler
        Linux:  sudo apt install tesseract-ocr tesseract-ocr-dan poppler-utils

    Note: Python 3.12 recommended -- spaCy does not yet support Python 3.14.

Recommended workflow:
    # 1. Dry run first to audit without writing anything
    python document_scanner.py /folder/ --anonymise --dry-run --log audit.json

    # 2. Run for real once satisfied
    python document_scanner.py /folder/ --anonymise --log run.json
"""

import argparse
import hashlib
import io
import json
import logging
import re
import sqlite3
import sys
from datetime import date, datetime, timedelta
from pathlib import Path

# Suppress pdfminer's noisy font-descriptor warnings that appear when PDFs
# contain malformed or incomplete font definitions.  These do not affect text
# extraction or CPR detection — the warning is informational only.
logging.getLogger("pdfminer").setLevel(logging.ERROR)
logging.getLogger("pdfminer.pdffont").setLevel(logging.ERROR)
logging.getLogger("pdfminer.pdfpage").setLevel(logging.ERROR)
logging.getLogger("pdfplumber").setLevel(logging.ERROR)
# ── Dependency checks ──────────────────────────────────────────────────────────

try:
    import pdfplumber
except ImportError:
    print("Missing dependency. Install with: pip install pdfplumber")
    sys.exit(1)

try:
    from pdf2image import convert_from_path
    PDF2IMAGE_OK = True
except ImportError:
    PDF2IMAGE_OK = False

try:
    import pytesseract
    TESSERACT_OK = True
except ImportError:
    TESSERACT_OK = False

OCR_AVAILABLE = PDF2IMAGE_OK and TESSERACT_OK

try:
    from pypdf import PdfReader, PdfWriter
    from reportlab.pdfgen import canvas as rl_canvas
    from reportlab.lib.colors import black as rl_black
    MASK_AVAILABLE = True
except ImportError:
    MASK_AVAILABLE = False

try:
    import fitz as _fitz          # PyMuPDF — for secure (sanitised) PDF redaction
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    import spacy
    SPACY_OK = True
except ImportError:
    SPACY_OK = False

try:
    from docx import Document as DocxDocument
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

try:
    import openpyxl
    XLSX_OK = True
except ImportError:
    XLSX_OK = False

# cv2 is imported lazily inside _get_cv2() to avoid macOS recursion errors.
# Never import cv2 at module level or from server.py.
CV2_OK = False

def _face_log(msg: str):
    """Debug logging — file output disabled."""
    import sys as _sys
    print(msg, file=_sys.stderr, flush=True)
_cv2_version = None
_cv2_import_error = None
_cv2_mod = None
_np_mod  = None

def _get_cv2():
    """Return (cv2, numpy) tuple, importing once on first call.

    In a PyInstaller bundle we exclude cv2/__init__.py entirely (it causes a
    macOS arm64 recursion crash) and load cv2.abi3.so directly instead.
    Outside the bundle, plain 'import cv2' works normally.
    """
    global CV2_OK, _cv2_version, _cv2_import_error, _cv2_mod, _np_mod
    if _cv2_mod is not None:
        return _cv2_mod, _np_mod
    if _cv2_import_error is not None:
        return None, None  # already tried and failed
    try:
        import sys as _sys
        import numpy as _np

        if getattr(_sys, "frozen", False):
            # Bundle has cv2.abi3.so but NOT cv2/__init__.py.
            # Load the .so directly and register it as 'cv2'.
            import importlib.util as _ilu
            import types as _types
            from pathlib import Path as _Path

            _so = _Path(_sys._MEIPASS) / "cv2" / "cv2.abi3.so"
            if not _so.exists():
                raise RuntimeError(f"cv2.abi3.so not found at {_so}")

            _spec = _ilu.spec_from_file_location("cv2", str(_so),
                        submodule_search_locations=[])
            _cv2 = _ilu.module_from_spec(_spec)
            _sys.modules["cv2"] = _cv2   # register before exec to break cycles
            _spec.loader.exec_module(_cv2)

            # Wire up cv2.data.haarcascades for cascade path resolution
            _data = _types.ModuleType("cv2.data")
            _data.haarcascades = str(_Path(_sys._MEIPASS) / "cv2" / "data") + "/"
            _sys.modules["cv2.data"] = _data
            _cv2.data = _data
        else:
            import cv2 as _cv2

        if not hasattr(_cv2, "imread"):
            raise RuntimeError(
                f"cv2 binary not loaded (file: {getattr(_cv2, '__file__', '?')})"
            )

        _cv2_version      = getattr(_cv2, "__version__", "unknown")
        CV2_OK            = True
        _cv2_mod          = _cv2
        _np_mod           = _np
        _cv2_import_error = None
    except Exception as e:
        CV2_OK            = False
        _cv2_import_error = str(e)
        import sys as _sys
        _sys.modules.pop("cv2", None)  # clean up partial registration
    return _cv2_mod, _np_mod

# spaCy model preference: large Danish → medium → small → multilingual → English fallback
SPACY_MODEL_PREFERENCE = [
    "da_core_news_lg", "da_core_news_md", "da_core_news_sm",
    "xx_ent_wiki_sm", "en_core_web_sm",
]
_NLP = None  # lazy-loaded singleton

def load_nlp():
    """Load the best available spaCy model. Returns model or None."""
    global _NLP
    if _NLP is not None:
        return _NLP
    if not SPACY_OK:
        return None

    import sys as _sys
    _frozen = getattr(_sys, "frozen", False)

    for model_name in SPACY_MODEL_PREFERENCE:
        try:
            import importlib as _il
            _mod = _il.import_module(model_name)
            _NLP = _mod.load()
            print(f"  [NER] Loaded spaCy model: {model_name}", flush=True)
            _face_log(f"[NER] Loaded spaCy model: {model_name}")
            return _NLP
        except Exception as _e:
            _face_log(f"[NER] {model_name} failed: {_e} (frozen={_frozen})")
            continue
    return None


# ── OCR page cache ───────────────────────────────────────────────────────────

_OCR_CACHE_PATH = Path.home() / ".document_scanner_ocr_cache.db"

class OCRCache:
    """
    SQLite-backed cache for OCR text extraction.

    Key:   SHA-256 of the raw page image bytes + lang string
    Value: extracted text string

    This means:
    - Rescanning the same file reuses cached text (near-instant).
    - Editing a file invalidates its pages (hash changes).
    - Different OCR language settings get separate cache entries.
    - The cache is shared across all processes (safe: writes are idempotent).
    """

    def __init__(self, path: Path = _OCR_CACHE_PATH):
        self._path = path
        self._conn: sqlite3.Connection | None = None

    def _connect(self) -> sqlite3.Connection:
        if self._conn is None:
            conn = sqlite3.connect(str(self._path), check_same_thread=False,
                                   timeout=10)
            conn.execute("""
                CREATE TABLE IF NOT EXISTS ocr_cache (
                    key   TEXT PRIMARY KEY,
                    text  TEXT NOT NULL,
                    ts    INTEGER NOT NULL
                )
            """)
            conn.execute("CREATE INDEX IF NOT EXISTS idx_ts ON ocr_cache(ts)")
            conn.commit()
            self._conn = conn
        return self._conn

    @staticmethod
    def _key(image_bytes: bytes, lang: str) -> str:
        h = hashlib.sha256(image_bytes)
        h.update(lang.encode())
        return h.hexdigest()

    def get(self, image_bytes: bytes, lang: str) -> str | None:
        key = self._key(image_bytes, lang)
        try:
            row = self._connect().execute(
                "SELECT text FROM ocr_cache WHERE key=?", (key,)
            ).fetchone()
            return row[0] if row else None
        except Exception:
            return None

    def put(self, image_bytes: bytes, lang: str, text: str) -> None:
        key = self._key(image_bytes, lang)
        ts  = int(datetime.now().timestamp())
        try:
            self._connect().execute(
                "INSERT OR REPLACE INTO ocr_cache(key, text, ts) VALUES(?,?,?)",
                (key, text, ts),
            )
            self._connect().commit()
        except Exception:
            pass

    def prune(self, max_entries: int = 50_000) -> None:
        """Delete oldest entries when the cache grows beyond max_entries."""
        try:
            conn = self._connect()
            n = conn.execute("SELECT COUNT(*) FROM ocr_cache").fetchone()[0]
            if n > max_entries:
                to_del = n - max_entries
                conn.execute("""
                    DELETE FROM ocr_cache
                    WHERE key IN (
                        SELECT key FROM ocr_cache ORDER BY ts ASC LIMIT ?
                    )
                """, (to_del,))
                conn.commit()
        except Exception:
            pass

    def clear(self) -> None:
        try:
            self._connect().execute("DELETE FROM ocr_cache")
            self._connect().commit()
        except Exception:
            pass

    def stats(self) -> dict:
        try:
            conn = self._connect()
            n    = conn.execute("SELECT COUNT(*) FROM ocr_cache").fetchone()[0]
            size = self._path.stat().st_size if self._path.exists() else 0
            return {"entries": n, "size_bytes": size}
        except Exception:
            return {"entries": 0, "size_bytes": 0}


# Module-level singleton — shared within a process
_ocr_cache = OCRCache()


def ocr_page_cached(image, lang: str) -> str:
    """
    Run Tesseract OCR on `image`, returning cached text when available.
    Falls back to uncached OCR if the cache is unavailable.
    """
    import io as _io
    # Serialise image to bytes for hashing (use PNG for lossless round-trip)
    buf = _io.BytesIO()
    image.save(buf, format="PNG")
    img_bytes = buf.getvalue()

    cached = _ocr_cache.get(img_bytes, lang)
    if cached is not None:
        return cached

    text = ocr_page(image, lang)
    _ocr_cache.put(img_bytes, lang, text)
    _ocr_cache.prune()
    return text


# ── Patterns ──────────────────────────────────────────────────────────────────

# Danish CPR: DDMMYY-XXXX  or  DDMMYYXXXX  (optional space/dash separator)
CPR_PATTERN = re.compile(r"\b(\d{2})(\d{2})(\d{2})[-\s]?(\d{4})\b")

DATE_PATTERNS = [
    (re.compile(r"\b(\d{4})[-/](\d{1,2})[-/](\d{1,2})\b"), "ISO YYYY-MM-DD"),
    (re.compile(r"\b(\d{1,2})[.\-/](\d{1,2})[.\-/](\d{4})\b"), "DD.MM.YYYY"),
    (re.compile(r"\b(\d{1,2})[.\-/](\d{1,2})[.\-/](\d{2})\b"), "DD.MM.YY"),
    (re.compile(
        r"\b(\d{1,2})\.\s*(januar|februar|marts|april|maj|juni|juli|"
        r"august|september|oktober|november|december)\s+(\d{4})\b", re.IGNORECASE),
     "D. maaned YYYY"),
    (re.compile(
        r"\b(\d{1,2})\s+(January|February|March|April|May|June|July|"
        r"August|September|October|November|December)\s+(\d{4})\b", re.IGNORECASE),
     "D Month YYYY"),
    (re.compile(
        r"\b(January|February|March|April|May|June|July|August|"
        r"September|October|November|December)\s+(\d{1,2}),?\s+(\d{4})\b", re.IGNORECASE),
     "Month D, YYYY"),
]

# ── Regex patterns for PII beyond CPR ─────────────────────────────────────────

# Danish phone: 8 digits, optionally grouped in pairs/fours with spaces or dashes
# Also matches +45 prefix
PHONE_PATTERN = re.compile(
    r"(?<!\d)(?:\+45[\s\-]?)?(?:\d{2}[\s\-]?){3}\d{2}(?!\d)"
)

EMAIL_PATTERN = re.compile(
    r"\b[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}\b"
)

# Danish IBAN: DKxx xxxx xxxx xxxx xx  (18 digits total, starts DK)
# Also plain REG/Konto format:  RRRR KKKKKKKKKK  (4-digit reg + up to 10-digit account)
IBAN_PATTERN = re.compile(
    r"\bDK\d{2}[\s\-]?\d{4}[\s\-]?\d{4}[\s\-]?\d{4}[\s\-]?\d{2}\b",
    re.IGNORECASE,
)
# Danish bank account REG no + account number (e.g. "1234 1234567890" or "REG: 1234 Konto: 12345")
REG_KONTO_PATTERN = re.compile(
    r"(?:reg(?:ister|\.|:)?\s*(?:nr\.?\s*)?)?\b(\d{4})[\s\-]+(\d{6,10})\b",
    re.IGNORECASE,
)

# Danish address: "Streetname 12" or "Streetname 12A, 2. tv" style
# Matches: word(s) ending in common suffix + number + optional letter + optional floor/door
DANISH_ADDRESS_PATTERN = re.compile(
    r"\b([A-ZÆØÅ][a-zæøåA-ZÆØÅ\-]+"           # Street name start (capitalised)
    r"(?:\s+[A-ZÆØÅ]?[a-zæøåA-ZÆØÅ\-]+)*"     # additional words (allé, vej, gade, …)
    r"(?:\s+(?:vej|gade|allé|alle|plads|torv|stræde|straede|boulevard|bvd|"
    r"have|park|skov|bakke|bjerg|dal|mark|eng|sø|dam|holm|bro|port))?)"  # optional suffix
    r"\s+(\d{1,4}[A-Za-z]?)"                   # house number (e.g. 12, 12A, 4B)
    r"(?:[,\s]+\d{1,2}\.?(?:\s*(?:sal|tv|th|mf|[0-9]{1,3}))?)?"  # optional floor/door
    r"(?:[,\s]+\d{4})?",                        # optional postcode
    re.UNICODE,
)

# Danish full name: two or more capitalised words (Firstname [Middle] Lastname).
# Covers Danish/Nordic letters. Excludes common non-name capitalisations by
# requiring at least 2 name parts and rejecting single-word matches.
# This is a reliable fallback for isolated names where spaCy has no context.
# Danish name pattern — matches "Firstname [Middle...] Lastname" where:
# - First name must be capitalised (A-ZÆØÅ start)
# - Middle/last parts can be lowercase (handles "Frode holm truelsen" style)
# - Particles like "la", "de", "van" are allowed
# - Accented chars (Bräuner, Buéno) supported via \w
DANISH_NAME_PATTERN = re.compile(
    r"\b([A-ZÆØÅ][\w\-]{1,})"              # First name — must be capitalised
    r"(?:\s+[\w\-]{2,})*"                   # Optional middle parts (any case)
    r"\s+([\w\-]{2,})\b",                   # Last name (any case, min 2 chars)
    re.UNICODE,
)

# Words that are NEVER part of a person name — used to filter regex false positives.
# NOTE: Month names that are also Danish first names (April, August, Juni, Juli,
# Mai/Maj) are intentionally excluded so we don't block real names.
_NAME_STOPWORDS = {
    # Calendar — only months not used as first names
    "januar","februar","marts","september","oktober","november","december",
    "january","february","march","september","october","november","december",
    # Days of week
    "mandag","tirsdag","onsdag","torsdag","fredag","lordag","sondag",
    "monday","tuesday","wednesday","thursday","friday","saturday","sunday",
    # Business / document words
    "dk","cvr","cpr","att","re","fwd","til","fra","dato","side","total",
    "faktura","invoice","bilag","nota","subtotal","moms","vat","inkl","ekskl",
    "afdeling","department","company","virksomhed","adresse","address",
    "telefon","email","website","homepage","tlf","mobil","mobile",
}

# Particles that can appear lowercase inside a name ("la", "de", "van", etc.)
_NAME_PARTICLES = {"la","le","de","du","van","von","af","of","the"}

def _is_name_match(m) -> bool:
    """Return True if a DANISH_NAME_PATTERN match looks like a real person name."""
    parts = m.group(0).split()
    if len(parts) < 2:
        return False
    # First and last parts must be at least 2 chars
    if len(parts[0]) < 2 or len(parts[-1]) < 2:
        return False
    # Reject if the first (capitalised) word is a document stopword
    if parts[0].lower() in _NAME_STOPWORDS:
        return False
    # Reject if ALL non-particle parts are stopwords
    real_parts = [p for p in parts if p.lower() not in _NAME_PARTICLES]
    if all(p.lower() in _NAME_STOPWORDS for p in real_parts):
        return False
    # Reject strings that are all-uppercase (acronyms, e.g. "CVR NR")
    if all(p.isupper() and len(p) > 1 for p in parts):
        return False
    # Require at least the first word to look like a name (starts uppercase, has lowercase)
    if not re.search(r'[a-zæøå]', parts[0]):
        return False
    return True


# Words that strongly suggest a nearby 10-digit sequence is a CPR number.
# Used by cpr_context_boost() to raise the risk score.
CPR_CONTEXT_WORDS = re.compile(
    r"\b(?:cpr|personnummer|person[\-\s]?nr|cpr[\-\s]?nr|"
    r"f\.?d\.?t\.?|fodt|fødselsdato|fdato|"
    r"born|date\s+of\s+birth|dob|"
    r"civil\s*registration|NemID|MitID)\b",
    re.IGNORECASE | re.UNICODE,
)


# ── False-positive exclusion: invoice / document-number context ───────────────
# If any of these words appear within ~120 characters of a candidate match,
# it is very likely an invoice number, order number, or part number — not a CPR.
CPR_FALSE_POSITIVE_WORDS = re.compile(
    r"\b(?:"
    # Invoice / order documents
    r"faktura(?:nr|nummer)?|invoice|invoicenr|invno|inv\.?\s*no"
    r"|ordre(?:nr|nummer)?|order(?:nr|number)?"
    r"|rekvisition|requisition"
    r"|tilbud(?:snr|snummer)?"
    r"|kvittering"
    r"|kreditnota|credit\s*note"
    # Item / part / product references
    r"|varenr|vare(?:nummer)?"
    r"|art(?:ikel)?(?:nr|nummer|no)?"
    r"|item\s*(?:nr|no|number|#)?"
    r"|part\s*(?:nr|no|number|#)?"
    r"|produkt(?:nr|nummer)?"
    r"|model(?:nr|number)?"
    r"|serial\s*(?:nr|no|number)?"
    r"|serie(?:nr|nummer)?"
    r"|lot\s*(?:nr|no|number)?"
    r"|batch\s*(?:nr|no|number)?"
    # Reference / document codes
    r"|referencenr|ref(?:erence)?\.?\s*(?:nr|no|number)?"
    r"|sagsnr|sags(?:nummer)?"
    r"|doc(?:ument)?\s*(?:nr|no|number|#)?"
    r"|bilag(?:snr|snummer)?"
    r"|bogf(?:øring)?"
    r"|kontonr|konto(?:nummer)?"
    r"|ean\s*(?:nr|no|number)?"
    r"|gln"
    r"|p(?:urchase)?\s*order"
    r"|po\s*(?:nr|no|number)?"
    r"|so\s*(?:nr|no)?"           # sales order
    # Typical invoice line columns
    r"|antal|quantity|qty"
    r"|stk\.|pcs\.|units?"
    r"|enhedspris|unit\s*price"
    r"|rabat|discount"
    r"|moms|vat|tax"
    r"|subtotal|i\s*alt|total\s*(?:ekskl|inkl)"
    r")\b",
    re.IGNORECASE,
)

# Characters that, if appearing immediately before the 10-digit match,
# indicate it's embedded in a longer document/product code — not a CPR.
# e.g. "REF-250312-4821", "ART250312-4821", "V250312-4821"
_CPR_PREFIX_NOISE = re.compile(r"[A-Za-z0-9]$")


def _is_false_positive(text: str, match_start: int, match_end: int,
                        window: int = 120) -> bool:
    """
    Return True if the 10-digit candidate is almost certainly NOT a CPR number.

    Two checks:
    1. Invoice/order/part-number keyword within `window` chars of the match.
    2. The character immediately preceding the match is alphanumeric
       (suggests the number is part of a product or reference code).
    """
    # Check 1 — surrounding keyword context
    lo = max(0, match_start - window)
    hi = min(len(text), match_end + window)
    if CPR_FALSE_POSITIVE_WORDS.search(text[lo:hi]):
        return True

    # Check 2 — prefix character (letter or digit immediately before match)
    if match_start > 0 and _CPR_PREFIX_NOISE.search(text[match_start - 1]):
        return True

    return False


def cpr_context_boost(text: str, cpr_match_start: int, cpr_match_end: int,
                      window: int = 80) -> bool:
    """
    Return True if a CPR-context keyword appears within `window` characters
    of the match — used to boost risk score for contextually confirmed CPRs.
    """
    lo = max(0, cpr_match_start - window)
    hi = min(len(text), cpr_match_end + window)
    return bool(CPR_CONTEXT_WORDS.search(text[lo:hi]))

# ── NER entity types to redact ─────────────────────────────────────────────────
# spaCy label → human label. Covers Danish (da_core_news) and multilingual models.
NER_REDACT_LABELS = {
    "PER":    "NAME",       # da_core_news
    "PERSON": "NAME",       # en_core_web / xx_ent_wiki
    "LOC":    "ADDRESS",    # da_core_news locations (includes addresses)
    "GPE":    "ADDRESS",    # geopolitical entity (en/xx models)
    "FAC":    "ADDRESS",    # facilities / addresses
    "ORG":    "ORG",        # organisations (optional — included for thoroughness)
}


# ── General helpers ───────────────────────────────────────────────────────────

# Official CPR mod-11 weights applied to digits 1-10
_MOD11_WEIGHTS = (4, 3, 2, 7, 6, 5, 4, 3, 2, 1)


def _passes_mod11(dd: str, mm: str, yy: str, seq: str) -> bool:
    """
    Return True if the 10-digit CPR passes the official Danish mod-11 checksum.

    Note: Denmark stopped issuing mod-11-valid CPR numbers around 2007 when the
    number space was exhausted.  Post-2007 births have CPR numbers that do NOT
    pass this check — so mod-11 failure does NOT prove a number is fake.
    Use this as a CONFIDENCE signal, not a hard gate.
    """
    digits = [int(c) for c in (dd + mm + yy + seq)]
    return sum(d * w for d, w in zip(digits, _MOD11_WEIGHTS)) % 11 == 0


def is_valid_cpr(dd, mm, yy, seq):
    """
    Validate a candidate CPR number.

    Returns:
      (False, False)  — fails date/range/century check — not a CPR
      (True, True)    — passes date check AND mod-11 checksum (high confidence)
      (True, False)   — passes date only, not mod-11 (post-2007 numbers are
                         legitimately valid but fail mod-11 — require context)

    Rules applied:
    - Month must be 01-12
    - Day must be 01-31 (or 41-71 for protected numbers where day += 40)
    - The date DDMMYY must be a real calendar date (e.g. 310200 is invalid)
    - Sequence (last 4 digits) must not be 0000
    - Century digit (first digit of seq) must be consistent with the year
      according to the official Danish CPR century table

    CPR century digit rules (7th digit → birth century):
      0-3 → always 1900s
      4   → 1937-1999 → 1900s ; 2000-2036 → 2000s
      5-8 → 1858-1899 → 1800s ; 1900-1999 → 1900s  (effectively 1900s for modern docs)
      9   → 1937-1999 → 1900s ; 2000-2036 → 2000s
    """
    try:
        d, m, y, s = int(dd), int(mm), int(yy), int(seq)
    except ValueError:
        return False, False

    # Reject all-zero sequence
    if s == 0:
        return False, False

    # Normalise protected numbers (day += 40)
    d_norm = d - 40 if d > 40 else d

    # Basic range checks
    if not (1 <= m <= 12):
        return False, False
    if not (1 <= d_norm <= 31):
        return False, False

    # Determine century from 7th digit (first digit of seq)
    c7 = s // 1000
    if c7 in (0, 1, 2, 3):
        century = 1900
    elif c7 == 4:
        century = 2000 if y <= 36 else 1900
    elif c7 in (5, 6, 7, 8):
        century = 1900
    elif c7 == 9:
        century = 2000 if y <= 36 else 1900
    else:
        return False, False

    # Validate actual calendar date (catches 310200, 290200 in non-leap years, etc.)
    try:
        date(century + y, m, d_norm)
    except ValueError:
        return False, False

    return True, _passes_mod11(dd, mm, yy, seq)


def is_text_page(page) -> bool:
    text = page.extract_text() or ""
    return len(text.replace(" ", "").replace("\n", "")) >= 20


def ocr_page(image, lang: str) -> str:
    config = "--oem 3 --psm 3"
    return pytesseract.image_to_string(image, lang=lang, config=config)


def extract_matches(text: str, page_num: int, source: str):
    """Extract CPR numbers and dates. Returns (cprs, dates)."""
    cprs, dates = [], []
    for m in CPR_PATTERN.finditer(text):
        dd, mm, yy, seq = m.groups()
        date_ok, mod11_ok = is_valid_cpr(dd, mm, yy, seq)
        if not date_ok:
            continue
        if _is_false_positive(text, m.start(), m.end()):
            continue
        ctx = cpr_context_boost(text, m.start(), m.end())
        # Gate: require mod-11 OR explicit CPR context keyword.
        # This rejects ~91% of random date-valid numbers (invoice/part numbers)
        # while keeping real post-2007 CPRs that appear with explicit labels.
        if not mod11_ok and not ctx:
            continue
        cprs.append({"page": page_num, "raw": m.group(0),
                     "formatted": f"{dd}{mm}{yy}-{seq}", "source": source,
                     "context_confirmed": ctx,
                     "mod11": mod11_ok})
    for pattern, fmt in DATE_PATTERNS:
        for m in pattern.finditer(text):
            dates.append({"page": page_num, "raw": m.group(0),
                          "format": fmt, "source": source})
    return cprs, dates


def dedup_dates(dates):
    seen, result = set(), []
    for d in dates:
        key = (d["page"], d["raw"].strip())
        if key not in seen:
            seen.add(key)
            result.append(d)
    return result


def count_pii_types(text: str, use_ner: bool = True) -> dict:
    """
    Count all PII types in text.
    Returns e.g. {"PHONE": 2, "EMAIL": 1, "IBAN": 0, "BANK_ACCOUNT": 1, "NAME": 3, "ADDRESS": 1, "ORG": 2}.
    NER (NAME/ADDRESS/ORG) is run when use_ner=True and the spaCy model is loaded.
    """
    counts: dict[str, int] = {
        "PHONE": 0, "EMAIL": 0, "IBAN": 0, "BANK_ACCOUNT": 0,
        "NAME": 0, "ADDRESS": 0, "ORG": 0,
    }

    for m in PHONE_PATTERN.finditer(text):
        raw = m.group(0).replace(" ", "").replace("-", "").lstrip("+")
        digits = re.sub(r"\D", "", raw)
        if len(digits) in (8, 10, 11):
            counts["PHONE"] += 1

    for _ in EMAIL_PATTERN.finditer(text):
        counts["EMAIL"] += 1

    for _ in IBAN_PATTERN.finditer(text):
        counts["IBAN"] += 1

    for m in REG_KONTO_PATTERN.finditer(text):
        reg, acct = m.group(1), m.group(2)
        if 1 <= int(reg) <= 9999 and len(acct) >= 6:
            counts["BANK_ACCOUNT"] += 1

    # NER-based counts — only run if model is loaded and text is non-trivial
    if use_ner and len(text.strip()) > 20:
        nlp = load_nlp()
        if nlp:
            NER_LIMIT = 20_000
            for chunk_start in range(0, min(len(text), NER_LIMIT * 10), NER_LIMIT):
                chunk = text[chunk_start:chunk_start + NER_LIMIT]
                if not chunk.strip():
                    continue
                doc = nlp(chunk)
                for ent in doc.ents:
                    mapped = NER_REDACT_LABELS.get(ent.label_)
                    if mapped in counts:
                        counts[mapped] += 1

    return counts


# ── Date parsing (for --older-than) ──────────────────────────────────────────

MONTH_DA = {"januar":1,"februar":2,"marts":3,"april":4,"maj":5,"juni":6,
            "juli":7,"august":8,"september":9,"oktober":10,"november":11,"december":12}
MONTH_EN = {"january":1,"february":2,"march":3,"april":4,"may":5,"june":6,
            "july":7,"august":8,"september":9,"october":10,"november":11,"december":12}

def parse_date(raw: str, fmt: str):
    raw = raw.strip()
    try:
        if fmt == "ISO YYYY-MM-DD":
            return datetime.strptime(raw, "%Y-%m-%d").date()
        if fmt in ("DD.MM.YYYY", "DD.MM.YY"):
            for sep in ".-/":
                try:
                    d, m, y = raw.split(sep)
                    y = int(y)
                    if fmt == "DD.MM.YY":
                        y += 2000 if y <= 30 else 1900
                    return date(y, int(m), int(d))
                except Exception:
                    pass
        if fmt == "D. maaned YYYY":
            mo = re.match(r"(\d{1,2})\.\s*(\w+)\s+(\d{4})", raw, re.IGNORECASE)
            if mo:
                d, mon, y = mo.groups()
                mn = MONTH_DA.get(mon.lower())
                if mn: return date(int(y), mn, int(d))
        if fmt == "D Month YYYY":
            mo = re.match(r"(\d{1,2})\s+(\w+)\s+(\d{4})", raw, re.IGNORECASE)
            if mo:
                d, mon, y = mo.groups()
                mn = MONTH_EN.get(mon.lower())
                if mn: return date(int(y), mn, int(d))
        if fmt == "Month D, YYYY":
            mo = re.match(r"(\w+)\s+(\d{1,2}),?\s+(\d{4})", raw, re.IGNORECASE)
            if mo:
                mon, d, y = mo.groups()
                mn = MONTH_EN.get(mon.lower())
                if mn: return date(int(y), mn, int(d))
    except Exception:
        pass
    return None


def older_than(d, days: int) -> bool:
    return d <= date.today() - timedelta(days=days)


def build_flagged_list(all_results, min_age_days):
    flagged = []
    for path, results in all_results:
        if not results["cprs"]:
            continue
        old_dates = []
        for hit in results["dates"]:
            d = parse_date(hit["raw"], hit["format"])
            if d and older_than(d, min_age_days):
                old_dates.append((d, hit["raw"], hit["page"]))
        if old_dates:
            old_dates.sort(key=lambda x: x[0])
            flagged.append({"path": path, "cpr_count": len(results["cprs"]),
                            "oldest_date": old_dates[0], "old_dates": old_dates})
    return flagged


def print_flagged(flagged, min_age_days):
    print(f"\n{'#'*62}")
    print(f"  FILES WITH CPR + DATES OLDER THAN {min_age_days} DAYS: {len(flagged)}")
    print(f"{'#'*62}")
    if not flagged:
        print("  None found.\n")
        return
    for i, entry in enumerate(flagged, 1):
        oldest_d, oldest_raw, oldest_page = entry["oldest_date"]
        print(f"\n  {i}. {entry['path']}")
        print(f"     CPR numbers : {entry['cpr_count']}")
        print(f"     Oldest date : {oldest_raw}  ({oldest_d.isoformat()}, page {oldest_page})")
        for d, raw, pg in entry["old_dates"][1:4]:
            print(f"                   {raw}  ({d.isoformat()}, page {pg})")
        if len(entry["old_dates"]) > 4:
            print(f"                   ... and {len(entry['old_dates'])-4} more")
    print()


# ── PII detection: text spans ─────────────────────────────────────────────────

def find_pii_spans_in_text(text: str, use_ner: bool = True) -> list[tuple[int, int, str]]:
    """
    Return list of (start, end, label) for all PII found in text.
    Covers: CPR, phone, email, and (if use_ner) NER entities.
    """
    spans = []

    # CPR
    for m in CPR_PATTERN.finditer(text):
        dd, mm, yy, seq = m.groups()
        date_ok, mod11_ok = is_valid_cpr(dd, mm, yy, seq)
        if not date_ok:
            continue
        if _is_false_positive(text, m.start(), m.end()):
            continue
        ctx = cpr_context_boost(text, m.start(), m.end())
        if not mod11_ok and not ctx:
            continue
        spans.append((m.start(), m.end(), "CPR"))

    # Phone
    for m in PHONE_PATTERN.finditer(text):
        raw = m.group(0).replace(" ", "").replace("-", "").lstrip("+")
        digits = re.sub(r"\D", "", raw)
        if len(digits) in (8, 10, 11):  # 8=DK, 10/11=with country code
            spans.append((m.start(), m.end(), "PHONE"))

    # Email
    for m in EMAIL_PATTERN.finditer(text):
        spans.append((m.start(), m.end(), "EMAIL"))

    # Danish IBAN
    for m in IBAN_PATTERN.finditer(text):
        spans.append((m.start(), m.end(), "IBAN"))

    # Danish REG/Konto bank account  (only when plausibly formatted as account)
    for m in REG_KONTO_PATTERN.finditer(text):
        reg, acct = m.group(1), m.group(2)
        if 1 <= int(reg) <= 9999 and len(acct) >= 6:
            spans.append((m.start(), m.end(), "BANK_ACCOUNT"))

    # Danish postal addresses
    for m in DANISH_ADDRESS_PATTERN.finditer(text):
        # Only include if the match is long enough to avoid false positives
        if len(m.group(0).strip()) >= 8:
            spans.append((m.start(), m.end(), "ADDRESS"))

    # Regex-based name detection — catches isolated "Firstname Lastname" cells
    # where spaCy has no surrounding context to work from.
    if use_ner:
        for m in DANISH_NAME_PATTERN.finditer(text):
            if _is_name_match(m):
                spans.append((m.start(), m.end(), "NAME"))

    # NER (names, addresses, orgs)
    # Cap at 20 000 chars per call — spaCy NER is O(n) but dense tabular text
    # (e.g. Excel-converted PDFs) can have thousands of tokens per page and stall.
    #
    # Context boosting: spaCy needs sentence context to recognise isolated names.
    # For short text (< 80 chars, e.g. a single cell or line) we prepend a label
    # so the model sees "Navn: Peter Hansen" instead of bare "Peter Hansen".
    # Matches are shifted back by the prefix length before being recorded.
    if use_ner:
        nlp = load_nlp()
        if nlp:
            NER_LIMIT = 20_000
            PREFIX = "Navn: "
            PLEN   = len(PREFIX)
            # Only inject prefix for short/isolated text
            if len(text.strip()) < 80:
                ner_input  = PREFIX + text
                ner_offset = -PLEN
            else:
                ner_input  = text
                ner_offset = 0
            for chunk_start in range(0, min(len(ner_input), NER_LIMIT * 10), NER_LIMIT):
                chunk = ner_input[chunk_start:chunk_start + NER_LIMIT]
                if not chunk.strip():
                    continue
                doc = nlp(chunk)
                for ent in doc.ents:
                    if ent.label_ in NER_REDACT_LABELS:
                        s = chunk_start + ent.start_char + ner_offset
                        e = chunk_start + ent.end_char   + ner_offset
                        if e <= 0:   # entity was entirely within the prefix
                            continue
                        spans.append((max(s, 0), e, NER_REDACT_LABELS[ent.label_]))

    # Merge overlapping spans
    spans.sort()
    merged = []
    for start, end, label in spans:
        if merged and start <= merged[-1][1]:
            prev_s, prev_e, prev_l = merged[-1]
            merged[-1] = (prev_s, max(prev_e, end), prev_l)
        else:
            merged.append((start, end, label))

    return merged


# ── Bounding box finders ──────────────────────────────────────────────────────

def find_pii_char_bboxes(page, use_ner: bool = True) -> list[tuple[float, float, float, float, str]]:
    """
    Return (x0, top, x1, bottom, label) for all PII on a text-based pdfplumber page.

    Uses extract_words() for bbox lookup, but extract_text() for the NER text so
    that spaCy sees newlines between lines — critical for name recognition. Without
    newlines, names from adjacent rows run together and spaCy misses them.
    """
    words = page.extract_words(keep_blank_chars=False, x_tolerance=3, y_tolerance=3)
    if not words:
        return []

    # Build a word-span index for bbox lookup (space-separated, no newlines)
    word_text = ""
    word_spans = []
    for w in words:
        ws = len(word_text)
        word_text += w["text"]
        word_spans.append((ws, len(word_text), w))
        word_text += " "

    # For PII/NER detection use extract_text() which preserves newlines between
    # lines — spaCy needs sentence structure to reliably recognise names.
    ner_text = page.extract_text() or word_text

    spans = find_pii_spans_in_text(ner_text, use_ner=use_ner)

    bboxes = []
    PAD = 1
    for span_start, span_end, label in spans:
        # The matched span is in ner_text coordinates. Map to word_text by
        # extracting the matched surface form and fuzzy-searching in word_text.
        matched_surface = ner_text[span_start:span_end].strip()
        if not matched_surface:
            continue

        # Search for the token sequence in the word list
        # Split matched surface into tokens (same split as extract_words uses)
        import re as _re
        tokens = _re.split(r'\s+', matched_surface)
        tokens = [t for t in tokens if t]

        hit_words = []
        if tokens:
            # Find the first word that starts with the first token
            for i, (ws, we, w) in enumerate(word_spans):
                if w["text"].startswith(tokens[0]) or tokens[0].startswith(w["text"]):
                    # Try to match the full token sequence from here
                    candidate = word_spans[i:i + len(tokens)]
                    if len(candidate) == len(tokens):
                        hit_words = [cw for (_, _, cw) in candidate]
                        break
                    # Partial match — just take as many words as match
                    hit_words = [cw for (_, _, cw) in candidate]
                    break

        if not hit_words:
            # Fallback: find words whose text overlaps with matched_surface tokens
            surface_lower = matched_surface.lower()
            hit_words = [w for (_, _, w) in word_spans
                         if w["text"].lower() in surface_lower
                         or surface_lower in w["text"].lower()]

        if not hit_words:
            continue

        bboxes.append((
            min(w["x0"]     for w in hit_words) - PAD,
            min(w["top"]    for w in hit_words) - PAD,
            max(w["x1"]     for w in hit_words) + PAD,
            max(w["bottom"] for w in hit_words) + PAD,
            label,
        ))
    return bboxes
def find_cpr_char_bboxes(page):
    """
    CPR-only version for --mask (no NER).

    Uses extract_words() to build the text string — the same tokenisation that
    extract_text() uses during scanning. Raw page.chars iteration fails on
    Excel-converted PDFs where chars have no inter-word spacing or are stored
    in a different order than reading order, causing CPR patterns to either
    not match or match at the wrong offsets.

    Strategy:
      1. Build a word list with bboxes via extract_words().
      2. Concatenate words (space-separated) and run CPR_PATTERN on that string.
      3. For each match, find which word(s) it falls in and union their bboxes.
         Add a small padding so the black box covers the full glyph.
    """
    words = page.extract_words(keep_blank_chars=False, x_tolerance=3, y_tolerance=3)
    if not words:
        return []

    # Build concatenated text and track each word's start offset
    full_text = ""
    word_spans = []  # (start_offset, end_offset, word_dict)
    for w in words:
        start = len(full_text)
        full_text += w["text"]
        word_spans.append((start, len(full_text), w))
        full_text += " "  # space separator between words

    bboxes = []
    for m in CPR_PATTERN.finditer(full_text):
        dd, mm, yy, seq = m.groups()
        date_ok, mod11_ok = is_valid_cpr(dd, mm, yy, seq)
        if not date_ok:
            continue
        if _is_false_positive(full_text, m.start(), m.end()):
            continue
        ctx = cpr_context_boost(full_text, m.start(), m.end())
        if not mod11_ok and not ctx:
            continue
        ms, me = m.start(), m.end()
        # Collect all words that overlap this match span
        hit_words = [w for (ws, we, w) in word_spans if ws < me and we > ms]
        if not hit_words:
            continue
        PAD = 1  # points of padding around the glyph
        bboxes.append((
            min(w["x0"]    for w in hit_words) - PAD,
            min(w["top"]   for w in hit_words) - PAD,
            max(w["x1"]    for w in hit_words) + PAD,
            max(w["bottom"]for w in hit_words) + PAD,
        ))
    return bboxes


def find_cpr_image_bboxes(image, lang: str):
    """CPR-only image bboxes for --mask."""
    raw_bboxes = find_pii_image_bboxes(image, lang, use_ner=False)
    return [(l, t, r, b) for (l, t, r, b, lbl) in raw_bboxes if lbl == "CPR"]


# ── Drawing helpers ───────────────────────────────────────────────────────────

def build_redaction_overlay(page_width, page_height, bboxes_pdfplumber) -> bytes:
    """Build a PDF overlay with black boxes. bboxes: (x0, top, x1, bottom[, label])."""
    buf = io.BytesIO()
    c = rl_canvas.Canvas(buf, pagesize=(page_width, page_height))
    c.setFillColor(rl_black)
    c.setStrokeColor(rl_black)
    pad = 1.5
    for bbox in bboxes_pdfplumber:
        x0, top, x1, bot = bbox[:4]
        rl_y = page_height - bot - pad
        rl_h = (bot - top) + pad * 2
        c.rect(x0 - pad, rl_y, (x1 - x0) + pad * 2, rl_h, fill=1, stroke=0)
    c.save()
    buf.seek(0)
    return buf.read()


def apply_overlay_to_page(writer, reader_page, bboxes):
    page_width  = float(reader_page.mediabox.width)
    page_height = float(reader_page.mediabox.height)
    overlay_bytes  = build_redaction_overlay(page_width, page_height, bboxes)
    overlay_page   = PdfReader(io.BytesIO(overlay_bytes)).pages[0]
    reader_page.merge_page(overlay_page)
    writer.add_page(reader_page)


def redact_image(image, bboxes_px):
    """Paint black rectangles over pixel bboxes in a PIL image."""
    from PIL import ImageDraw
    img = image.copy()
    draw = ImageDraw.Draw(img)
    for bbox in bboxes_px:
        left, top, right, bottom = bbox[:4]
        draw.rectangle([left, top, right, bottom], fill="black")
    return img


def image_to_pdf_page(image, dpi=300) -> bytes:
    buf = io.BytesIO()
    image.convert("RGB").save(buf, format="PDF", resolution=dpi)
    buf.seek(0)
    return buf.read()


# ── Secure PDF redaction (PyMuPDF) ───────────────────────────────────────────

def redact_pdf_secure(input_path: Path, output_path: Path, results: dict,
                      force_ocr: bool, lang: str, dpi: int, poppler_path,
                      use_ner: bool = False) -> "int | bool":
    """
    Physically-secure PDF redaction using PyMuPDF (fitz).

    Unlike the reportlab overlay approach, PyMuPDF:
      1. Draws opaque redaction annotations over the target character bboxes.
      2. Calls page.apply_redactions() which physically REMOVES the underlying
         text/image data — not just paints over it.
      3. Saves with garbage collection and compression to strip orphaned objects.

    This means a user cannot recover the redacted text by:
      - Selecting text under the black box in a viewer
      - Extracting the PDF text layer programmatically
      - Inspecting raw PDF object streams

    Falls back to the reportlab overlay method if PyMuPDF is not installed.
    """
    if not PYMUPDF_AVAILABLE:
        return redact_pdf(input_path, output_path, results,
                          force_ocr, lang, dpi, poppler_path, use_ner)

    page_methods = results["page_methods"]

    images = None
    ocr_pages = [p for p, m in page_methods.items() if m == "ocr"]
    if ocr_pages and OCR_AVAILABLE:
        images = convert_from_path(str(input_path), dpi=dpi, poppler_path=poppler_path)

    total = 0
    doc = _fitz.open(str(input_path))

    with pdfplumber.open(input_path) as plumb_pdf:
        for page_num, plumb_page in enumerate(plumb_pdf.pages, start=1):
            method    = page_methods.get(page_num, "text")
            fitz_page = doc[page_num - 1]

            # Get bboxes in pdfplumber coordinates (origin top-left, y increases down)
            if method == "text":
                bboxes = (find_pii_char_bboxes(plumb_page, use_ner=use_ner)
                          if use_ner else find_cpr_char_bboxes(plumb_page))
            elif method == "ocr" and images is not None:
                img    = images[page_num - 1]
                bboxes = (find_pii_image_bboxes(img, lang, use_ner=use_ner)
                          if use_ner else find_cpr_image_bboxes(img, lang))
            else:
                bboxes = []

            # pdfplumber char coords: origin top-left of CropBox, y increases DOWN.
            # fitz Rect coords:       origin top-left of MediaBox, y increases DOWN.
            # Both already have y=0 at the top — no flip needed.
            # Add the CropBox offset so boxes land correctly when CropBox != MediaBox.
            cb = fitz_page.cropbox
            mb = fitz_page.mediabox
            crop_x0 = cb.x0 - mb.x0
            crop_y0 = cb.y0 - mb.y0

            for bbox in bboxes:
                x0, top, x1, bottom = bbox[:4]
                rect = _fitz.Rect(
                    x0     + crop_x0,
                    top    + crop_y0,
                    x1     + crop_x0,
                    bottom + crop_y0,
                )
                annot = fitz_page.add_redact_annot(rect, fill=(0, 0, 0))
                _ = annot  # silence linter

            # Apply redactions — physically removes text/image data under rects
            # PDF_REDACT_IMAGE_REMOVE / PDF_REDACT_LINE_ART_REMOVE were added in
            # PyMuPDF 1.22; fall back to their integer values (2) on older builds.
            _img_flag  = getattr(_fitz, "PDF_REDACT_IMAGE_REMOVE",    2)
            _art_flag  = getattr(_fitz, "PDF_REDACT_LINE_ART_REMOVE", 2)
            fitz_page.apply_redactions(images=_img_flag, graphics=_art_flag)
            total += len(bboxes)

    # Save with full garbage collection (removes orphaned objects/streams)
    doc.save(
        str(output_path),
        garbage=4,          # maximum GC: also removes unused xref entries
        deflate=True,       # compress streams
        clean=True,         # sanitise content streams
        linear=False,
    )
    doc.close()
    return total


# ── Generic redact-PDF engine (reportlab overlay — visual only) ───────────────

def redact_pdf(input_path: Path, output_path: Path, results: dict,
               force_ocr: bool, lang: str, dpi: int, poppler_path,
               use_ner: bool = False) -> int | bool:
    """
    Write a redacted PDF to output_path.
    If use_ner=False: CPR only (--mask).
    If use_ner=True:  all PII (--anonymise).
    Returns count of redacted regions, or False on error.
    """
    if not MASK_AVAILABLE:
        print("  Requires: pip install pypdf reportlab")
        return False

    page_methods = results["page_methods"]
    reader = PdfReader(str(input_path))
    writer = PdfWriter()

    images = None
    ocr_pages = [p for p, m in page_methods.items() if m == "ocr"]
    if ocr_pages and OCR_AVAILABLE:
        images = convert_from_path(str(input_path), dpi=dpi, poppler_path=poppler_path)

    total = 0
    with pdfplumber.open(input_path) as plumb_pdf:
        for page_num, plumb_page in enumerate(plumb_pdf.pages, start=1):
            method = page_methods.get(page_num, "text")
            reader_page = reader.pages[page_num - 1]

            if method == "text":
                bboxes = (find_pii_char_bboxes(plumb_page, use_ner=use_ner)
                          if use_ner else find_cpr_char_bboxes(plumb_page))
                if bboxes:
                    apply_overlay_to_page(writer, reader_page, bboxes)
                    total += len(bboxes)
                else:
                    writer.add_page(reader_page)

            elif method == "ocr" and images is not None:
                img = images[page_num - 1]
                bboxes = (find_pii_image_bboxes(img, lang, use_ner=use_ner)
                          if use_ner else find_cpr_image_bboxes(img, lang))
                if bboxes:
                    writer.add_page(
                        PdfReader(io.BytesIO(
                            image_to_pdf_page(redact_image(img, bboxes), dpi)
                        )).pages[0]
                    )
                    total += len(bboxes)
                else:
                    writer.add_page(reader_page)
            else:
                writer.add_page(reader_page)

    with open(output_path, "wb") as f:
        writer.write(f)
    return total



# ── Word document support ─────────────────────────────────────────────────────

def _iter_docx_runs(doc):
    """Yield every run in a docx Document: body, tables, headers, footers."""
    def _from_paragraphs(paragraphs):
        for para in paragraphs:
            for run in para.runs:
                yield run

    yield from _from_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from _from_paragraphs(cell.paragraphs)
    for section in doc.sections:
        for hf in [section.header, section.footer,
                   section.even_page_header, section.even_page_footer,
                   section.first_page_header, section.first_page_footer]:
            try:
                yield from _from_paragraphs(hf.paragraphs)
            except Exception:
                pass


def scan_docx(docx_path: Path) -> dict:
    """
    Scan a .docx file for CPR numbers and dates.
    Returns the same results dict shape as scan_document(), plus internal
    _doc / _run_map / _full_text keys used by redact_docx().
    """
    if not DOCX_OK:
        print("  .docx support requires: pip install python-docx")
        return {"cprs": [], "dates": [], "page_methods": {1: "docx"},
                "_doc": None, "_run_map": [], "_full_text": ""}

    doc = DocxDocument(str(docx_path))

    # Build full text + run map (global_start, global_end, run)
    full_text = ""
    run_map = []
    for run in _iter_docx_runs(doc):
        if run.text:
            start = len(full_text)
            full_text += run.text
            run_map.append((start, len(full_text), run))

    cprs, dates = extract_matches(full_text, 1, "docx")
    return {
        "cprs": cprs,
        "dates": dates,
        "page_methods": {1: "docx"},
        "_full_text": full_text,
        "_run_map": run_map,
        "_doc": doc,
    }


def _redact_runs(run_map: list, spans: list):
    """
    Replace characters in the given spans with block characters (█).
    Modifies runs in-place.
    spans: list of (start, end, label) in full_text coordinates.
    """
    if not spans:
        return

    # Build char → (run, index_within_run) lookup
    char_owner = []   # index = position in full_text, value = (run, char_pos_in_run)
    for (gs, ge, run) in run_map:
        for i in range(ge - gs):
            char_owner.append((run, i))

    # Apply redactions (process in reverse so earlier spans aren't shifted)
    for span_start, span_end, _label in sorted(spans, key=lambda s: s[0], reverse=True):
        # Group by run
        by_run = {}
        for pos in range(span_start, min(span_end, len(char_owner))):
            run_obj, char_pos = char_owner[pos]
            rid = id(run_obj)
            if rid not in by_run:
                by_run[rid] = {"run": run_obj, "positions": []}
            by_run[rid]["positions"].append(char_pos)
        for entry in by_run.values():
            run_obj = entry["run"]
            chars = list(run_obj.text)
            for p in entry["positions"]:
                if p < len(chars):
                    chars[p] = "█"
            run_obj.text = "".join(chars)


def redact_docx(input_path: Path, output_path: Path, results: dict,
                use_ner: bool = False) -> int:
    """
    Write a redacted copy of a .docx.
    use_ner=False → CPR only; use_ner=True → all PII.
    Returns number of spans redacted.
    """
    doc      = results.get("_doc")
    run_map  = results.get("_run_map", [])
    text     = results.get("_full_text", "")

    if doc is None:
        return 0

    spans = find_pii_spans_in_text(text, use_ner=use_ner)

    # If CPR-only, filter to CPR spans
    if not use_ner:
        spans = [(s, e, l) for s, e, l in spans if l == "CPR"]

    _redact_runs(run_map, spans)
    doc.save(str(output_path))
    return len(spans)


def print_docx_results(docx_path: Path, results: dict):
    cprs  = results["cprs"]
    dates = results["dates"]
    print(f"\n{'='*62}")
    print(f"File : {docx_path}  [Word document]")
    print(f"{'='*62}")
    print(f"\n  CPR Numbers found: {len(cprs)}")
    if cprs:
        for hit in cprs:
            print(f"    {hit['formatted']:<16}  (raw: \"{hit['raw']}\")")
    else:
        print("    None found.")
    print(f"\n  Dates found: {len(dates)}")
    if dates:
        for hit in dates:
            print(f"    {hit['raw']:<28}  [{hit['format']}]")
    else:
        print("    None found.")
    print()


# ── Logging ───────────────────────────────────────────────────────────────────

# Module-level logger — handlers are added in main() based on --log argument
logger = logging.getLogger("scanner")
logger.setLevel(logging.DEBUG)

_log_records: list[dict] = []   # in-memory log, flushed to JSON at end

def _log(level: str, path: Path | None, event: str, **kwargs):
    """
    Append a structured log record and emit to the logger.
    level: "INFO" | "WARNING" | "ACTION" | "DRY_RUN" | "ERROR"
    """
    record = {
        "time":  datetime.now().isoformat(timespec="seconds"),
        "level": level,
        "file":  str(path) if path else None,
        "event": event,
        **kwargs,
    }
    _log_records.append(record)
    msg = f"[{level}] {path.name if path else ''} — {event}"
    if kwargs:
        extras = "  " + "  ".join(f"{k}={v}" for k, v in kwargs.items())
        msg += extras
    if level == "ERROR":
        logger.error(msg)
    elif level == "WARNING":
        logger.warning(msg)
    else:
        logger.info(msg)


def flush_log(log_path: Path):
    """Write all accumulated log records to a JSON file."""
    with open(log_path, "w", encoding="utf-8") as f:
        json.dump(_log_records, f, ensure_ascii=False, indent=2, default=str)
    print(f"\nLog written to: {log_path}  ({len(_log_records)} records)")



# ── Excel / CSV support ───────────────────────────────────────────────────────

def _cell_text(cell) -> str:
    """Return a string representation of a cell value, or empty string."""
    if cell.value is None:
        return ""
    return str(cell.value)


def scan_xlsx(path: Path) -> dict:
    """
    Scan an .xlsx / .xlsm file for CPR numbers and dates across all sheets.
    Returns results dict compatible with the rest of the pipeline, plus
    _wb (workbook) for use by redact_xlsx().
    Each CPR/date hit carries sheet + row + col in the "page" field
    (formatted as "Sheet!R{row}C{col}").
    """
    if not XLSX_OK:
        print("  .xlsx support requires: pip install openpyxl")
        return {"cprs": [], "dates": [], "page_methods": {1: "xlsx"}, "_wb": None}

    wb = openpyxl.load_workbook(str(path), data_only=True)
    all_cprs, all_dates = [], []

    for sheet in wb.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                val = _cell_text(cell)
                if not val:
                    continue
                location = f"{sheet.title}!R{cell.row}C{cell.column}"
                cprs, dates = extract_matches(val, location, "xlsx")
                all_cprs.extend(cprs)
                all_dates.extend(dates)

    return {
        "cprs": all_cprs,
        "dates": all_dates,
        "page_methods": {1: "xlsx"},
        "_wb": wb,
        "_path": path,
    }


def scan_csv(path: Path) -> dict:
    """
    Scan a .csv file for CPR numbers and dates.
    Returns results dict compatible with the rest of the pipeline.
    """
    import csv as _csv

    all_cprs, all_dates = [], []
    try:
        with open(path, newline="", encoding="utf-8-sig", errors="replace") as f:
            reader = _csv.reader(f)
            for row_num, row in enumerate(reader, start=1):
                for col_num, cell in enumerate(row, start=1):
                    if not cell.strip():
                        continue
                    location = f"R{row_num}C{col_num}"
                    cprs, dates = extract_matches(cell, location, "csv")
                    all_cprs.extend(cprs)
                    all_dates.extend(dates)
    except Exception as e:
        print(f"  Warning: could not read CSV: {e}")

    return {
        "cprs": all_cprs,
        "dates": all_dates,
        "page_methods": {1: "csv"},
        "_wb": None,
        "_path": path,
    }


def scan_text(text: str, source: str = "text") -> dict:
    """
    Scan a plain text string for CPR numbers and dates.
    Returns a results dict compatible with the rest of the pipeline.
    False-positive suppression (invoice/part-number context) is applied
    via extract_matches → extract_cpr_and_dates → _is_false_positive.
    """
    cprs, dates = extract_cpr_and_dates(text, page_num=1, source=source)
    return {
        "cprs": cprs,
        "dates": dates,
        "page_methods": {1: "text"},
    }


def scan_image(path: Path, lang: str = "dan+eng") -> dict:
    """
    OCR an image file and scan the resulting text for CPR numbers.
    Requires Tesseract and pytesseract.
    """
    try:
        import pytesseract as _tess
        from PIL import Image as _PILImage
        img = _PILImage.open(path)
        text = _tess.image_to_string(img, lang=lang, config="--oem 3 --psm 3")
        return scan_text(text, source="image-ocr")
    except ImportError:
        return {"cprs": [], "dates": [], "error": "pytesseract/PIL not available"}
    except Exception as e:
        return {"cprs": [], "dates": [], "error": str(e)}


def redact_xlsx(input_path: Path, output_path: Path, results: dict,
                use_ner: bool = False) -> int:
    """
    Write a redacted copy of an .xlsx file.
    Cells containing PII are overwritten with "████████".
    use_ner=False -> CPR only; use_ner=True -> all PII.
    Returns number of cells redacted.
    """
    wb = results.get("_wb")
    if wb is None:
        return 0

    redacted = 0
    for sheet in wb.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                val = _cell_text(cell)
                if not val:
                    continue
                # Wrap cell in a context sentence so spaCy NER can recognise
                # names that appear in isolation (e.g. a name-only cell has no
                # surrounding text to provide the model with PER entity context).
                PREFIX = "Navn: "
                ctx = PREFIX + val
                raw_spans = find_pii_spans_in_text(ctx, use_ner=use_ner)
                # Shift spans back by prefix length; discard any that start in prefix
                plen = len(PREFIX)
                spans = [(s - plen, e - plen, l) for s, e, l in raw_spans if e > plen]
                spans = [(max(s, 0), e, l) for s, e, l in spans]
                if not use_ner:
                    spans = [(s, e, l) for s, e, l in spans if l == "CPR"]
                if spans:
                    # Replace the whole cell value with redaction marker
                    # (partial in-cell redaction is not reliably possible in xlsx)
                    cell.value = "████████"
                    redacted += 1

    wb.save(str(output_path))
    return redacted


def redact_csv(input_path: Path, output_path: Path, use_ner: bool = False) -> int:
    """
    Write a redacted copy of a .csv file.
    Cells containing PII are overwritten with "████████".
    Returns number of cells redacted.
    """
    import csv as _csv

    rows_out = []
    redacted = 0
    try:
        with open(input_path, newline="", encoding="utf-8-sig", errors="replace") as f:
            reader = _csv.reader(f)
            for row in reader:
                new_row = []
                for cell in row:
                    if cell.strip():
                        PREFIX = "Navn: "
                        ctx = PREFIX + cell
                        plen = len(PREFIX)
                        raw_spans = find_pii_spans_in_text(ctx, use_ner=use_ner)
                        spans = [(max(s - plen, 0), e - plen, l)
                                 for s, e, l in raw_spans if e > plen]
                    else:
                        spans = []
                    if not use_ner:
                        spans = [(s, e, l) for s, e, l in spans if l == "CPR"]
                    if spans:
                        new_row.append("████████")
                        redacted += 1
                    else:
                        new_row.append(cell)
                rows_out.append(new_row)
    except Exception as e:
        print(f"  Warning: could not read CSV for redaction: {e}")
        return 0

    with open(output_path, "w", newline="", encoding="utf-8") as f:
        _csv.writer(f).writerows(rows_out)
    return redacted


def print_xlsx_results(path: Path, results: dict, file_type: str = "xlsx"):
    cprs  = results["cprs"]
    dates = results["dates"]
    label = "Excel spreadsheet" if file_type == "xlsx" else "CSV file"
    print(f"\n{'='*62}")
    print(f"File : {path}  [{label}]")
    print(f"{'='*62}")
    print(f"\n  CPR Numbers found: {len(cprs)}")
    if cprs:
        for hit in cprs:
            print(f"    {hit['page']:<20}  {hit['formatted']:<16}  (raw: \"{hit['raw']}\")")
    else:
        print("    None found.")
    print(f"\n  Dates found: {len(dates)}")
    if dates:
        for hit in dates:
            print(f"    {hit['page']:<20}  {hit['raw']:<28}  [{hit['format']}]")
    else:
        print("    None found.")
    print()


# ── Face detection & pixelation ───────────────────────────────────────────────

# Use both frontal and profile cascades for better coverage
_FACE_CASCADES = None

def _get_face_cascades():
    global _FACE_CASCADES
    if _FACE_CASCADES is not None:
        return _FACE_CASCADES
    cv2, np = _get_cv2()
    if cv2 is None:
        return []

    def _find_cascade(name: str):
        """Try multiple locations to find a Haar cascade XML file."""
        import sys as _sys
        candidates = []
        # 1. PyInstaller bundle — check FIRST so bundle path wins over stale install paths
        if hasattr(_sys, "_MEIPASS"):
            candidates.append(str(Path(_sys._MEIPASS) / "cv2" / "data" / name))
            candidates.append(str(Path(_sys._MEIPASS) / name))
        # 2. cv2.data attribute (standard install / venv)
        try:
            candidates.append(cv2.data.haarcascades + name)
        except Exception:
            pass
        # 3. Relative to cv2 package directory
        try:
            candidates.append(str(Path(cv2.__file__).parent / "data" / name))
        except Exception:
            pass
        # 4. Common system paths
        for base in ["/usr/share/opencv4", "/usr/share/opencv",
                     "/usr/local/share/opencv4", "/usr/local/share/opencv"]:
            candidates.append(str(Path(base) / "haarcascades" / name))

        for p in candidates:
            if p and Path(p).exists():
                c = cv2.CascadeClassifier(p)
                if not c.empty():
                    _face_log(f"  [+] Cascade: {p}")
                    return c
        # Nothing worked — log all paths tried so it shows in the app console
        _face_log(f"  [!] Cascade not found: {name}")
        for p in candidates:
            _face_log(f"      {p}  exists={Path(p).exists()}")
        return None

    cascades = []
    for name in ["haarcascade_frontalface_default.xml", "haarcascade_profileface.xml"]:
        c = _find_cascade(name)
        if c is not None:
            cascades.append(c)

    if not cascades:
        _face_log("  [!] No Haar cascade XML files found — face detection disabled")

    _FACE_CASCADES = cascades
    return cascades


def detect_faces_cv2(img_cv2, min_size: int = 40, neighbors: int = 4,
                     strict: bool = False):
    """
    Detect faces in a BGR cv2 image using Haar cascades (frontal + profile).
    Returns list of (x, y, w, h) in pixel coordinates.

    Parameters
    ----------
    min_size   : minimum face side in pixels
    neighbors  : minNeighbors for detectMultiScale (higher = stricter, fewer detections)
    strict     : unused, kept for API compatibility
    """
    cv2, np = _get_cv2()
    if cv2 is None:
        return []
    gray = cv2.cvtColor(img_cv2, cv2.COLOR_BGR2GRAY)
    # Equalise histogram to improve detection on dark or low-contrast images
    gray = cv2.equalizeHist(gray)

    cascades = _get_face_cascades()
    if not cascades:
        return []

    found = []
    seen = set()

    def _add(x, y, w, h):
        key = (x // 10, y // 10, w // 10, h // 10)
        if key not in seen:
            seen.add(key)
            found.append((x, y, w, h))

    for cascade in cascades:
        for img in [gray, cv2.flip(gray, 1)]:
            faces = cascade.detectMultiScale(
                img, scaleFactor=1.1, minNeighbors=neighbors,
                minSize=(min_size, min_size), flags=cv2.CASCADE_SCALE_IMAGE
            )
            if faces is not None and len(faces) > 0:
                if img is not gray:  # flip back x coords
                    w_img = img.shape[1]
                    faces = [(w_img - x - w, y, w, h) for (x, y, w, h) in faces]
                for face in faces:
                    _add(*face)
    return found


def pixelate_region(img_cv2, x: int, y: int, w: int, h: int, blocks: int = 6):
    """Pixelate a rectangular region in a cv2 image. Returns modified copy.
    Lower blocks = larger pixels = stronger anonymisation.
    A Gaussian blur is applied on top to prevent edge-sharpening attacks.
    """
    cv2, np = _get_cv2()
    out = img_cv2.copy()
    roi = out[y:y+h, x:x+w]
    bw = max(1, w // blocks)
    bh = max(1, h // blocks)
    small = cv2.resize(roi, (bw, bh), interpolation=cv2.INTER_LINEAR)
    pixelated = cv2.resize(small, (w, h), interpolation=cv2.INTER_NEAREST)
    ksize = max(3, (min(w, h) // blocks) | 1)
    pixelated = cv2.GaussianBlur(pixelated, (ksize, ksize), 0)
    out[y:y+h, x:x+w] = pixelated
    return out


def blur_faces_in_image(img_cv2, min_size: int = 30, blocks: int = 6):
    """
    Detect faces and apply pixelation to each. Returns (modified_img, face_count).
    """
    cv2, np = _get_cv2()
    if cv2 is None:
        return img_cv2, 0
    faces = detect_faces_cv2(img_cv2, min_size=min_size)
    out = img_cv2.copy()
    for (x, y, w, h) in faces:
        pad_x = int(w * 0.1)
        pad_y = int(h * 0.1)
        x2 = max(0, x - pad_x)
        y2 = max(0, y - pad_y)
        w2 = min(out.shape[1] - x2, w + pad_x * 2)
        h2 = min(out.shape[0] - y2, h + pad_y * 2)
        out = pixelate_region(out, x2, y2, w2, h2, blocks=blocks)
    return out, len(faces)


def pil_to_cv2(pil_img):
    cv2, np = _get_cv2()
    return cv2.cvtColor(np.array(pil_img.convert("RGB")), cv2.COLOR_RGB2BGR)


def cv2_to_pil(img_cv2):
    cv2, np = _get_cv2()
    from PIL import Image as PILImage
    return PILImage.fromarray(cv2.cvtColor(img_cv2, cv2.COLOR_BGR2RGB))


def cv2_to_bytes(img_cv2, fmt: str = "JPEG") -> bytes:
    """Encode cv2 image to bytes in given format."""
    cv2, np = _get_cv2()
    ext = {"JPEG": ".jpg", "PNG": ".png", "WEBP": ".webp"}.get(fmt.upper(), ".jpg")
    ok, buf = cv2.imencode(ext, img_cv2)
    if not ok:
        raise RuntimeError(f"cv2.imencode failed for format {fmt}")
    return buf.tobytes()


# ── Face blur: standalone image files ─────────────────────────────────────────

def blur_faces_image_file(input_path: Path, output_path: Path,
                           blocks: int = 6) -> int:
    """
    Detect and pixelate faces in a standalone image file.
    Returns number of faces blurred.
    """
    cv2, np = _get_cv2()
    if cv2 is None:
        raise RuntimeError("OpenCV not available")
    img = cv2.imread(str(input_path))
    if img is None:
        raise ValueError(f"Could not read image: {input_path}")
    result, count = blur_faces_in_image(img, blocks=blocks)
    cv2.imwrite(str(output_path), result)
    return count


# ── Face blur: PDF pages ───────────────────────────────────────────────────────

def blur_faces_pdf(input_path: Path, output_path: Path,
                   dpi: int = 150, poppler_path=None,
                   blocks: int = 6) -> int:
    """
    Render each PDF page, detect faces, draw pixelated overlay back onto the
    original page (preserving the text layer), save as new PDF.
    Returns total number of faces blurred across all pages.
    """
    if not OCR_AVAILABLE:
        raise RuntimeError("pdf2image required: pip install pdf2image")

    cv2, np = _get_cv2()
    if cv2 is None:
        raise RuntimeError("OpenCV not available")

    from PIL import Image as PILImage
    images = convert_from_path(str(input_path), dpi=dpi, poppler_path=poppler_path)

    reader  = PdfReader(str(input_path))
    writer  = PdfWriter()
    total_faces = 0

    for page_num, (pil_img, reader_page) in enumerate(zip(images, reader.pages), start=1):
        page_w = float(reader_page.mediabox.width)   # PDF points
        page_h = float(reader_page.mediabox.height)

        img_px_w, img_px_h = pil_img.size
        scale_x = page_w / img_px_w
        scale_y = page_h / img_px_h

        img_cv2 = pil_to_cv2(pil_img)
        _, face_count = blur_faces_in_image(img_cv2, blocks=blocks)

        if face_count == 0:
            writer.add_page(reader_page)
            continue

        # Build a pixelated patch for each face and compose into a reportlab overlay
        faces = detect_faces_cv2(img_cv2)
        buf = io.BytesIO()
        c = rl_canvas.Canvas(buf, pagesize=(page_w, page_h))

        for (x, y, w, h) in faces:
            pad_x = int(w * 0.1)
            pad_y = int(h * 0.1)
            x2, y2 = max(0, x - pad_x), max(0, y - pad_y)
            w2 = min(img_px_w - x2, w + pad_x * 2)
            h2 = min(img_px_h - y2, h + pad_y * 2)

            # Pixelate just this region from the rendered page image
            face_roi = img_cv2[y2:y2+h2, x2:x2+w2]
            bw = max(1, w2 // blocks)
            bh = max(1, h2 // blocks)
            small = cv2.resize(face_roi, (bw, bh), interpolation=cv2.INTER_LINEAR)
            pixelated_roi = cv2.resize(small, (w2, h2), interpolation=cv2.INTER_NEAREST)

            # Convert to PIL for reportlab
            roi_pil = cv2_to_pil(pixelated_roi)
            roi_buf = io.BytesIO()
            roi_pil.save(roi_buf, format="PNG")
            roi_buf.seek(0)

            # PDF coords: reportlab origin is bottom-left; image origin is top-left
            pdf_x  = x2 * scale_x
            pdf_y  = page_h - (y2 + h2) * scale_y
            pdf_w  = w2 * scale_x
            pdf_h  = h2 * scale_y
            c.drawImage(
                __import__("reportlab.lib.utils", fromlist=["ImageReader"]).ImageReader(roi_buf),
                pdf_x, pdf_y, width=pdf_w, height=pdf_h
            )

        c.save()
        buf.seek(0)
        overlay_page = PdfReader(buf).pages[0]
        reader_page.merge_page(overlay_page)
        writer.add_page(reader_page)
        total_faces += face_count

    with open(output_path, "wb") as f:
        writer.write(f)
    return total_faces


# ── Face blur: Word documents ─────────────────────────────────────────────────

def blur_faces_docx(input_path: Path, output_path: Path,
                    blocks: int = 6) -> int:
    """
    Detect and pixelate faces in images embedded in a .docx file.
    Replaces the image part bytes in-place and saves as a new file.
    Returns number of faces blurred.
    """
    if not DOCX_OK:
        raise RuntimeError("python-docx required: pip install python-docx")

    cv2, np = _get_cv2()
    if cv2 is None:
        raise RuntimeError("OpenCV not available")

    import shutil
    from docx import Document
    from docx.oxml.ns import qn
    from docx.enum.shape import WD_INLINE_SHAPE
    from PIL import Image as PILImage

    shutil.copy2(str(input_path), str(output_path))
    doc = Document(str(output_path))
    total_faces = 0

    for shape in doc.inline_shapes:
        try:
            if shape.type != WD_INLINE_SHAPE.PICTURE:
                continue

            blip = shape._inline.graphic.graphicData.pic.blipFill.blip
            rId = blip.embed
            image_part = doc.part.related_parts[rId]

            # Decode image bytes → cv2
            img_data = image_part.blob
            np_arr   = np.frombuffer(img_data, dtype=np.uint8)
            img_cv2  = cv2.imdecode(np_arr, cv2.IMREAD_COLOR)
            if img_cv2 is None:
                continue

            result, count = blur_faces_in_image(img_cv2, blocks=blocks)
            if count == 0:
                continue

            # Re-encode with same format (try JPEG first, fall back to PNG)
            ct = image_part.content_type
            fmt = "PNG" if "png" in ct.lower() else "JPEG"
            new_bytes = cv2_to_bytes(result, fmt=fmt)

            # Monkey-patch blob on the part object
            image_part._blob = new_bytes
            total_faces += count

        except Exception as e:
            pass  # skip shapes that can't be processed

    doc.save(str(output_path))
    return total_faces


# ── Face blur: Excel workbooks ────────────────────────────────────────────────

def blur_faces_xlsx(input_path: Path, output_path: Path,
                    blocks: int = 6) -> int:
    """
    Detect and pixelate faces in images embedded in an .xlsx workbook.
    Returns number of faces blurred.
    """
    if not XLSX_OK:
        raise RuntimeError("openpyxl required: pip install openpyxl")

    cv2, np = _get_cv2()
    if cv2 is None:
        raise RuntimeError("OpenCV not available")

    import shutil
    shutil.copy2(str(input_path), str(output_path))

    # openpyxl stores images as _images list on each worksheet
    wb = openpyxl.load_workbook(str(output_path))
    total_faces = 0

    for sheet in wb.worksheets:
        for img_obj in getattr(sheet, "_images", []):
            try:
                # img_obj.ref is the image data (BytesIO or bytes)
                raw = img_obj.ref
                if hasattr(raw, "read"):
                    raw = raw.read()
                np_arr  = np.frombuffer(raw, dtype=np.uint8)
                img_cv2 = cv2.imdecode(np_arr, cv2.IMREAD_COLOR)
                if img_cv2 is None:
                    continue

                result, count = blur_faces_in_image(img_cv2, blocks=blocks)
                if count == 0:
                    continue

                # Re-encode and replace
                new_bytes = cv2_to_bytes(result, fmt="PNG")
                img_obj.ref = io.BytesIO(new_bytes)
                total_faces += count
            except Exception:
                pass

    wb.save(str(output_path))
    return total_faces

# ── Core scanner ──────────────────────────────────────────────────────────────

def scan_pdf(pdf_path: Path, force_ocr=False, lang="dan+eng",
             dpi=300, poppler_path=None) -> dict:
    results = {"cprs": [], "dates": [], "page_methods": {}}

    with pdfplumber.open(pdf_path) as pdf:
        images = None
        if OCR_AVAILABLE:
            needs_ocr = (list(range(len(pdf.pages))) if force_ocr
                         else [i for i, p in enumerate(pdf.pages) if not is_text_page(p)])
            if needs_ocr:
                print(f"  Rendering pages to images for OCR (DPI={dpi})...", flush=True)
                images = convert_from_path(str(pdf_path), dpi=dpi, poppler_path=poppler_path)

        for page_num, page in enumerate(pdf.pages, start=1):
            use_text = not force_ocr and is_text_page(page)
            if use_text:
                method = "text"
                text = page.extract_text() or ""
                cprs, dates = extract_matches(text, page_num, "text")
            elif OCR_AVAILABLE and images is not None:
                method = "ocr"
                _img = images[page_num-1]
                images[page_num-1] = None  # release PIL image as soon as OCR is done
                cprs, dates = extract_matches(ocr_page_cached(_img, lang), page_num, "ocr")
                del _img
            else:
                method = "skipped"
                if not OCR_AVAILABLE:
                    print(f"  Page {page_num}: image-based but OCR unavailable.")
                cprs, dates = [], []

            results["page_methods"][page_num] = method
            results["cprs"].extend(cprs)
            results["dates"].extend(dates)

    results["dates"] = dedup_dates(results["dates"])
    return results


# ── Output ────────────────────────────────────────────────────────────────────

def print_results(pdf_path: Path, results: dict):
    methods    = results["page_methods"]
    text_pages = [p for p, m in methods.items() if m == "text"]
    ocr_pages  = [p for p, m in methods.items() if m == "ocr"]
    skip_pages = [p for p, m in methods.items() if m == "skipped"]

    print(f"\n{'='*62}")
    print(f"File : {pdf_path}")
    print(f"Pages: {len(methods)}  |  text: {len(text_pages)}  |  OCR: {len(ocr_pages)}  |  skipped: {len(skip_pages)}")
    print(f"{'='*62}")
    if ocr_pages:
        print(f"  [OCR]  Applied to page(s): {', '.join(map(str, ocr_pages))}")
    if skip_pages:
        print(f"  [SKIP] Skipped page(s): {', '.join(map(str, skip_pages))}")

    cprs  = results["cprs"]
    dates = results["dates"]

    print(f"\n  CPR Numbers found: {len(cprs)}")
    if cprs:
        for hit in cprs:
            tag = " [OCR]" if hit["source"] == "ocr" else ""
            print(f"    Page {hit['page']:>3}: {hit['formatted']:<16}  (raw: \"{hit['raw']}\"){tag}")
    else:
        print("    None found.")

    print(f"\n  Dates found: {len(dates)}")
    if dates:
        for hit in dates:
            tag = " [OCR]" if hit["source"] == "ocr" else ""
            print(f"    Page {hit['page']:>3}: {hit['raw']:<28}  [{hit['format']}]{tag}")
    else:
        print("    None found.")
    print()


# ── Entry point ───────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Scan PDF and Word documents for Danish CPR numbers, dates and personal data.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    parser.add_argument("pdfs", nargs="+", metavar="FILE", help="PDF/Word file(s) or folder(s) to scan")
    parser.add_argument("--ocr", action="store_true", help="Force OCR on every page")
    parser.add_argument("--lang", default="dan+eng", metavar="LANG", help="Tesseract language(s), default: dan+eng")
    parser.add_argument("--dpi", type=int, default=300, metavar="DPI", help="Rendering DPI for OCR, default: 300")
    parser.add_argument("--poppler", default=None, metavar="PATH", help="Path to Poppler bin folder (Windows)")
    parser.add_argument("--older-than", type=int, default=None, metavar="DAYS",
                        help="List files with CPR numbers AND dates older than DAYS")
    parser.add_argument("--mask", action="store_true",
                        help="Black out CPR numbers -> <n>_masked.pdf/.docx")
    parser.add_argument("--anonymise", action="store_true",
                        help="Black out ALL personal data -> <n>_anonymised.pdf/.docx")
    parser.add_argument("--dry-run", action="store_true",
                        help="Scan and report findings without writing any output files")
    parser.add_argument("--log", default=None, metavar="FILE",
                        help="Write a structured JSON log of all findings to FILE")
    parser.add_argument("--blur-faces", action="store_true",
                        help="Detect and pixelate portrait photos -> <n>_faces.pdf/.docx/.xlsx/.jpg")
    parser.add_argument("--blur-strength", type=int, default=6, metavar="N",
                        help="Face blur strength: lower = stronger (default: 6, range: 2-20)")
    args = parser.parse_args()

    dry_run = args.dry_run

    # Logging setup
    console_handler = logging.StreamHandler(sys.stdout)
    console_handler.setFormatter(logging.Formatter("%(message)s"))
    logger.addHandler(console_handler)

    if dry_run:
        print("=" * 62)
        print("  DRY RUN - no files will be written")
        print("=" * 62 + "\n")
        _log("INFO", None, "dry_run_started")

    # Dependency warnings
    if not OCR_AVAILABLE:
        missing = [m for m, ok in [("pdf2image", PDF2IMAGE_OK), ("pytesseract", TESSERACT_OK)] if not ok]
        msg = f"OCR disabled - pip install {' '.join(missing)}"
        print(f"WARNING: {msg}\n")
        _log("WARNING", None, msg)

    if (args.mask or args.anonymise) and not MASK_AVAILABLE:
        msg = "--mask/--anonymise require: pip install pypdf reportlab"
        print(f"WARNING: {msg}\n")
        _log("WARNING", None, msg)

    if not DOCX_OK:
        print("INFO: python-docx not installed - .docx files will be skipped.")
        print("      Install with: pip install python-docx\n")
        _log("WARNING", None, "python-docx not installed - .docx files skipped")

    if not XLSX_OK:
        print("INFO: openpyxl not installed - .xlsx/.csv files will be skipped.")
        print("      Install with: pip install openpyxl\n")
        _log("WARNING", None, "openpyxl not installed - .xlsx files skipped")

    if args.blur_faces and not CV2_OK:
        print("WARNING: --blur-faces requires OpenCV: pip install opencv-python\n")

    if args.anonymise:
        if not SPACY_OK:
            msg = "--anonymise requires spaCy: pip install spacy"
            print(f"WARNING: {msg}\n")
            _log("WARNING", None, msg)
        else:
            nlp = load_nlp()
            if nlp is None:
                msg = "No spaCy model found - falling back to regex-only"
                print(f"WARNING: {msg}\n")
                _log("WARNING", None, msg)

    # Collect files
    SUPPORTED = {".pdf", ".docx", ".xlsx", ".xlsm", ".csv",
                 ".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".tif", ".webp"}
    all_paths = []
    for entry in args.pdfs:
        path = Path(entry)
        if not path.exists():
            print(f"Not found: {path}")
            _log("WARNING", path, "file_not_found")
        elif path.is_dir():
            found = sorted(p for p in path.rglob("*") if p.suffix.lower() in SUPPORTED)
            pdf_count  = sum(1 for p in found if p.suffix.lower() == ".pdf")
            docx_count = sum(1 for p in found if p.suffix.lower() == ".docx")
            xlsx_count = sum(1 for p in found if p.suffix.lower() in {".xlsx", ".xlsm", ".csv"})
            img_count  = sum(1 for p in found if p.suffix.lower() in {".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".tif", ".webp"})
            print(f"Found {pdf_count} PDF(s), {docx_count} Word doc(s), {xlsx_count} spreadsheet(s) and {img_count} image(s) in: {path}")
            _log("INFO", path, "folder_scanned", pdf_count=pdf_count, docx_count=docx_count, xlsx_count=xlsx_count, img_count=img_count)
            all_paths.extend(found)
        elif path.suffix.lower() in SUPPORTED:
            all_paths.append(path)
        else:
            print(f"Unsupported file type, skipping: {path}")
            _log("WARNING", path, "unsupported_type")

    if not all_paths:
        print("No supported files to process.")
        _log("INFO", None, "no_files_found")
        if args.log:
            flush_log(Path(args.log))
        return

    # Process files
    all_results = []
    for path in all_paths:
        try:
            ext = path.suffix.lower()

            if ext in {".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".tif", ".webp"}:
                # Standalone image — face blur only (triggered by --blur-faces, --mask, or --anonymise)
                print(f"\n{'='*62}")
                print(f"File : {path}  [image]")
                print(f"{'='*62}")
                _log("INFO", path, "scanned", file_type="image")

                do_blur = args.blur_faces or args.mask or args.anonymise
                if do_blur:
                    if not CV2_OK:
                        print(f"  [FACE] Skipping - opencv-python not installed.")
                        print(f"         pip install opencv-python\n")
                        _log("WARNING", path, "skipped_no_opencv")
                        continue
                    out = path.with_stem(path.stem + "_faces")
                    if dry_run:
                        print(f"  [DRY-RUN] Would write -> {out.name}  (face blur)\n")
                        _log("DRY_RUN", path, "face_blur_skipped_dry_run", output=str(out))
                    else:
                        print(f"  [FACE] Scanning for faces ...", flush=True)
                        n = blur_faces_image_file(path, out, blocks=args.blur_strength)
                        if n:
                            print(f"  [FACE] Done - {n} face(s) blurred -> {out.name}\n")
                            _log("ACTION", path, "faces_blurred", output=str(out), faces=n)
                        else:
                            out.unlink(missing_ok=True)
                            print(f"  [FACE] No faces detected - no output written.\n")
                            _log("INFO", path, "no_faces_detected")
                else:
                    print(f"  Image file: use --blur-faces, --mask, or --anonymise to pixelate portraits.\n")
                    _log("INFO", path, "image_no_action_requested")
                # Images have no CPR/date data — don't add to all_results
                continue

            elif ext == ".docx":
                if not DOCX_OK:
                    print(f"Skipping {path.name} - python-docx not installed.")
                    _log("WARNING", path, "skipped_no_python_docx")
                    continue

                results = scan_docx(path)
                print_docx_results(path, results)
                all_results.append((path, results))
                _log("INFO", path, "scanned",
                     file_type="docx",
                     cpr_count=len(results["cprs"]),
                     date_count=len(results["dates"]),
                     cprs=[h["formatted"] for h in results["cprs"]])

                if args.mask:
                    out = path.with_stem(path.stem + "_masked")
                    if results["cprs"]:
                        if dry_run:
                            print(f"  [DRY-RUN] Would write -> {out.name}  ({len(results['cprs'])} CPR region(s))")
                            _log("DRY_RUN", path, "mask_skipped_dry_run",
                                 output=str(out), cpr_count=len(results["cprs"]))
                        else:
                            print(f"  [MASK] Writing -> {out.name} ...", flush=True)
                            n = redact_docx(path, out, results, use_ner=False)
                            print(f"  [MASK] Done - {n} region(s) redacted.\n")
                            _log("ACTION", path, "masked", output=str(out), regions=n)
                    else:
                        print("  [MASK] No CPR numbers found - skipping.\n")
                        _log("INFO", path, "mask_skipped_no_cpr")

                if args.anonymise:
                    out = path.with_stem(path.stem + "_anonymised")
                    if dry_run:
                        spans = find_pii_spans_in_text(results["_full_text"], use_ner=True)
                        label_counts = {}
                        for _, _, lbl in spans:
                            label_counts[lbl] = label_counts.get(lbl, 0) + 1
                        summary = "  ".join(f"{lbl}:{c}" for lbl, c in sorted(label_counts.items()))
                        print(f"  [DRY-RUN] Would write -> {out.name}  ({len(spans)} region(s): {summary})")
                        _log("DRY_RUN", path, "anonymise_skipped_dry_run",
                             output=str(out), total_regions=len(spans), by_label=label_counts)
                    else:
                        print(f"  [ANON] Writing -> {out.name} ...", flush=True)
                        n = redact_docx(path, out, results, use_ner=True)
                        print(f"  [ANON] Done - {n} region(s) redacted.\n")
                        _log("ACTION", path, "anonymised", output=str(out), regions=n)

                if args.blur_faces:
                    if not CV2_OK:
                        print(f"  [FACE] Skipping - opencv-python not installed.")
                    else:
                        out = path.with_stem(path.stem + "_faces")
                        if dry_run:
                            print(f"  [DRY-RUN] Would write -> {out.name}  (face blur)")
                            _log("DRY_RUN", path, "face_blur_skipped_dry_run", output=str(out))
                        else:
                            print(f"  [FACE] Scanning for faces ...", flush=True)
                            n = blur_faces_docx(path, out, blocks=args.blur_strength)
                            if n:
                                print(f"  [FACE] Done - {n} face(s) blurred -> {out.name}\n")
                                _log("ACTION", path, "faces_blurred", output=str(out), faces=n)
                            else:
                                out.unlink(missing_ok=True)
                                print(f"  [FACE] No faces detected.\n")
                                _log("INFO", path, "no_faces_detected")

            elif ext in {".xlsx", ".xlsm"}:
                if not XLSX_OK:
                    print(f"Skipping {path.name} - openpyxl not installed.")
                    _log("WARNING", path, "skipped_no_openpyxl")
                    continue

                results = scan_xlsx(path)
                print_xlsx_results(path, results, "xlsx")
                all_results.append((path, results))
                _log("INFO", path, "scanned",
                     file_type="xlsx",
                     cpr_count=len(results["cprs"]),
                     date_count=len(results["dates"]),
                     cprs=[h["formatted"] for h in results["cprs"]])

                if args.mask:
                    out = path.with_stem(path.stem + "_masked")
                    if results["cprs"]:
                        if dry_run:
                            print(f"  [DRY-RUN] Would write -> {out.name}  ({len(results['cprs'])} CPR cell(s))")
                            _log("DRY_RUN", path, "mask_skipped_dry_run",
                                 output=str(out), cpr_count=len(results["cprs"]))
                        else:
                            print(f"  [MASK] Writing -> {out.name} ...", flush=True)
                            n = redact_xlsx(path, out, results, use_ner=False)
                            print(f"  [MASK] Done - {n} cell(s) redacted.\n")
                            _log("ACTION", path, "masked", output=str(out), regions=n)
                    else:
                        print("  [MASK] No CPR numbers found - skipping.\n")
                        _log("INFO", path, "mask_skipped_no_cpr")

                if args.anonymise:
                    out = path.with_stem(path.stem + "_anonymised")
                    if dry_run:
                        full_text = " ".join(
                            _cell_text(c)
                            for s in results["_wb"].worksheets
                            for row in s.iter_rows()
                            for c in row
                        )
                        spans = find_pii_spans_in_text(full_text, use_ner=True)
                        label_counts = {}
                        for _, _, lbl in spans:
                            label_counts[lbl] = label_counts.get(lbl, 0) + 1
                        summary = "  ".join(f"{lbl}:{c}" for lbl, c in sorted(label_counts.items()))
                        print(f"  [DRY-RUN] Would write -> {out.name}  ({len(spans)} region(s): {summary})")
                        _log("DRY_RUN", path, "anonymise_skipped_dry_run",
                             output=str(out), total_regions=len(spans), by_label=label_counts)
                    else:
                        print(f"  [ANON] Writing -> {out.name} ...", flush=True)
                        n = redact_xlsx(path, out, results, use_ner=True)
                        print(f"  [ANON] Done - {n} cell(s) redacted.\n")
                        _log("ACTION", path, "anonymised", output=str(out), regions=n)

                if args.blur_faces:
                    if not CV2_OK:
                        print(f"  [FACE] Skipping - opencv-python not installed.")
                    else:
                        out = path.with_stem(path.stem + "_faces")
                        if dry_run:
                            print(f"  [DRY-RUN] Would write -> {out.name}  (face blur)")
                            _log("DRY_RUN", path, "face_blur_skipped_dry_run", output=str(out))
                        else:
                            print(f"  [FACE] Scanning for faces ...", flush=True)
                            n = blur_faces_xlsx(path, out, blocks=args.blur_strength)
                            if n:
                                print(f"  [FACE] Done - {n} face(s) blurred -> {out.name}\n")
                                _log("ACTION", path, "faces_blurred", output=str(out), faces=n)
                            else:
                                out.unlink(missing_ok=True)
                                print(f"  [FACE] No faces detected.\n")
                                _log("INFO", path, "no_faces_detected")

            elif ext == ".csv":
                results = scan_csv(path)
                print_xlsx_results(path, results, "csv")
                all_results.append((path, results))
                _log("INFO", path, "scanned",
                     file_type="csv",
                     cpr_count=len(results["cprs"]),
                     date_count=len(results["dates"]),
                     cprs=[h["formatted"] for h in results["cprs"]])

                if args.mask:
                    out = path.with_stem(path.stem + "_masked")
                    if results["cprs"]:
                        if dry_run:
                            print(f"  [DRY-RUN] Would write -> {out.name}  ({len(results['cprs'])} CPR cell(s))")
                            _log("DRY_RUN", path, "mask_skipped_dry_run",
                                 output=str(out), cpr_count=len(results["cprs"]))
                        else:
                            print(f"  [MASK] Writing -> {out.name} ...", flush=True)
                            n = redact_csv(path, out, use_ner=False)
                            print(f"  [MASK] Done - {n} cell(s) redacted.\n")
                            _log("ACTION", path, "masked", output=str(out), regions=n)
                    else:
                        print("  [MASK] No CPR numbers found - skipping.\n")
                        _log("INFO", path, "mask_skipped_no_cpr")

                if args.anonymise:
                    out = path.with_stem(path.stem + "_anonymised")
                    if dry_run:
                        import csv as _csv
                        full_text = ""
                        with open(path, newline="", encoding="utf-8-sig", errors="replace") as f:
                            for row in _csv.reader(f):
                                full_text += " ".join(row) + " "
                        spans = find_pii_spans_in_text(full_text, use_ner=True)
                        label_counts = {}
                        for _, _, lbl in spans:
                            label_counts[lbl] = label_counts.get(lbl, 0) + 1
                        summary = "  ".join(f"{lbl}:{c}" for lbl, c in sorted(label_counts.items()))
                        print(f"  [DRY-RUN] Would write -> {out.name}  ({len(spans)} region(s): {summary})")
                        _log("DRY_RUN", path, "anonymise_skipped_dry_run",
                             output=str(out), total_regions=len(spans), by_label=label_counts)
                    else:
                        print(f"  [ANON] Writing -> {out.name} ...", flush=True)
                        n = redact_csv(path, out, use_ner=True)
                        print(f"  [ANON] Done - {n} cell(s) redacted.\n")
                        _log("ACTION", path, "anonymised", output=str(out), regions=n)

            else:
                results = scan_pdf(path, force_ocr=args.ocr, lang=args.lang,
                                   dpi=args.dpi, poppler_path=args.poppler)
                print_results(path, results)
                all_results.append((path, results))
                _log("INFO", path, "scanned",
                     file_type="pdf",
                     pages=len(results["page_methods"]),
                     ocr_pages=sum(1 for m in results["page_methods"].values() if m == "ocr"),
                     cpr_count=len(results["cprs"]),
                     date_count=len(results["dates"]),
                     cprs=[h["formatted"] for h in results["cprs"]])

                if args.mask:
                    out = path.with_stem(path.stem + "_masked")
                    if results["cprs"]:
                        if dry_run:
                            print(f"  [DRY-RUN] Would write -> {out.name}  ({len(results['cprs'])} CPR region(s))")
                            _log("DRY_RUN", path, "mask_skipped_dry_run",
                                 output=str(out), cpr_count=len(results["cprs"]))
                        else:
                            print(f"  [MASK] Writing -> {out.name} ...", flush=True)
                            n = redact_pdf(path, out, results, args.ocr, args.lang,
                                           args.dpi, args.poppler, use_ner=False)
                            if n is not False:
                                print(f"  [MASK] Done - {n} region(s) redacted.\n")
                                _log("ACTION", path, "masked", output=str(out), regions=n)
                    else:
                        print("  [MASK] No CPR numbers found - skipping.\n")
                        _log("INFO", path, "mask_skipped_no_cpr")

                if args.anonymise:
                    out = path.with_stem(path.stem + "_anonymised")
                    if dry_run:
                        full_text = ""
                        with pdfplumber.open(path) as _pdf:
                            for _page in _pdf.pages:
                                full_text += (_page.extract_text() or "") + " "
                        spans = find_pii_spans_in_text(full_text, use_ner=True)
                        label_counts = {}
                        for _, _, lbl in spans:
                            label_counts[lbl] = label_counts.get(lbl, 0) + 1
                        summary = "  ".join(f"{lbl}:{c}" for lbl, c in sorted(label_counts.items()))
                        print(f"  [DRY-RUN] Would write -> {out.name}  ({len(spans)} region(s): {summary})")
                        _log("DRY_RUN", path, "anonymise_skipped_dry_run",
                             output=str(out), total_regions=len(spans), by_label=label_counts)
                    else:
                        print(f"  [ANON] Writing -> {out.name} ...", flush=True)
                        n = redact_pdf(path, out, results, args.ocr, args.lang,
                                       args.dpi, args.poppler, use_ner=True)
                        if n is not False:
                            print(f"  [ANON] Done - {n} region(s) redacted.\n")
                            _log("ACTION", path, "anonymised", output=str(out), regions=n)

                if args.blur_faces:
                    if not CV2_OK:
                        print(f"  [FACE] Skipping - opencv-python not installed.")
                    elif not OCR_AVAILABLE:
                        print(f"  [FACE] Skipping - pdf2image required for PDF face blur.")
                    else:
                        out = path.with_stem(path.stem + "_faces")
                        if dry_run:
                            print(f"  [DRY-RUN] Would write -> {out.name}  (face blur)")
                            _log("DRY_RUN", path, "face_blur_skipped_dry_run", output=str(out))
                        else:
                            print(f"  [FACE] Scanning pages for faces ...", flush=True)
                            n = blur_faces_pdf(path, out, poppler_path=args.poppler, blocks=args.blur_strength)
                            if n:
                                print(f"  [FACE] Done - {n} face(s) blurred -> {out.name}\n")
                                _log("ACTION", path, "faces_blurred", output=str(out), faces=n)
                            else:
                                out.unlink(missing_ok=True)
                                print(f"  [FACE] No faces detected.\n")
                                _log("INFO", path, "no_faces_detected")

        except Exception as e:
            print(f"Error processing {path}: {e}")
            _log("ERROR", path, str(e))

    if args.older_than is not None:
        flagged = build_flagged_list(all_results, args.older_than)
        print_flagged(flagged, args.older_than)
        _log("INFO", None, "flagged_summary",
             older_than_days=args.older_than,
             flagged_count=len(flagged),
             flagged_files=[str(f["path"]) for f in flagged])

    # Final summary
    total_cprs     = sum(len(r["cprs"])  for _, r in all_results)
    total_dates    = sum(len(r["dates"]) for _, r in all_results)
    files_with_cpr = sum(1 for _, r in all_results if r["cprs"])
    print(f"{'--'*31}")
    print(f"  Scanned : {len(all_results)} file(s)")
    print(f"  CPR nos : {total_cprs} found in {files_with_cpr} file(s)")
    print(f"  Dates   : {total_dates} found")
    if dry_run:
        print("  Mode    : DRY RUN - no files written")
    print(f"{'--'*31}\n")
    _log("INFO", None, "scan_complete",
         files_scanned=len(all_results),
         total_cprs=total_cprs,
         total_dates=total_dates,
         files_with_cpr=files_with_cpr,
         dry_run=dry_run)

    if args.log:
        flush_log(Path(args.log))

if __name__ == "__main__":
    main()


def count_faces_in_file(path, poppler_path=None, neighbors: int = 4) -> int:
    """
    Return the number of faces detected in a file (image, PDF, docx, xlsx).
    Uses only this module's cv2/numpy — never triggers a second import from
    outside (avoids the 'recursion detected during loading cv2' error on macOS).
    neighbors controls detection strictness: higher = fewer false positives.
    """
    import sys as _sys
    cv2, np = _get_cv2()
    if cv2 is None:
        _face_log(f"[face] cv2 unavailable: {_cv2_import_error}")
        return 0

    ext = Path(path).suffix.lower()
    total = 0
    cascades = _get_face_cascades()
    _face_log(f"[face] {Path(path).name}  ext={ext}  cascades={len(cascades)}  neighbors={neighbors}")

    try:
        if ext in {".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".tif", ".webp"}:
            img = cv2.imread(str(path))
            _face_log(f"[face]   imread={img is not None}  shape={getattr(img, 'shape', None)}")
            if img is not None:
                total = len(detect_faces_cv2(img, neighbors=neighbors))
                _face_log(f"[face]   detected={total}")

        elif ext == ".pdf":
            if PYMUPDF_AVAILABLE:
                import fitz as _fitz
                doc = _fitz.open(str(path))
                for page_idx in range(min(5, len(doc))):
                    pix = doc[page_idx].get_pixmap(dpi=100)
                    arr = cv2.imdecode(
                        np.frombuffer(pix.tobytes("jpeg"), np.uint8),
                        cv2.IMREAD_COLOR)
                    if arr is not None:
                        total += len(detect_faces_cv2(arr, neighbors=neighbors))
                    if total > 0:
                        break
                doc.close()
            else:
                from pdf2image import convert_from_path
                pages = convert_from_path(str(path), dpi=100,
                                          first_page=1, last_page=5,
                                          poppler_path=poppler_path)
                for page in pages:
                    arr = cv2.cvtColor(np.array(page), cv2.COLOR_RGB2BGR)
                    total += len(detect_faces_cv2(arr, neighbors=neighbors))
                    if total > 0:
                        break

        elif ext == ".docx":
            from docx import Document
            from docx.enum.shape import WD_INLINE_SHAPE
            doc = Document(str(path))
            for shape in doc.inline_shapes:
                try:
                    if shape.type != WD_INLINE_SHAPE.PICTURE:
                        continue
                    blip = shape._inline.graphic.graphicData.pic.blipFill.blip
                    blob = doc.part.related_parts[blip.embed].blob
                    arr  = np.frombuffer(blob, dtype=np.uint8)
                    img  = cv2.imdecode(arr, cv2.IMREAD_COLOR)
                    if img is not None:
                        total += len(detect_faces_cv2(img, neighbors=neighbors))
                except Exception:
                    pass

        elif ext in {".xlsx", ".xlsm"}:
            import openpyxl
            wb = openpyxl.load_workbook(str(path), read_only=False, data_only=True)
            for sname in wb.sheetnames:
                for img_obj in wb[sname]._images:
                    try:
                        blob = img_obj._data()
                        arr  = np.frombuffer(blob, dtype=np.uint8)
                        img  = cv2.imdecode(arr, cv2.IMREAD_COLOR)
                        if img is not None:
                            total += len(detect_faces_cv2(img, neighbors=neighbors))
                    except Exception:
                        pass
            wb.close()

    except Exception:
        pass

    return total
