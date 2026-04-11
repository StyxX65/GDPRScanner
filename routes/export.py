"""
Excel and Article 30 export, bulk delete
"""
from __future__ import annotations
import json, io, re, traceback, logging
from pathlib import Path
from flask import Blueprint, Response, jsonify, request
from routes import state
from app_config import _GUID_RE, _resolve_display_name

try:
    from gdpr_db import get_db as _get_db
    DB_OK = True
except ImportError:
    DB_OK = False
    def _get_db(*a, **kw): return None  # type: ignore[misc]

try:
    from m365_connector import M365PermissionError
except ImportError:
    class M365PermissionError(Exception): pass  # type: ignore[no-redef]

bp = Blueprint("export", __name__)
logger = logging.getLogger(__name__)


def _build_excel_bytes() -> tuple[bytes, str]:
    """Build the M365 scan Excel workbook and return (bytes, filename).
    Raises on error. Used by export_excel() and send_report()."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    HEADER_BG  = "1F3864"
    HEADER_FG  = "FFFFFF"
    ALT_BG     = "EEF2FF"
    SOURCE_MAP = {
        "email":      ("📧 Outlook",       "D6E4F7"),
        "onedrive":   ("💾 OneDrive",     "D6F7E4"),
        "sharepoint": ("🌐 SharePoint",   "FFF0D6"),
        "teams":      ("💬 Teams",        "F7D6F0"),
        "gmail":      ("📧 Gmail",        "D6EAF8"),
        "gdrive":     ("💾 Google Drive", "D5F5E3"),
        "local":      ("📁 Local",        "E6F7E6"),
        "smb":        ("🌐 Network",      "E0F0FA"),
    }
    COLS = [
        ("Name / Subject",    45),
        ("CPR Hits",           9),
        ("Face count",         9),
        ("GPS",                6),
        ("Special category",  22),
        ("EXIF author",       18),
        ("Folder",            30),
        ("Account",           24),
        ("Role",              10),
        ("Disposition",       18),
        ("Date Modified",     14),
        ("Size (KB)",         10),
        ("URL",               50),
    ]

    thin   = Side(style="thin", color="CCCCCC")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    def _fill(hex_col):
        return PatternFill("solid", fgColor=hex_col)

    def _write_sheet(ws, rows, tab_color):
        ws.sheet_properties.tabColor = tab_color
        for col_idx, (col_name, col_w) in enumerate(COLS, 1):
            cell = ws.cell(row=1, column=col_idx, value=col_name)
            cell.font      = Font(name="Arial", bold=True, color=HEADER_FG, size=10)
            cell.fill      = _fill(HEADER_BG)
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            cell.border    = border
            ws.column_dimensions[get_column_letter(col_idx)].width = col_w
        ws.row_dimensions[1].height = 20
        ws.freeze_panes = "A2"

        for r_idx, item in enumerate(rows, 2):
            row_fill = _fill(ALT_BG if r_idx % 2 == 0 else "FFFFFF")
            _disp = ""
            if DB_OK:
                try:
                    _d = _get_db().get_disposition(item.get("id", ""))
                    _disp = (_d.get("status", "") if _d else "")
                except Exception:
                    pass
            _sc = item.get("special_category", [])
            _sc_str = ", ".join(
                s for s in (_sc if isinstance(_sc, list) else [str(_sc or "")])
                if s not in ("gps_location", "exif_pii")
            )
            _exif   = item.get("exif") or {}
            _gps    = _exif.get("gps")
            _author = _exif.get("author") or ""
            values = [
                item.get("name", ""),
                item.get("cpr_count", 0),
                item.get("face_count", 0),
                "✔" if _gps else "",
                _sc_str,
                _author,
                item.get("folder", ""),
                item.get("account_name", "") or item.get("source", ""),
                item.get("user_role", ""),
                _disp,
                item.get("modified", ""),
                item.get("size_kb", ""),
                item.get("url", ""),
            ]
            for col_idx, val in enumerate(values, 1):
                is_url = col_idx == 13 and val
                cell = ws.cell(row=r_idx, column=col_idx, value=val)
                cell.font      = Font(name="Arial", size=10,
                                     color="1155CC" if is_url else "000000",
                                     underline="single" if is_url else None)
                cell.fill      = row_fill
                cell.alignment = Alignment(vertical="center", wrap_text=(col_idx == 1))
                cell.border    = border
            ws.row_dimensions[r_idx].height = 16

        if rows:
            tr = len(rows) + 2
            ws.cell(row=tr, column=1, value="Total").font = Font(name="Arial", bold=True, size=10)
            ws.cell(row=tr, column=2, value=f"=SUM(B2:B{tr-1})").font = Font(name="Arial", bold=True, size=10)
            for col_idx in range(1, len(COLS) + 1):
                ws.cell(row=tr, column=col_idx).fill   = _fill("D0D8F0")
                ws.cell(row=tr, column=col_idx).border = border

        ws.auto_filter.ref = f"A1:{get_column_letter(len(COLS))}1"

    wb     = Workbook()
    ws_sum = wb.active
    ws_sum.title = "Summary"
    ws_sum.sheet_properties.tabColor = "1F3864"
    ws_sum["A1"] = "GDPRScanner — Export"
    ws_sum["A1"].font = Font(name="Arial", bold=True, size=14, color=HEADER_FG)
    ws_sum["A1"].fill = _fill(HEADER_BG)
    ws_sum.merge_cells("A1:D1")
    ws_sum["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws_sum.row_dimensions[1].height = 28

    import datetime as _dt
    ws_sum["A2"] = "Generated:"
    ws_sum["B2"] = _dt.datetime.now().strftime("%Y-%m-%d %H:%M")
    ws_sum["A3"] = "Total flagged items:"
    ws_sum["B3"] = len(state.flagged_items)
    gps_count = sum(1 for i in state.flagged_items if (i.get("exif") or {}).get("gps"))
    if gps_count:
        ws_sum["A4"] = "Items with GPS data:"
        ws_sum["B4"] = gps_count
    for cell in (ws_sum["A2"], ws_sum["A3"], ws_sum["A4"]):
        cell.font = Font(name="Arial", bold=True, size=10)
    for cell in (ws_sum["B2"], ws_sum["B3"], ws_sum["B4"]):
        cell.font = Font(name="Arial", size=10)
    ws_sum.column_dimensions["A"].width = 22
    ws_sum.column_dimensions["B"].width = 20

    for ci, h in enumerate(["Source", "Items", "Total CPR Hits"], 1):
        cell = ws_sum.cell(row=6, column=ci, value=h)
        cell.font      = Font(name="Arial", bold=True, color=HEADER_FG, size=10)
        cell.fill      = _fill(HEADER_BG)
        cell.border    = border
        cell.alignment = Alignment(horizontal="center", vertical="center")
    ws_sum.row_dimensions[6].height = 18
    ws_sum.column_dimensions["C"].width = 16

    by_source: dict = {}
    for item in state.flagged_items:
        by_source.setdefault(item.get("source_type", "other"), []).append(item)

    sum_row = 7
    for src_key, (label, tab_bg) in SOURCE_MAP.items():
        items = by_source.get(src_key, [])
        if not items:
            continue
        ws_sum.cell(row=sum_row, column=1, value=label).font = Font(name="Arial", size=10)
        ws_sum.cell(row=sum_row, column=2, value=len(items)).font = Font(name="Arial", size=10)
        ws_sum.cell(row=sum_row, column=3, value=sum(i.get("cpr_count", 0) for i in items)).font = Font(name="Arial", size=10)
        for ci in range(1, 4):
            ws_sum.cell(row=sum_row, column=ci).border = border
            ws_sum.cell(row=sum_row, column=ci).fill = _fill("EEF2FF" if sum_row % 2 == 0 else "FFFFFF")
        sum_row += 1

    for src_key, (label, tab_bg) in SOURCE_MAP.items():
        items = by_source.get(src_key, [])
        if not items:
            continue
        clean_label = label.split(" ", 1)[1]
        _write_sheet(wb.create_sheet(title=clean_label), items, tab_bg)

    # GPS items sheet
    gps_items = [i for i in state.flagged_items if (i.get("exif") or {}).get("gps")]
    if gps_items:
        ws_gps = wb.create_sheet(title="GPS locations")
        ws_gps.sheet_properties.tabColor = "1A7A6E"
        GPS_COLS = [
            ("Name", 40), ("Latitude", 14), ("Longitude", 14),
            ("Maps link", 50), ("Account", 24), ("Date Modified", 14),
        ]
        for col_idx, (col_name, col_w) in enumerate(GPS_COLS, 1):
            cell = ws_gps.cell(row=1, column=col_idx, value=col_name)
            cell.font      = Font(name="Arial", bold=True, color=HEADER_FG, size=10)
            cell.fill      = _fill("1A7A6E")
            cell.alignment = Alignment(horizontal="center", vertical="center")
            cell.border    = border
            ws_gps.column_dimensions[get_column_letter(col_idx)].width = col_w
        ws_gps.freeze_panes = "A2"
        for r_idx, item in enumerate(gps_items, 2):
            _exif = item.get("exif") or {}
            _gps  = _exif.get("gps") or {}
            row_fill = _fill("E0F7F4" if r_idx % 2 == 0 else "FFFFFF")
            for col_idx, val in enumerate([
                item.get("name", ""),
                _gps.get("lat", ""),
                _gps.get("lon", ""),
                _gps.get("maps_url", ""),
                item.get("account_name", "") or item.get("source", ""),
                item.get("modified", ""),
            ], 1):
                is_link = col_idx == 4 and val
                cell = ws_gps.cell(row=r_idx, column=col_idx, value=val)
                cell.font   = Font(name="Arial", size=10,
                                   color="1155CC" if is_link else "000000",
                                   underline="single" if is_link else None)
                cell.fill   = row_fill
                cell.border = border
        ws_gps.auto_filter.ref = f"A1:{get_column_letter(len(GPS_COLS))}1"

    # External transfers sheet
    ext_items = [i for i in state.flagged_items
                 if i.get("transfer_risk") in ("external-recipient", "external-share", "shared")]
    if ext_items:
        ws_ext = wb.create_sheet(title="External transfers")
        _write_sheet(ws_ext, ext_items, "E74C3C")
        ws_ext.sheet_properties.tabColor = "E74C3C"
        ws_sum.cell(row=sum_row, column=1, value="⚠ External transfers").font = Font(name="Arial", size=10, bold=True, color="E74C3C")
        ws_sum.cell(row=sum_row, column=2, value=len(ext_items)).font = Font(name="Arial", size=10, bold=True, color="E74C3C")
        ws_sum.cell(row=sum_row, column=3, value=sum(i.get("cpr_count", 0) for i in ext_items)).font = Font(name="Arial", size=10, bold=True, color="E74C3C")
        for ci in range(1, 4):
            ws_sum.cell(row=sum_row, column=ci).border = border
            ws_sum.cell(row=sum_row, column=ci).fill = _fill("FDE8E8")

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    fname = f"m365_scan_{_dt.datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    return buf.read(), fname
@bp.route("/api/export_excel")
def export_excel():
    """Export flagged items as an Excel workbook with per-source tabs."""
    # If in-memory list is empty (e.g. after page reload), try loading from DB.
    # Use get_session_items() so concurrent M365 + Google + File scans (each with
    # their own scan_id) are all included, not just the single latest scan_id.
    if not state.flagged_items and DB_OK:
        try:
            db = _get_db()
            if db:
                db_items = db.get_session_items()
                if db_items:
                    state.flagged_items[:] = db_items
        except Exception:
            pass
    try:
        xl_bytes, fname = _build_excel_bytes()
        return Response(
            xl_bytes,
            mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f"attachment; filename={fname}"}
        )
    except ImportError:
        return jsonify({"error": "openpyxl not installed — run: pip install openpyxl"}), 500
    except Exception as e:
        import traceback
        logger.error("export_excel error: %s\n%s", e, traceback.format_exc())
        return jsonify({"error": str(e)}), 500


# ── Article 30 report ─────────────────────────────────────────────────────────

def _build_article30_docx() -> tuple[bytes, str]:
    """Generate a GDPR Article 30 Register of Processing Activities as .docx.
    Returns (bytes, filename). Strings are translated using the active state.LANG dict."""
    try:
        from docx import Document as _Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise ImportError("python-docx not installed — run: pip install python-docx")

    import datetime as _dt

    # Translate helper — falls back to English default if key missing
    def L(key: str, default: str = "") -> str:
        return state.LANG.get(key, default)

    # ── Data ─────────────────────────────────────────────────────────────────
    db    = _get_db() if DB_OK else None
    stats   = db.get_stats() if db else {}
    items   = db.get_session_items() if db else list(state.flagged_items)
    trend   = db.get_trend(10) if db else []
    overdue = db.get_overdue_items(5) if db else []

    # Build account_id → display_name map from the scan's stored user_ids
    # This lets us resolve GUIDs and "Microsoft Konto" placeholders that
    # were stored in account_name before _resolve_display_name was applied.
    _acct_map: dict[str, str] = {}
    if db:
        try:
            scan_id = stats.get("scan_id") or db.latest_scan_id()
            if scan_id:
                row = db._connect().execute(
                    "SELECT user_count, options FROM scans WHERE id=?", (scan_id,)
                ).fetchone()
                # user_ids are stored in the options JSON column
                opts_json = json.loads(row["options"] or "{}") if row else {}
                for u in opts_json.get("user_ids", []):
                    uid  = u.get("id", "")
                    name = u.get("displayName", "")
                    if uid and name:
                        _acct_map[uid] = name
        except Exception:
            pass
    # Also seed from in-memory state.flagged_items (catches current scan not yet in DB)
    for item in state.flagged_items:
        aid  = item.get("account_id", "")
        name = item.get("account_name", "")
        if aid and name and not _GUID_RE.match(name.strip()):
            _acct_map.setdefault(aid, name)

    def _acct_label(item: dict) -> str:
        """Return the best human-readable account label for an item."""
        aid  = item.get("account_id", "")
        name = item.get("account_name", "")
        # Try the lookup map first (most reliable — built from scan user_ids)
        if aid and aid in _acct_map:
            return _acct_map[aid]
        # Fall back to stored name, resolving GUIDs/placeholders against account_id
        return _resolve_display_name(name, aid)
    overdue_ids = {o["id"] for o in overdue}

    now_str   = _dt.datetime.now().strftime("%Y-%m-%d %H:%M")
    date_str  = _dt.datetime.now().strftime("%Y-%m-%d")
    fname     = f"article30_{date_str}.docx"

    # Aggregate by source
    by_source: dict = {}
    for item in items:
        st = item.get("source_type", "other")
        by_source.setdefault(st, []).append(item)

    SOURCE_LABELS = {
        "email":      "Exchange (Outlook)",
        "onedrive":   "OneDrive",
        "sharepoint": "SharePoint",
        "teams":      "Teams",
        "gmail":      "Gmail",
        "gdrive":     "Google Drive",
        "local":      "Local files",
        "smb":        "Network / SMB",
    }

    # ── Colour palette ────────────────────────────────────────────────────────
    DARK_BLUE  = RGBColor(0x1F, 0x38, 0x64)
    MID_BLUE   = RGBColor(0x00, 0x78, 0xD4)
    LIGHT_GREY = RGBColor(0xF2, 0xF2, 0xF2)
    RED        = RGBColor(0xC0, 0x39, 0x2B)
    ORANGE     = RGBColor(0xC5, 0x5A, 0x00)
    WHITE      = RGBColor(0xFF, 0xFF, 0xFF)

    def _hex(c: RGBColor) -> str:
        return f"{c[0]:02X}{c[1]:02X}{c[2]:02X}"

    # ── Document setup ────────────────────────────────────────────────────────
    doc = _Document()
    doc.core_properties.title   = "GDPR Article 30 — Register of Processing Activities"
    doc.core_properties.author  = "GDPRScanner"
    doc.core_properties.subject = "GDPR Compliance"

    # Page margins — A4 with 2.5 cm margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Helper: set cell background ──────────────────────────────────────────
    def _cell_bg(cell, hex_color: str):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        tcPr.append(shd)

    def _set_cell_border(cell, **kwargs):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for edge in ("top", "left", "bottom", "right"):
            cfg = kwargs.get(edge, {})
            el  = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"),   cfg.get("val",   "single"))
            el.set(qn("w:sz"),    cfg.get("sz",    "4"))
            el.set(qn("w:space"), cfg.get("space", "0"))
            el.set(qn("w:color"), cfg.get("color", "CCCCCC"))
            tcBorders.append(el)
        tcPr.append(tcBorders)

    def _para(text: str = "", bold=False, size=11, color=None,
              align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6) -> object:
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        if text:
            run = p.add_run(text)
            run.bold      = bold
            run.font.size = Pt(size)
            if color:
                run.font.color.rgb = color
        return p

    def _heading(text: str, level: int = 1):
        p  = doc.add_heading(text, level=level)
        r  = p.runs[0] if p.runs else p.add_run(text)
        r.font.color.rgb = DARK_BLUE
        r.font.size      = Pt(16 if level == 1 else 13)
        r.bold           = True
        p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
        p.paragraph_format.space_after  = Pt(4)
        return p

    def _kv(label: str, value: str, label_width=2.5, bold=False, highlight=False):
        """Two-column key-value paragraph using a 2-cell table row."""
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        w_label = int(label_width * 1440)
        w_value = int((16.0 - label_width) * 1440 * 0.6)  # approx content width
        c1, c2 = tbl.rows[0].cells
        _cell_bg(c1, "FFF3E0" if highlight else "F2F2F2")
        _cell_bg(c2, "FFF3E0" if highlight else "FFFFFF")
        c1.width = Inches(label_width)
        c2.width = Inches(16.0 - label_width)
        p1 = c1.paragraphs[0]; p1.clear()
        r1 = p1.add_run(label); r1.bold = True; r1.font.size = Pt(10)
        p2 = c2.paragraphs[0]; p2.clear()
        r2 = p2.add_run(value); r2.font.size = Pt(10); r2.bold = bold
        if highlight:
            r1.font.color.rgb = RGBColor(0x6B, 0x00, 0x6B)
            r2.font.color.rgb = RGBColor(0x6B, 0x00, 0x6B)
        for cell in (c1, c2):
            _set_cell_border(cell, top={"color": "E0E0E0"}, bottom={"color": "E0E0E0"},
                             left={"color": "E0E0E0"}, right={"color": "E0E0E0"})
        return tbl

    # ── Cover page ────────────────────────────────────────────────────────────
    _para()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(40)
    r = title_p.add_run(L("a30_title", "GDPR Article 30"))
    r.bold = True; r.font.size = Pt(28); r.font.color.rgb = DARK_BLUE

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub_p.add_run(L("a30_subtitle", "Register of Processing Activities"))
    r2.font.size = Pt(16); r2.font.color.rgb = MID_BLUE

    _para()
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = meta_p.add_run(f"{L('a30_generated','Generated')}: {now_str}  ·  GDPRScanner")
    r3.font.size = Pt(10); r3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Divider line
    _para()
    div = doc.add_paragraph()
    div_fmt = div.paragraph_format
    div_fmt.space_after = Pt(20)
    pPr = div._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6")
    bot.set(qn("w:color"), _hex(MID_BLUE))
    pBdr.append(bot); pPr.append(pBdr)

    doc.add_page_break()

    # ── Section 1: Summary ────────────────────────────────────────────────────
    _heading(L("a30_s1", "1. Summary"))

    total_items    = len(items)
    total_cpr      = sum(i.get("cpr_count", 0) for i in items)
    special_items  = [i for i in items if i.get("special_category") and
                      i["special_category"] not in ("[]", "", None, [])]
    photo_items    = [i for i in items if i.get("face_count", 0) > 0]
    gps_items      = [i for i in items if "gps_location" in (i.get("special_category") or [])]
    exif_pii_items = [i for i in items if "exif_pii" in (i.get("special_category") or [])]
    unique_subj    = stats.get("unique_subjects", 0)
    total_scanned  = stats.get("total_scanned", 0)
    scan_date      = _dt.datetime.fromtimestamp(
        stats.get("started_at", 0)).strftime("%Y-%m-%d %H:%M") if stats.get("started_at") else "—"
    special_items  = [i for i in items if i.get("special_category") and
                      i["special_category"] not in ("[]", "", None, [])]

    _kv(L("a30_scan_date",       "Scan date"),                scan_date)
    _kv(L("a30_items_scanned",   "Items scanned"),            str(total_scanned))
    _kv(L("a30_flagged",         "Flagged items"),            str(total_items))
    _kv(L("a30_cpr_hits",        "Total CPR hits"),           str(total_cpr))
    _kv(L("a30_data_subjects",   "Estimated data subjects"),  str(unique_subj))
    _kv(L("a30_overdue",         "Overdue items (>5 yrs)"),   str(len(overdue_ids)))
    if gps_items:
        _kv(L("a30_gps_items", "Items with GPS location data (Art. 4 — location = personal data)"),
            str(len(gps_items)))
    if exif_pii_items:
        _kv(L("a30_exif_pii_items", "Items with EXIF PII (author, description, keywords)"),
            str(len(exif_pii_items)))
    if photo_items:
        total_faces = sum(i.get("face_count", 0) for i in photo_items)
        _kv(L("a30_photo_items", "Photos with detected faces (Art. 9 biometric)"),
            f"{len(photo_items)} items / {total_faces} faces")
        _para(L("a30_photo_note",
                "Photographs of identifiable persons are biometric data under Art. 9 GDPR. "
                "Retention requires a documented legal basis under Art. 9(2). "
                "For school photographs of pupils under 15, parental consent is required "
                "(Databeskyttelsesloven §6). See Datatilsynet guidance on school photography."),
              size=9, space_after=4)
    if special_items:
        _kv(L("a30_special_cat", "Art. 9 special category items"),
            str(len(special_items)))
        _para(L("a30_special_cat_note",
                "These items contain health, criminal, biometric, religious, ethnic, "
                "trade union, political, or sexual orientation data. "
                "An explicit legal basis (Art. 9(2)) and possibly a DPIA (Art. 35) is required."),
              size=9, space_after=4)

    _para()

    # Per-source breakdown table
    _para(L("a30_by_source", "Breakdown by source"), bold=True, size=11, space_before=10)

    src_tbl = doc.add_table(rows=1, cols=5)
    src_tbl.style = "Table Grid"
    hdr_cells = src_tbl.rows[0].cells
    for cell, txt in zip(hdr_cells, [L("a30_col_source","Source"), L("a30_col_items","Items"),
                                     L("a30_col_cpr","CPR hits"), L("a30_col_overdue","Overdue"),
                                     L("a30_col_special","Art. 9")]):
        _cell_bg(cell, _hex(DARK_BLUE))
        p = cell.paragraphs[0]; p.clear()
        r = p.add_run(txt); r.bold = True
        r.font.size = Pt(10); r.font.color.rgb = WHITE

    for src_key in ("email", "onedrive", "sharepoint", "teams", "gmail", "gdrive", "local", "smb"):
        src_items = by_source.get(src_key, [])
        if not src_items:
            continue
        row   = src_tbl.add_row().cells
        n_ov   = sum(1 for i in src_items if i.get("id") in overdue_ids)
        n_cpr  = sum(i.get("cpr_count", 0) for i in src_items)
        n_spec = sum(1 for i in src_items if i.get("special_category") and
                     i["special_category"] not in ("[]", "", None, []))
        for cell, val in zip(row, [
            SOURCE_LABELS.get(src_key, src_key),
            str(len(src_items)), str(n_cpr), str(n_ov),
            str(n_spec) if n_spec else "—"
        ]):
            p = cell.paragraphs[0]; p.clear()
            r = p.add_run(val); r.font.size = Pt(10)
            if val != "0" and cell == row[3]:
                r.font.color.rgb = ORANGE
            if n_spec and cell == row[4]:
                r.font.color.rgb = RGBColor(0x7B, 0x00, 0x82)
                r.bold = True

    # ── Section 2: Data categories ────────────────────────────────────────────
    doc.add_page_break()
    _heading(L("a30_s2", "2. Personal Data Categories Identified"))

    _para(L("a30_s2_intro", "The following categories of personal data were detected during scanning."),
          size=10, space_after=8)

    # Aggregate PII from DB or from items
    pii_totals: dict = {}
    if db:
        rows = db._connect().execute(
            """SELECT pii_type, SUM(hit_count) FROM pii_hits
               WHERE scan_id=? GROUP BY pii_type""",
            (stats.get("scan_id") or db.latest_scan_id() or 0,)
        ).fetchall()
        for pii_type, count in rows:
            pii_totals[pii_type] = count

    PII_LABELS = {
        "PHONE":        L("a30_pii_phone",        "Phone numbers"),
        "EMAIL":        L("a30_pii_email",        "Email addresses"),
        "IBAN":         L("a30_pii_iban",         "IBAN bank numbers"),
        "BANK_ACCOUNT": L("a30_pii_bank",         "Bank account numbers"),
        "NAME":         L("a30_pii_name",         "Personal names (NER)"),
        "ADDRESS":      L("a30_pii_address",      "Addresses (NER)"),
        "ORG":          L("a30_pii_org",          "Organisations (NER)"),
    }

    pii_tbl = doc.add_table(rows=1, cols=3)
    pii_tbl.style = "Table Grid"
    for cell, txt in zip(pii_tbl.rows[0].cells,
                          [L("a30_col_category","Data category"), L("a30_col_count","Count"), L("a30_col_gdpr_class","GDPR classification")]):
        _cell_bg(cell, _hex(DARK_BLUE))
        p = cell.paragraphs[0]; p.clear()
        r = p.add_run(txt); r.bold = True
        r.font.size = Pt(10); r.font.color.rgb = WHITE

    # CPR row first — always
    cpr_row = pii_tbl.add_row().cells
    for cell, val in zip(cpr_row, [L("a30_cpr_label", "CPR numbers (Danish personal ID)"), str(total_cpr),
                                    L("a30_cpr_class", "Art. 9 — national identifier")]):
        p = cell.paragraphs[0]; p.clear()
        r = p.add_run(val); r.font.size = Pt(10)
        _cpr_class = L("a30_cpr_class", "Art. 9 — national identifier")
        if val == _cpr_class:
            r.font.color.rgb = RED; r.bold = True

    for pii_type, label in PII_LABELS.items():
        count = pii_totals.get(pii_type, 0)
        if not count:
            continue
        cls = L("a30_pii_class_9", "Art. 9 — health/sensitive") if pii_type in ("NAME", "ADDRESS") else L("a30_pii_class_4", "Art. 4 — personal data")
        row = pii_tbl.add_row().cells
        for cell, val in zip(row, [label, str(count), cls]):
            p = cell.paragraphs[0]; p.clear()
            r = p.add_run(val); r.font.size = Pt(10)

    # ── Section 3: Data inventory ─────────────────────────────────────────────
    doc.add_page_break()
    _heading(L("a30_s3", "3. Data Inventory"))

    _para(L("a30_s3_intro", "All flagged items are listed below with location, retention status, and compliance disposition."),
          size=10, space_after=8)

    # Split by user role for separate presentation
    student_items = [i for i in items if i.get("user_role") == "student"]
    staff_items   = [i for i in items if i.get("user_role") != "student"]

    _disp_map = {
        "unreviewed":       L("a30_disp_unreviewed",      "Unreviewed"),
        "retain-legal":     L("a30_disp_retain_legal",    "Retain — Legal obligation"),
        "retain-legitimate": L("a30_disp_retain_legit",   "Retain — Legitimate interest"),
        "retain-contract":  L("a30_disp_retain_contract", "Retain — Contract"),
        "delete-scheduled": L("a30_disp_delete_sched",    "Delete — Scheduled"),
        "deleted":          L("a30_disp_deleted",         "Deleted"),
        "personal-use":     L("a30_disp_personal_use",    "Personal use — out of GDPR scope (Art. 2(2)(c))"),
    }

    def _inv_table(tbl_items: list):
        tbl = doc.add_table(rows=1, cols=6)
        tbl.style = "Table Grid"
        col_hdrs = [L("a30_col_name","Name / Subject"), L("a30_col_source","Source"),
                    L("a30_col_account","Account"), L("a30_col_modified","Modified"),
                    L("a30_col_cpr_short","CPR"), L("a30_col_disp","Disposition")]
        for cell, txt in zip(tbl.rows[0].cells, col_hdrs):
            _cell_bg(cell, _hex(DARK_BLUE))
            p = cell.paragraphs[0]; p.clear()
            r = p.add_run(txt); r.bold = True
            r.font.size = Pt(9); r.font.color.rgb = WHITE
        sorted_tbl = sorted(tbl_items,
            key=lambda x: (0 if x.get("id") in overdue_ids else 1, -x.get("cpr_count", 0)))
        for idx, item in enumerate(sorted_tbl[:500]):
            disp_rec = db.get_disposition(item["id"]) if db else None
            raw_disp = disp_rec.get("status", "unreviewed") if disp_rec else "unreviewed"
            disp_str = _disp_map.get(raw_disp, raw_disp.replace("-", " ").title())
            is_ov    = item.get("id") in overdue_ids
            row = tbl.add_row().cells
            vals = [
                (item.get("name", "")[:60] + ("…" if len(item.get("name", "")) > 60 else "")),
                SOURCE_LABELS.get(item.get("source_type", ""), item.get("source_type", "")),
                _acct_label(item),
                item.get("modified", ""),
                str(item.get("cpr_count", 0)),
                disp_str,
            ]
            bg = "FFF8F0" if is_ov else ("FFFFFF" if idx % 2 == 0 else "F8F8F8")
            for cell, val in zip(row, vals):
                _cell_bg(cell, bg)
                p = cell.paragraphs[0]; p.clear()
                r = p.add_run(val); r.font.size = Pt(8)
                if is_ov and cell == row[3]:
                    r.font.color.rgb = ORANGE
        if len(tbl_items) > 500:
            _para(f"… {len(tbl_items) - 500} {L('a30_more_items', 'additional items not shown.')}",
                  size=9, color=RGBColor(0x88, 0x88, 0x88), space_before=4)

    if staff_items:
        if student_items:
            _para(L("a30_inv_staff", "👔 Staff / Faculty"), bold=True, size=11, space_before=6, space_after=4)
        _inv_table(staff_items)

    if student_items:
        _para(L("a30_inv_students", "🎓 Students"), bold=True, size=11, space_before=14, space_after=2)
        _para(L("a30_student_consent_note",
                "Note: Student accounts in Danish folkeskole (pupils under age 15) require parental "
                "consent for processing of personal data under Databeskyttelsesloven §6. "
                "Items in student accounts must not be auto-deleted — any action requires "
                "review by school administration and, for pupils under 15, notification of parents "
                "or guardians as rights holders under GDPR Article 8."),
              size=9, color=RGBColor(0x88, 0x44, 0x00), space_after=6)
        _inv_table(student_items)

    # ── Section 4: Retention analysis ────────────────────────────────────────
    if overdue:
        doc.add_page_break()
        _heading(L("a30_s4", "4. Retention Analysis"))

        _para(L("a30_s4_intro", "The following items exceed the 5-year retention threshold and should be reviewed for deletion under GDPR Article 5(1)(e) — storage limitation."),
              size=10, space_after=8)

        ret_tbl = doc.add_table(rows=1, cols=5)
        ret_tbl.style = "Table Grid"
        for cell, txt in zip(ret_tbl.rows[0].cells,
                              [L("a30_col_name","Name"), L("a30_col_source","Source"), L("a30_col_account","Account"), L("a30_col_modified","Modified"), L("a30_col_cpr","CPR hits")]):
            _cell_bg(cell, _hex(ORANGE))
            p = cell.paragraphs[0]; p.clear()
            r = p.add_run(txt); r.bold = True
            r.font.size = Pt(9); r.font.color.rgb = WHITE

        for item in overdue[:200]:
            row = ret_tbl.add_row().cells
            for cell, val in zip(row, [
                item.get("name", "")[:55],
                SOURCE_LABELS.get(item.get("source_type", ""), ""),
                _acct_label(item),
                item.get("modified", ""),
                str(item.get("cpr_count", 0)),
            ]):
                p = cell.paragraphs[0]; p.clear()
                r = p.add_run(val); r.font.size = Pt(8)

    # ── Section 5: Scan history ───────────────────────────────────────────────
    if trend:
        sec_num = "5" if overdue else "4"
        doc.add_page_break()
        _heading(f"{sec_num}. {L('a30_s5','Compliance Trend').split('. ',1)[-1]}")

        _para(L("a30_s5_intro", "Flagged item counts over the last scans (most recent first)."),
              size=10, space_after=8)

        trend_tbl = doc.add_table(rows=1, cols=4)
        trend_tbl.style = "Table Grid"
        for cell, txt in zip(trend_tbl.rows[0].cells,
                              [L("a30_col_scan_date","Scan date"), L("a30_col_flagged","Flagged"), L("a30_col_overdue","Overdue"), L("a30_col_scan_type","Scan type")]):
            _cell_bg(cell, _hex(DARK_BLUE))
            p = cell.paragraphs[0]; p.clear()
            r = p.add_run(txt); r.bold = True
            r.font.size = Pt(9); r.font.color.rgb = WHITE

        for t in reversed(trend):
            row = trend_tbl.add_row().cells
            for cell, val in zip(row, [
                t.get("scan_date", ""),
                str(t.get("flagged_count", 0)),
                str(t.get("overdue_count", 0)),
                L("a30_scan_delta", "Delta") if t.get("delta") else L("a30_scan_full", "Full"),
            ]):
                p = cell.paragraphs[0]; p.clear()
                r = p.add_run(val); r.font.size = Pt(9)

    # ── Section: Deletion audit log ───────────────────────────────────────────
    del_log   = db.get_deletion_log(limit=500) if db else []
    del_stats = db.deletion_log_stats() if db else {}

    # Running section counter — starts at 3 (summary, categories, inventory always present)
    last_sec  = 3
    last_sec += 1 if overdue  else 0   # retention analysis
    last_sec += 1 if trend    else 0   # compliance trend

    if del_log:
        del_sec   = last_sec
        last_sec += 1
        doc.add_page_break()
        _heading(f"{del_sec}. {L('a30_s_dellog', 'Deletion Audit Log')}")

        _para(L("a30_dellog_intro",
                f"A total of {del_stats.get('total', len(del_log))} item(s) containing personal data "
                f"have been deleted via GDPRScanner. "
                f"CPR hits removed: {del_stats.get('cpr_hits_deleted', 0)}. "
                f"This log satisfies the accountability obligation under GDPR Article 5(2)."),
              size=10, space_after=8)

        # Summary by reason
        by_reason = del_stats.get("by_reason", {})
        if by_reason:
            _para(L("a30_dellog_by_reason", "Deletions by reason"), bold=True, size=10, space_before=4, space_after=4)
            reason_tbl = doc.add_table(rows=1, cols=2)
            reason_tbl.style = "Table Grid"
            for cell, txt in zip(reason_tbl.rows[0].cells,
                                  [L("a30_col_reason", "Reason"), L("a30_col_count", "Count")]):
                _cell_bg(cell, _hex(DARK_BLUE))
                p = cell.paragraphs[0]; p.clear()
                r = p.add_run(txt); r.bold = True
                r.font.size = Pt(9); r.font.color.rgb = WHITE
            REASON_LABELS = {
                "manual":               L("a30_reason_manual",    "Manual (individual card delete)"),
                "bulk":                 L("a30_reason_bulk",       "Bulk delete"),
                "retention":            L("a30_reason_retention",  "Retention policy enforcement"),
                "data-subject-request": L("a30_reason_dsr",        "Data subject erasure request (Art. 17)"),
            }
            for reason, count in sorted(by_reason.items()):
                row = reason_tbl.add_row().cells
                for cell, val in zip(row, [REASON_LABELS.get(reason, reason), str(count)]):
                    p = cell.paragraphs[0]; p.clear()
                    r = p.add_run(val); r.font.size = Pt(9)

        # Full log table
        _para(L("a30_dellog_records", "Deletion records"), bold=True, size=10, space_before=10, space_after=4)
        log_tbl = doc.add_table(rows=1, cols=7)
        log_tbl.style = "Table Grid"
        for cell, txt in zip(log_tbl.rows[0].cells, [
            L("a30_col_deleted_at",  "Deleted at"),
            L("a30_col_name",        "Name"),
            L("a30_col_source",      "Source"),
            L("a30_col_account",     "Account"),
            L("a30_col_cpr",         "CPR hits"),
            L("a30_col_reason",      "Reason"),
            L("a30_col_deleted_by",  "Deleted by"),
        ]):
            _cell_bg(cell, _hex(DARK_BLUE))
            p = cell.paragraphs[0]; p.clear()
            r = p.add_run(txt); r.bold = True
            r.font.size = Pt(8); r.font.color.rgb = WHITE

        for idx, entry in enumerate(del_log):
            ts  = _dt.datetime.fromtimestamp(entry.get("deleted_at", 0)).strftime("%Y-%m-%d %H:%M")
            bg  = "FFFFFF" if idx % 2 == 0 else "F8F8F8"
            row = log_tbl.add_row().cells
            for cell, val in zip(row, [
                ts,
                entry.get("item_name", "")[:40],
                SOURCE_LABELS.get(entry.get("source_type", ""), entry.get("source_type", "")),
                _acct_map.get(entry.get("account_id", "")) or _resolve_display_name(entry.get("account_name", ""), entry.get("account_id", "")),
                str(entry.get("cpr_count", 0)),
                REASON_LABELS.get(entry.get("reason", ""), entry.get("reason", "")),
                entry.get("deleted_by", "") or "—",
            ]):
                _cell_bg(cell, bg)
                p = cell.paragraphs[0]; p.clear()
                r = p.add_run(val); r.font.size = Pt(7)

    # ── Section: Article 9 special categories ────────────────────────────────
    if special_items:
        last_sec += 1
        doc.add_page_break()
        _heading(f"{last_sec}. {L('a30_s_special', 'Special Category Data (Article 9)')}")

        _para(L("a30_special_intro",
                f"{len(special_items)} item(s) were detected as containing special category "
                f"data under GDPR Article 9. These require an explicit legal basis beyond "
                f"Article 6, and processing should be covered by a Data Protection Impact "
                f"Assessment (DPIA) under Article 35."),
              size=10, space_after=8)

        # Category breakdown table
        from collections import Counter as _Counter
        cat_counts: dict = _Counter()
        for item in special_items:
            sc = item.get("special_category", [])
            if isinstance(sc, str):
                import json as _scjson
                try:
                    sc = _scjson.loads(sc)
                except Exception:
                    sc = []
            for c in sc:
                cat_counts[c] += 1

        if cat_counts:
            _para(L("a30_special_by_cat", "Detected categories"), bold=True, size=10,
                  space_before=4, space_after=4)
            cat_tbl = doc.add_table(rows=1, cols=2)
            cat_tbl.style = "Table Grid"
            for cell, txt in zip(cat_tbl.rows[0].cells,
                                  [L("a30_col_category", "Category"),
                                   L("a30_col_count", "Items")]):
                _cell_bg(cell, _hex(DARK_BLUE))
                p = cell.paragraphs[0]; p.clear()
                r = p.add_run(txt); r.bold = True
                r.font.size = Pt(9); r.font.color.rgb = WHITE
            CAT_LABELS = {
                "health":           L("a30_cat_health",    "Health data (Art. 9)"),
                "mental_health":    L("a30_cat_mental",    "Mental health (Art. 9)"),
                "criminal":         L("a30_cat_criminal",  "Criminal records (Art. 10)"),
                "trade_union":      L("a30_cat_union",     "Trade union membership (Art. 9)"),
                "religion":         L("a30_cat_religion",  "Religious beliefs (Art. 9)"),
                "ethnicity":        L("a30_cat_ethnicity", "Racial/ethnic origin (Art. 9)"),
                "political":        L("a30_cat_political", "Political opinions (Art. 9)"),
                "biometric":        L("a30_cat_biometric", "Biometric data (Art. 9)"),
                "sexual_orientation": L("a30_cat_sexual",  "Sexual orientation (Art. 9)"),
            }
            for cat, count in sorted(cat_counts.items(), key=lambda x: -x[1]):
                row = cat_tbl.add_row().cells
                for cell, val in zip(row, [CAT_LABELS.get(cat, cat), str(count)]):
                    p = cell.paragraphs[0]; p.clear()
                    r = p.add_run(val); r.font.size = Pt(9)

        # Item list (capped at 50)
        _para(L("a30_special_items", "Affected items (up to 50)"), bold=True, size=10,
              space_before=10, space_after=4)
        sc_tbl = doc.add_table(rows=1, cols=5)
        sc_tbl.style = "Table Grid"
        for cell, txt in zip(sc_tbl.rows[0].cells, [
            L("a30_col_name",     "Name"),
            L("a30_col_account",  "Account"),
            L("a30_col_source",   "Source"),
            L("a30_col_category", "Category"),
            L("a30_col_cpr",      "CPR hits"),
        ]):
            _cell_bg(cell, _hex(DARK_BLUE))
            p = cell.paragraphs[0]; p.clear()
            r = p.add_run(txt); r.bold = True
            r.font.size = Pt(8); r.font.color.rgb = WHITE

        for idx, item in enumerate(special_items[:50]):
            bg = "FFFFFF" if idx % 2 == 0 else "FFF0F8"
            sc = item.get("special_category", [])
            if isinstance(sc, str):
                try:
                    import json as _scj2; sc = _scj2.loads(sc)
                except Exception:
                    sc = []
            row = sc_tbl.add_row().cells
            for cell, val in zip(row, [
                item.get("name", "")[:35],
                _acct_map.get(item.get("account_id", "")) or item.get("account_name", ""),
                SOURCE_LABELS.get(item.get("source_type", ""), item.get("source_type", "")),
                ", ".join(CAT_LABELS.get(c, c) for c in sc)[:45],
                str(item.get("cpr_count", 0)),
            ]):
                _cell_bg(cell, bg)
                p = cell.paragraphs[0]; p.clear()
                r = p.add_run(val); r.font.size = Pt(7)

    # ── Section: Photographs / biometric data (#9) ───────────────────────────
    if photo_items:
        last_sec += 1
        doc.add_page_break()
        _heading(f"{last_sec}. {L('a30_s_photos', 'Photographs and Biometric Data (Article 9)')}")

        total_faces = sum(i.get("face_count", 0) for i in photo_items)
        _para(L("a30_photo_intro",
                f"{len(photo_items)} image file(s) containing {total_faces} detected face(s) "
                f"were found in the scan. Photographs of identifiable persons constitute "
                f"biometric data under GDPR Article 9 and are subject to the same "
                f"heightened protection as health or criminal records data."),
              size=10, space_after=8)

        _para(L("a30_photo_guidance", "Retention guidance"), bold=True, size=10,
              space_before=4, space_after=4)
        for line in [
            L("a30_photo_g1",
              "Photos may only be retained while the original purpose remains valid "
              "(Art. 5(1)(b) — purpose limitation)."),
            L("a30_photo_g2",
              "Pupils under 15 require parental consent (Databeskyttelsesloven §6). "
              "Consent must be freely given, specific, and documented."),
            L("a30_photo_g3",
              "Photos on public-facing websites must be removed promptly after a person "
              "leaves the organisation or withdraws consent (Art. 17 — right to erasure)."),
            L("a30_photo_g4",
              "Historical/archive use may justify longer retention under Art. 89 only "
              "with specific safeguards and case-by-case assessment."),
        ]:
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(line); r.font.size = Pt(9)

        # GPS items sub-section
        if gps_items:
            _para(L("a30_gps_title", "Items with GPS location data"), bold=True, size=10,
                  space_before=10, space_after=4)
            _para(L("a30_gps_intro",
                    "The following files contain GPS coordinates embedded in EXIF metadata. "
                    "Location data constitutes personal data under Art. 4 GDPR. For photos of children "
                    "or staff, GPS data may reveal sensitive patterns (home address, health institution, "
                    "religious site). Consider stripping EXIF before sharing or publishing."),
                  size=9, space_after=6)
            gps_tbl = doc.add_table(rows=1, cols=4)
            gps_tbl.style = "Table Grid"
            for cell, txt in zip(gps_tbl.rows[0].cells, [
                L("a30_col_name", "Name"),
                L("a30_gps_col_lat", "Latitude"),
                L("a30_gps_col_lon", "Longitude"),
                L("a30_col_date", "Modified"),
            ]):
                _cell_bg(cell, _hex(DARK_BLUE))
                p = cell.paragraphs[0]; p.clear()
                r = p.add_run(txt); r.bold = True
                r.font.size = Pt(8); r.font.color.rgb = WHITE
            for idx, item in enumerate(gps_items[:50]):
                bg = "FFFFFF" if idx % 2 == 0 else "E8F7FF"
                row = gps_tbl.add_row().cells
                exif = item.get("exif") or {}
                gps  = exif.get("gps") or {}
                for cell, val in zip(row, [
                    item.get("name", "")[:40],
                    str(gps.get("lat", ""))[:12],
                    str(gps.get("lon", ""))[:12],
                    item.get("modified", ""),
                ]):
                    _cell_bg(cell, bg)
                    p = cell.paragraphs[0]; p.clear()
                    r = p.add_run(val); r.font.size = Pt(7)

        # Photo item list (capped at 50)
        _para(L("a30_photo_items", "Detected photo items (up to 50)"), bold=True, size=10,
              space_before=10, space_after=4)
        ph_tbl = doc.add_table(rows=1, cols=6)
        ph_tbl.style = "Table Grid"
        for cell, txt in zip(ph_tbl.rows[0].cells, [
            L("a30_col_name",    "Name"),
            L("a30_col_account", "Account"),
            L("a30_col_source",  "Source"),
            L("a30_photo_col_faces", "Faces"),
            L("a30_gps_col",     "GPS"),
            L("a30_col_date",    "Modified"),
        ]):
            _cell_bg(cell, _hex(DARK_BLUE))
            p = cell.paragraphs[0]; p.clear()
            r = p.add_run(txt); r.bold = True
            r.font.size = Pt(8); r.font.color.rgb = WHITE

        for idx, item in enumerate(photo_items[:50]):
            bg = "FFFFFF" if idx % 2 == 0 else "E8F7FF"
            row = ph_tbl.add_row().cells
            for cell, val in zip(row, [
                item.get("name", "")[:40],
                _acct_map.get(item.get("account_id", "")) or item.get("account_name", ""),
                SOURCE_LABELS.get(item.get("source_type", ""), item.get("source_type", "")),
                str(item.get("face_count", 0)),
                "✔" if (item.get("exif") or {}).get("gps") else "",
                item.get("modified", ""),
            ]):
                _cell_bg(cell, bg)
                p = cell.paragraphs[0]; p.clear()
                r = p.add_run(val); r.font.size = Pt(7)

    # ── Section: Methodology ─────────────────────────────────────────────────
    # last_sec already reflects all optional sections that were added above
    doc.add_page_break()
    _heading(f"{last_sec}. {L('a30_s6_short', 'Methodology and Legal Basis')}")

    _para(L("a30_method_title", "Scanning methodology"), bold=True, size=11, space_before=6, space_after=4)
    for line in [
        L("a30_method_1", "CPR numbers are detected using pattern matching against the official Danish CPR format (DDMMYY-XXXX)."),
        L("a30_method_2", "Additional personal data (phone numbers, email addresses, IBANs, bank accounts, names, addresses, and organisations) is detected using regular expressions and spaCy NER."),
        L("a30_method_3", "CPR numbers stored in this document's database are SHA-256 hashed and never stored in plaintext."),
        L("a30_method_4", "Scanning covers Exchange mailboxes (all folders including Sent Items), OneDrive, SharePoint, and Microsoft Teams channel files via the Microsoft Graph API. When connected, Google Workspace scanning covers Gmail and Google Drive via a service account with domain-wide delegation. Local and network (SMB) file shares are scanned directly."),
        L("a30_method_5", "When photo scanning is enabled, image files are analysed using OpenCV Haar cascade face detection to identify photographs of persons (Art. 9 biometric data)."),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(line); r.font.size = Pt(10)

    _para(L("a30_gdpr_title", "GDPR Articles referenced"), bold=True, size=11, space_before=10, space_after=4)
    for line in [
        L("a30_gdpr_1", "Article 5(1)(c) — Data minimisation: only necessary data should be retained"),
        L("a30_gdpr_2", "Article 5(1)(e) — Storage limitation: data must not be kept longer than necessary"),
        L("a30_gdpr_3", "Article 9 — Special categories: health, criminal, trade union, and similar data require explicit legal basis"),
        L("a30_gdpr_4", "Article 15 — Right of access: data subjects may request information about their data"),
        L("a30_gdpr_5", "Article 17 — Right to erasure: data subjects may request deletion"),
        L("a30_gdpr_6", "Article 30 — Records of processing activities: this document satisfies the obligation"),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(line); r.font.size = Pt(10)

    _para(f"{L('a30_generated','Generated')}: {now_str}  ·  GDPRScanner  ·  {L('a30_confidential','Confidential — GDPR compliance document')}",
          size=9, color=RGBColor(0x88, 0x88, 0x88), align=WD_ALIGN_PARAGRAPH.CENTER, space_before=20)

    # ── Serialise ─────────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read(), fname


@bp.route("/api/export_article30")
def export_article30():
    """Generate and return an Article 30 Word document."""
    # Pre-populate in-memory list from DB session so _build_article30_docx()
    # has state.flagged_items available for the account-name seed (line ~318).
    if not state.flagged_items and DB_OK:
        try:
            db = _get_db()
            if db:
                db_items = db.get_session_items()
                if db_items:
                    state.flagged_items[:] = db_items
        except Exception:
            pass
    if not state.flagged_items:
        return jsonify({"error": "No results to export — run a scan first"}), 400
    try:
        docx_bytes, fname = _build_article30_docx()
        return Response(
            docx_bytes,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={fname}"}
        )
    except ImportError as e:
        return jsonify({"error": str(e)}), 500
    except Exception as e:
        import traceback
        logger.error("export_article30 error: %s\n%s", e, traceback.format_exc())
        return jsonify({"error": str(e)}), 500


def delete_item():
    """Delete a single flagged item. Returns {ok, error}."""
    if not state.connector:
        return jsonify({"ok": False, "error": "not authenticated"}), 401
    data        = request.get_json() or {}
    item_id     = data.get("id", "")
    source_type = data.get("source_type", "")
    account_id  = data.get("account_id", "") or "me"
    drive_id    = data.get("drive_id", "")

    if not item_id:
        return jsonify({"ok": False, "error": "id required"}), 400

    try:
        if source_type == "email":
            ok = state.connector.delete_message(account_id, item_id)
        elif drive_id:
            ok = state.connector.delete_drive_item(drive_id, item_id)
        else:
            ok = state.connector.delete_drive_item_for_user(account_id, item_id)

        if ok or ok is False:  # False = already gone, treat as success
            # Retrieve full item for audit log before removing it
            item_meta = next((x for x in state.flagged_items if x.get("id") == item_id), {})
            state.flagged_items = [x for x in state.flagged_items if x.get("id") != item_id]
            _db = _get_db() if DB_OK else None
            if _db:
                try:
                    _db.log_deletion(item_meta or {"id": item_id, "source_type": source_type},
                                     reason="manual")
                    _db.delete_item_record(item_id)
                except Exception: pass
            return jsonify({"ok": True})
        return jsonify({"ok": False, "error": "Delete returned unexpected result"})
    except M365PermissionError:
        return jsonify({"ok": False, "error":
            "Permission denied (403) — deletion requires Mail.ReadWrite / Files.ReadWrite.All / Sites.ReadWrite.All. "
            "Go to Azure → App registrations → API permissions → add these and Grant admin consent."})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)})


@bp.route("/api/delete_bulk", methods=["POST"])
def delete_bulk():
    """Delete multiple items matching criteria. Streams progress as SSE."""
    if not state.connector:
        return jsonify({"ok": False, "error": "not authenticated"}), 401
    data    = request.get_json() or {}
    item_ids = data.get("ids", [])   # explicit list of ids, or empty = use filters
    filters  = data.get("filters", {})
    del_reason = data.get("reason", "bulk")  # manual/bulk/retention/data-subject-request

    # Build target list
    if item_ids:
        targets = [x for x in state.flagged_items if x.get("id") in set(item_ids)]
    else:
        targets = list(state.flagged_items)
        # Apply filters
        if filters.get("source_type"):
            targets = [x for x in targets if x.get("source_type") == filters["source_type"]]
        if filters.get("min_cpr"):
            targets = [x for x in targets if x.get("cpr_count", 0) >= int(filters["min_cpr"])]
        if filters.get("older_than_date"):
            targets = [x for x in targets if x.get("modified", "9999") <= filters["older_than_date"]]

    deleted_ids  = []
    failed_items = []

    for item in targets:
        iid         = item.get("id", "")
        source_type = item.get("source_type", "")
        account_id  = item.get("account_id", "") or "me"
        drive_id    = item.get("drive_id", "")
        try:
            if source_type == "email":
                state.connector.delete_message(account_id, iid)
            elif drive_id:
                state.connector.delete_drive_item(drive_id, iid)
            else:
                state.connector.delete_drive_item_for_user(account_id, iid)
            deleted_ids.append(iid)
        except M365PermissionError:
            failed_items.append({"id": iid, "name": item.get("name", ""), "error":
                "403 — requires Mail.ReadWrite / Files.ReadWrite.All / Sites.ReadWrite.All (Azure admin consent)"})
        except Exception as e:
            failed_items.append({"id": iid, "name": item.get("name", ""), "error": str(e)})

    # Build id->item map for audit log
    _deleted_meta = {x.get("id"): x for x in targets if x.get("id") in set(deleted_ids)}
    state.flagged_items = [x for x in state.flagged_items if x.get("id") not in set(deleted_ids)]
    _db = _get_db() if DB_OK else None
    if _db:
        for _did in deleted_ids:
            try:
                _db.log_deletion(_deleted_meta.get(_did, {"id": _did}), reason=del_reason)
                _db.delete_item_record(_did)
            except Exception: pass

    return jsonify({
        "ok":      True,
        "deleted": len(deleted_ids),
        "failed":  len(failed_items),
        "errors":  failed_items[:10],  # cap error list
    })


