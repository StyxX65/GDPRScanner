#!/usr/bin/env python3
"""
GDPRScanner — Scan Exchange, OneDrive, SharePoint & Teams for CPR numbers.
Run with:  python gdpr_scanner.py [--port 5100]

Requires:
    pip install flask msal requests pillow

Optional (better PDF scanning):
    pip install pymupdf
"""

import argparse
import base64
import hashlib
import io
import json
import socket
import logging
import logging.handlers
import os
import queue
from collections import deque
import re
import sys
import tempfile
import concurrent.futures
import threading
import time
from pathlib import Path

# ── Logging setup ─────────────────────────────────────────────────────────────
def _configure_logging() -> None:
    """Configure root logger with console + rotating file handler.

    Called once at startup before any module-level code uses the logger.
    Idempotent — skipped if handlers are already attached (e.g. under pytest).
    """
    _data_dir = Path.home() / ".gdprscanner"
    _data_dir.mkdir(exist_ok=True)
    root = logging.getLogger()
    if root.handlers:
        return
    root.setLevel(logging.INFO)
    _fmt = logging.Formatter(
        "%(asctime)s %(levelname)-8s %(name)s — %(message)s",
        datefmt="%Y-%m-%d %H:%M:%S",
    )
    _sh = logging.StreamHandler()
    _sh.setFormatter(_fmt)
    _fh = logging.handlers.RotatingFileHandler(
        _data_dir / "gdpr_scanner.log",
        maxBytes=2 * 1024 * 1024,
        backupCount=3,
        encoding="utf-8",
    )
    _fh.setFormatter(_fmt)
    root.addHandler(_sh)
    root.addHandler(_fh)
    # Suppress noisy third-party loggers
    logging.getLogger("pdfminer").setLevel(logging.ERROR)
    logging.getLogger("pdfplumber").setLevel(logging.ERROR)
    logging.getLogger("werkzeug").setLevel(logging.WARNING)

_configure_logging()
logger = logging.getLogger(__name__)

# ── Module identity fix ───────────────────────────────────────────────────────
# When run as `python gdpr_scanner.py`, Python loads this module as `__main__`.
# When scan_scheduler.py does `import gdpr_scanner`, Python would load a SECOND
# copy with its own _sse_queues, broadcast(), etc. — so scheduled scan events
# would never reach the browser's SSE connection.
# Fix: register this module under both names so all imports share one instance.
if __name__ == "__main__":
    sys.modules["gdpr_scanner"] = sys.modules[__name__]


# ── One-time migration shim: rename ~/.m365_scanner_* → ~/.gdpr_scanner_* ────
# Runs silently on first startup after upgrading from v1.5.x.
# Safe to re-run — only moves files that don't already exist at the new path.
def _migrate_legacy_files():
    _LEGACY = [
        (".m365_scanner_config.json",       ".gdpr_scanner_config.json"),
        (".m365_scanner.db",                ".gdpr_scanner.db"),
        (".m365_scanner_token.json",        ".gdpr_scanner_token.json"),
        (".m365_scanner_delta.json",        ".gdpr_scanner_delta.json"),
        (".m365_scanner_settings.json",     ".gdpr_scanner_settings.json"),
        (".m365_scanner_smtp.json",         ".gdpr_scanner_smtp.json"),
        (".m365_scanner_role_overrides.json",".gdpr_scanner_role_overrides.json"),
        (".m365_scanner_file_sources.json", ".gdpr_scanner_file_sources.json"),
        (".m365_scanner_machine_id",        ".gdpr_scanner_machine_id"),
        (".m365_scanner_checkpoint.json",   ".gdpr_scanner_checkpoint.json"),
        (".m365_scanner_schedule.json",     ".gdpr_scanner_schedule.json"),
        (".m365_scanner_msal_cache.bin",    ".gdpr_scanner_msal_cache.bin"),
        (".m365_scanner_lang",              ".gdpr_scanner_lang"),
    ]
    home = Path.home()
    for old_name, new_name in _LEGACY:
        old = home / old_name
        new = home / new_name
        if old.exists() and not new.exists():
            try:
                old.rename(new)
                logger.info("[migrate] %s → %s", old_name, new_name)
            except Exception as _e:
                logger.warning("[migrate] Could not rename %s: %s", old_name, _e)

_migrate_legacy_files()

# ── One-time migration: move ~/.gdpr_scanner_* → ~/.gdprscanner/ ────────────
# Runs silently on first startup after upgrading from v1.6.2 or earlier.
def _migrate_to_data_dir():
    _DATA_DIR = Path.home() / ".gdprscanner"
    _DATA_DIR.mkdir(exist_ok=True)
    _MOVES = [
        (".gdpr_scanner_config.json",        "config.json"),
        (".gdpr_scanner.db",                 "scanner.db"),
        (".gdpr_scanner_token.json",         "token.json"),
        (".gdpr_scanner_delta.json",         "delta.json"),
        (".gdpr_scanner_settings.json",      "settings.json"),
        (".gdpr_scanner_smtp.json",          "smtp.json"),
        (".gdpr_scanner_role_overrides.json","role_overrides.json"),
        (".gdpr_scanner_file_sources.json",  "file_sources.json"),
        (".gdpr_scanner_machine_id",         "machine_id"),
        (".gdpr_scanner_checkpoint.json",    "checkpoint.json"),
        (".gdpr_scanner_schedule.json",      "schedule.json"),
        (".gdpr_scanner_msal_cache.bin",     "msal_cache.bin"),
        (".gdpr_scanner_lang",               "lang"),
        (".gdpr_scanner_google.json",        "google.json"),
        (".gdpr_scanner_google_sa.json",     "google_sa.json"),
        (".gdpr_scanner_src_toggles.json",   "src_toggles.json"),
    ]
    home = Path.home()
    for old_name, new_name in _MOVES:
        old = home / old_name
        new = _DATA_DIR / new_name
        if old.exists() and not new.exists():
            try:
                old.rename(new)
                logger.info("[migrate] ~/%s → ~/.gdprscanner/%s", old_name, new_name)
            except Exception as _e:
                logger.warning("[migrate] Could not move %s: %s", old_name, _e)

_migrate_to_data_dir()


# ── Flask ─────────────────────────────────────────────────────────────────────
try:
    from flask import Flask, Response, jsonify, render_template, request, session
except ImportError:
    print("Flask required: pip install flask")
    sys.exit(1)

# ── PIL ───────────────────────────────────────────────────────────────────────
try:
    from PIL import Image as PILImage
    PIL_OK = True
except ImportError:
    PIL_OK = False

# ── Scanner ───────────────────────────────────────────────────────────────────
sys.path.insert(0, str(Path(__file__).parent))
try:
    import document_scanner as ds
    SCANNER_OK = True
except ImportError as e:
    logger.warning("document_scanner not found: %s", e)
    SCANNER_OK = False

try:
    from file_scanner import FileScanner, store_smb_password, SMB_OK as _SMB_OK
    FILE_SCANNER_OK = True
except ImportError:
    FILE_SCANNER_OK = False
    _SMB_OK = False

# ── Connector ─────────────────────────────────────────────────────────────────
try:
    from m365_connector import M365Connector, M365Error, M365PermissionError, M365DeltaTokenExpired, MSAL_OK, REQUESTS_OK
    CONNECTOR_OK = True
except ImportError as e:
    logger.warning("m365_connector not found: %s", e)
    CONNECTOR_OK = False

try:
    from google_connector import GoogleConnector as _GoogleConnector, GOOGLE_AUTH_OK, load_saved_key as _load_google_key
    GOOGLE_CONNECTOR_OK = True
except ImportError:
    GOOGLE_CONNECTOR_OK = False
    GOOGLE_AUTH_OK = False
    def _load_google_key(): return None

try:
    from gdpr_db import get_db as _get_db, ScanDB as _ScanDB
    DB_OK = True
except ImportError:
    DB_OK = False
    def _get_db(*a, **kw): return None

# Single source of truth — read from VERSION file alongside this script
APP_VERSION = (Path(__file__).parent / "VERSION").read_text().strip()

# Compiled once — matches a bare UUID with no surrounding text
_GUID_RE = re.compile(
    r'^[0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{12}$', re.I
)
# Localised variants of the generic guest-account placeholder
_GENERIC_DISPLAY_NAMES = {
    "microsoft konto", "microsoft account", "microsoftkonto",
    "microsoft-konto", "compte microsoft", "cuenta de microsoft",
}

def _resolve_display_name(display_name: str,
                           email: str = "",
                           upn: str = "") -> str:
    """Return the best human-readable name for a Microsoft 365 user.

    Guest accounts (personal Microsoft accounts invited to the tenant) often
    have their displayName set to either:
      - A raw GUID  (e.g. "c710b7e1-4f9a-4066-a66f-f8b0b1b0ade3")
      - A generic localised placeholder  (e.g. "Microsoft Konto")
    In those cases we fall back to the email address or UPN, which is always
    human-readable and uniquely identifies the account.
    """
    dn = (display_name or "").strip()
    if not dn or _GUID_RE.match(dn) or dn.lower() in _GENERIC_DISPLAY_NAMES:
        return email or upn or dn
    return dn


# ── Sub-module imports — re-export everything for blueprint __getattr__ ────────
from app_config import (
    _load_src_toggles, _save_src_toggles,
    LANG, _load_lang, _load_lang_forced, _lang_override, _set_lang_override,
    _load_keywords, _check_special_category,
    _compiled_keywords, _keyword_data, _keyword_flat,
    _load_config, _save_config,
    _get_admin_pin_hash, _set_admin_pin, _verify_admin_pin, _admin_pin_is_set,
    _profiles_load, _profiles_write, _profiles_save_all, _profile_from_settings,
    _profile_get, _profile_save, _profile_delete, _profile_touch,
    _save_settings, _load_settings,
    _load_role_overrides, _save_role_overrides,
    _load_file_sources, _save_file_sources,
    _get_fernet, _encrypt_password, _decrypt_password,
    _load_smtp_config, _save_smtp_config,
    _SETTINGS_PATH, _SMTP_CONFIG_PATH, _ROLE_OVERRIDES_PATH,
    _FILE_SOURCES_PATH, _MACHINE_ID_PATH,
)
# _load_keywords already called by app_config at import time

from checkpoint import (
    _checkpoint_key, _save_checkpoint, _load_checkpoint, _clear_checkpoint,
    _load_delta_tokens, _save_delta_tokens,
    _CHECKPOINT_PATH, _DELTA_PATH,
)

from sse import broadcast, _sse_queues, _sse_buffer
import sse as _sse_mod  # for _current_scan_id access at call time

from cpr_detector import (
    _scan_bytes, _scan_bytes_timeout, _scan_text_direct, _html_esc, _get_pii_counts,
    _make_thumb, _placeholder_svg,
    _extract_exif, _detect_photo_faces,
    SUPPORTED_EXTS, PHOTO_EXTS,
    _EXIF_PII_TAGS,
)
# Inject runtime deps into cpr_detector
import cpr_detector as _cprd
_cprd.ds             = ds
_cprd.SCANNER_OK     = SCANNER_OK
_cprd.PILImage       = PILImage if PIL_OK else None
_cprd.PIL_OK         = PIL_OK
_cprd.LANG           = LANG
_cprd._check_special_category = _check_special_category

from scan_engine import run_scan, run_file_scan
# Inject runtime deps into scan_engine
import scan_engine as _se
_se.broadcast        = broadcast
_se._sse_buffer      = _sse_buffer
_se.LANG             = LANG
_se.SCANNER_OK       = SCANNER_OK
_se.PIL_OK           = PIL_OK
_se.FILE_SCANNER_OK  = FILE_SCANNER_OK
_se.CONNECTOR_OK     = CONNECTOR_OK
_se.DB_OK            = DB_OK
_se.PHOTO_EXTS       = PHOTO_EXTS
_se.SUPPORTED_EXTS   = SUPPORTED_EXTS
# cpr helpers
_se._scan_bytes              = _scan_bytes
_se._scan_bytes_timeout      = _scan_bytes_timeout
_se._detect_photo_faces      = _detect_photo_faces
_se._extract_exif            = _extract_exif
_se._make_thumb              = _make_thumb
_se._placeholder_svg         = _placeholder_svg
_se._check_special_category  = _check_special_category
_se._get_pii_counts          = _get_pii_counts
_se._html_esc                = _html_esc
# checkpoint
_se._load_checkpoint    = _load_checkpoint
_se._save_checkpoint    = _save_checkpoint
_se._clear_checkpoint   = _clear_checkpoint
_se._checkpoint_key     = _checkpoint_key
_se._load_delta_tokens  = _load_delta_tokens
_se._save_delta_tokens  = _save_delta_tokens

# ── App state ─────────────────────────────────────────────────────────────────
import os as _os
_BASE_DIR = _os.path.dirname(_os.path.abspath(__file__))
if getattr(sys, "frozen", False):  # PyInstaller bundle
    _BASE_DIR = sys._MEIPASS
app = Flask(__name__,
            template_folder=_os.path.join(_BASE_DIR, "templates"),
            static_folder=_os.path.join(_BASE_DIR, "static"))

# Session secret — derived from machine_id so it survives restarts without a separate file.
# machine_id is also the Fernet key (base64-encoded 32 bytes); we use its raw bytes as the secret.
try:
    from app_config import _MACHINE_ID_PATH as _mid_path  # type: ignore[attr-defined]
    import base64 as _b64
    _mid_bytes = _mid_path.read_bytes() if _mid_path.exists() else None
    app.secret_key = _b64.b64decode(_mid_bytes) if _mid_bytes else _os.urandom(32)
except Exception:
    app.secret_key = _os.urandom(32)

_connector:  "M365Connector | None" = None
# _scan_lock and _scan_abort live in routes/state.py
from routes.state import _scan_lock, _scan_abort
# _sse_queues, _sse_buffer, _current_scan_id live in sse.py
flagged_items: list = []
scan_meta:    dict = {}

# ── Checkpoint (incremental / resumable scans) ────────────────────────────────
# ── HTML ──────────────────────────────────────────────────────────────────────

# ── Shared state (imported by route blueprints) ───────────────────────────────
from routes import state as _state
# Wire the mutable globals to the state module so blueprints share the same objects
# These assignments run once at startup; blueprints use state.X to read/write them.
_state.LANG              = LANG
_state.connector         = _connector
_state.flagged_items     = flagged_items
_state.scan_meta         = scan_meta
_state.compiled_keywords = _compiled_keywords
_state.keyword_data      = _keyword_data
_state.keyword_flat      = _keyword_flat

# ── Auto-restore Google Workspace connector from saved key ────────────────────
if GOOGLE_CONNECTOR_OK:
    try:
        _gkey = _load_google_key()
        if _gkey:
            from routes.google_auth import _load_google_config as _lgcfg
            _gcfg = _lgcfg()
            _state.google_connector = _GoogleConnector(_gkey, admin_email=_gcfg.get("admin_email", ""))
    except Exception as _ge:
        logger.warning("[google] Could not restore connector: %s", _ge)

# Helper so scan engine can update state.flagged_items in-place rather than rebind
def _sync_state():
    """Called after scan updates flagged_items/scan_meta to sync the state module."""
    _state.flagged_items[:] = flagged_items
    _state.scan_meta.clear()
    _state.scan_meta.update(scan_meta)

# ── HTML template ────────────────────────────────────────────────────────────
# Served from templates/index.html via Flask render_template().
# JavaScript served from static/app.js via Flask static file handling.


# ── Auth state ─────────────────────────────────────────────────────────────────
# ── Routes ────────────────────────────────────────────────────────────────────

@app.route("/")
def index():
    return render_template("index.html", app_version=APP_VERSION,
                            lang_json=json.dumps(LANG, ensure_ascii=False),
                            viewer_mode=False)


@app.route("/view")
def viewer():
    from app_config import validate_viewer_token, get_viewer_pin_hash
    token = request.args.get("token", "").strip()
    if token:
        if validate_viewer_token(token) is None:
            return render_template("viewer_denied.html"), 403
        # Bind a session so the viewer doesn't need the token on every navigation
        session["viewer_ok"] = True
        return render_template("index.html", app_version=APP_VERSION,
                                lang_json=json.dumps(LANG, ensure_ascii=False),
                                viewer_mode=True)
    if session.get("viewer_ok"):
        return render_template("index.html", app_version=APP_VERSION,
                                lang_json=json.dumps(LANG, ensure_ascii=False),
                                viewer_mode=True)
    # No token, no session — show PIN form if a PIN is configured, else deny
    pin_hash = get_viewer_pin_hash()
    if pin_hash:
        return render_template("viewer_pin.html")
    return render_template("viewer_denied.html"), 403


def _build_excel_bytes() -> tuple[bytes, str]:
    """Build the M365 scan Excel workbook and return (bytes, filename).
    Raises on error. Used by export_excel() and send_report()."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    HEADER_BG  = "1F3864"
    HEADER_FG  = "FFFFFF"
    ALT_BG     = "EEF2FF"
    SOURCE_MAP = {
        "email":      ("📧 Outlook",     "D6E4F7"),
        "onedrive":   ("💾 OneDrive",   "D6F7E4"),
        "sharepoint": ("🌐 SharePoint", "FFF0D6"),
        "teams":      ("💬 Teams",      "F7D6F0"),
        "local":      ("📁 Local",      "E6F7E6"),
        "smb":        ("🌐 Network",    "E0F0FA"),
    }
    COLS = [
        ("Name / Subject",    45),
        ("CPR Hits",           9),
        ("Face count",         9),
        ("GPS",                6),
        ("Special category",  22),
        ("EXIF author",       18),
        ("Folder",            30),
        ("Account",           24),
        ("Role",              10),
        ("Disposition",       18),
        ("Date Modified",     14),
        ("Size (KB)",         10),
        ("URL",               50),
    ]

    thin   = Side(style="thin", color="CCCCCC")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    def _fill(hex_col):
        return PatternFill("solid", fgColor=hex_col)

    def _write_sheet(ws, rows, tab_color):
        ws.sheet_properties.tabColor = tab_color
        for col_idx, (col_name, col_w) in enumerate(COLS, 1):
            cell = ws.cell(row=1, column=col_idx, value=col_name)
            cell.font      = Font(name="Arial", bold=True, color=HEADER_FG, size=10)
            cell.fill      = _fill(HEADER_BG)
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            cell.border    = border
            ws.column_dimensions[get_column_letter(col_idx)].width = col_w
        ws.row_dimensions[1].height = 20
        ws.freeze_panes = "A2"

        for r_idx, item in enumerate(rows, 2):
            row_fill = _fill(ALT_BG if r_idx % 2 == 0 else "FFFFFF")
            _disp = ""
            if DB_OK:
                try:
                    _d = _get_db().get_disposition(item.get("id", ""))
                    _disp = (_d.get("status", "") if _d else "")
                except Exception:
                    pass
            _sc = item.get("special_category", [])
            _sc_str = ", ".join(
                s for s in (_sc if isinstance(_sc, list) else [str(_sc or "")])
                if s not in ("gps_location", "exif_pii")
            )
            _exif   = item.get("exif") or {}
            _gps    = _exif.get("gps")
            _author = _exif.get("author") or ""
            values = [
                item.get("name", ""),
                item.get("cpr_count", 0),
                item.get("face_count", 0),
                "✔" if _gps else "",
                _sc_str,
                _author,
                item.get("folder", ""),
                item.get("account_name", "") or item.get("source", ""),
                item.get("user_role", ""),
                _disp,
                item.get("modified", ""),
                item.get("size_kb", ""),
                item.get("url", ""),
            ]
            for col_idx, val in enumerate(values, 1):
                is_url = col_idx == 13 and val
                cell = ws.cell(row=r_idx, column=col_idx, value=val)
                cell.font      = Font(name="Arial", size=10,
                                     color="1155CC" if is_url else "000000",
                                     underline="single" if is_url else None)
                cell.fill      = row_fill
                cell.alignment = Alignment(vertical="center", wrap_text=(col_idx == 1))
                cell.border    = border
            ws.row_dimensions[r_idx].height = 16

        if rows:
            tr = len(rows) + 2
            ws.cell(row=tr, column=1, value="Total").font = Font(name="Arial", bold=True, size=10)
            ws.cell(row=tr, column=2, value=f"=SUM(B2:B{tr-1})").font = Font(name="Arial", bold=True, size=10)
            for col_idx in range(1, len(COLS) + 1):
                ws.cell(row=tr, column=col_idx).fill   = _fill("D0D8F0")
                ws.cell(row=tr, column=col_idx).border = border

        ws.auto_filter.ref = f"A1:{get_column_letter(len(COLS))}1"

    wb     = Workbook()
    ws_sum = wb.active
    ws_sum.title = "Summary"
    ws_sum.sheet_properties.tabColor = "1F3864"
    ws_sum["A1"] = "GDPRScanner — Export"
    ws_sum["A1"].font = Font(name="Arial", bold=True, size=14, color=HEADER_FG)
    ws_sum["A1"].fill = _fill(HEADER_BG)
    ws_sum.merge_cells("A1:D1")
    ws_sum["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws_sum.row_dimensions[1].height = 28

    import datetime as _dt
    ws_sum["A2"] = "Generated:"
    ws_sum["B2"] = _dt.datetime.now().strftime("%Y-%m-%d %H:%M")
    ws_sum["A3"] = "Total flagged items:"
    ws_sum["B3"] = len(flagged_items)
    gps_count = sum(1 for i in flagged_items if (i.get("exif") or {}).get("gps"))
    if gps_count:
        ws_sum["A4"] = "Items with GPS data:"
        ws_sum["B4"] = gps_count
    for cell in (ws_sum["A2"], ws_sum["A3"], ws_sum["A4"]):
        cell.font = Font(name="Arial", bold=True, size=10)
    for cell in (ws_sum["B2"], ws_sum["B3"], ws_sum["B4"]):
        cell.font = Font(name="Arial", size=10)
    ws_sum.column_dimensions["A"].width = 22
    ws_sum.column_dimensions["B"].width = 20

    for ci, h in enumerate(["Source", "Items", "Total CPR Hits"], 1):
        cell = ws_sum.cell(row=6, column=ci, value=h)
        cell.font      = Font(name="Arial", bold=True, color=HEADER_FG, size=10)
        cell.fill      = _fill(HEADER_BG)
        cell.border    = border
        cell.alignment = Alignment(horizontal="center", vertical="center")
    ws_sum.row_dimensions[6].height = 18
    ws_sum.column_dimensions["C"].width = 16

    by_source: dict = {}
    for item in flagged_items:
        by_source.setdefault(item.get("source_type", "other"), []).append(item)

    sum_row = 7
    for src_key, (label, tab_bg) in SOURCE_MAP.items():
        items = by_source.get(src_key, [])
        if not items:
            continue
        ws_sum.cell(row=sum_row, column=1, value=label).font = Font(name="Arial", size=10)
        ws_sum.cell(row=sum_row, column=2, value=len(items)).font = Font(name="Arial", size=10)
        ws_sum.cell(row=sum_row, column=3, value=sum(i.get("cpr_count", 0) for i in items)).font = Font(name="Arial", size=10)
        for ci in range(1, 4):
            ws_sum.cell(row=sum_row, column=ci).border = border
            ws_sum.cell(row=sum_row, column=ci).fill = _fill("EEF2FF" if sum_row % 2 == 0 else "FFFFFF")
        sum_row += 1

    for src_key, (label, tab_bg) in SOURCE_MAP.items():
        items = by_source.get(src_key, [])
        if not items:
            continue
        clean_label = label.split(" ", 1)[1]
        _write_sheet(wb.create_sheet(title=clean_label), items, tab_bg)

    # GPS items sheet
    gps_items = [i for i in flagged_items if (i.get("exif") or {}).get("gps")]
    if gps_items:
        ws_gps = wb.create_sheet(title="GPS locations")
        ws_gps.sheet_properties.tabColor = "1A7A6E"
        GPS_COLS = [
            ("Name", 40), ("Latitude", 14), ("Longitude", 14),
            ("Maps link", 50), ("Account", 24), ("Date Modified", 14),
        ]
        for col_idx, (col_name, col_w) in enumerate(GPS_COLS, 1):
            cell = ws_gps.cell(row=1, column=col_idx, value=col_name)
            cell.font      = Font(name="Arial", bold=True, color=HEADER_FG, size=10)
            cell.fill      = _fill("1A7A6E")
            cell.alignment = Alignment(horizontal="center", vertical="center")
            cell.border    = border
            ws_gps.column_dimensions[get_column_letter(col_idx)].width = col_w
        ws_gps.freeze_panes = "A2"
        for r_idx, item in enumerate(gps_items, 2):
            _exif = item.get("exif") or {}
            _gps  = _exif.get("gps") or {}
            row_fill = _fill("E0F7F4" if r_idx % 2 == 0 else "FFFFFF")
            for col_idx, val in enumerate([
                item.get("name", ""),
                _gps.get("lat", ""),
                _gps.get("lon", ""),
                _gps.get("maps_url", ""),
                item.get("account_name", "") or item.get("source", ""),
                item.get("modified", ""),
            ], 1):
                is_link = col_idx == 4 and val
                cell = ws_gps.cell(row=r_idx, column=col_idx, value=val)
                cell.font   = Font(name="Arial", size=10,
                                   color="1155CC" if is_link else "000000",
                                   underline="single" if is_link else None)
                cell.fill   = row_fill
                cell.border = border
        ws_gps.auto_filter.ref = f"A1:{get_column_letter(len(GPS_COLS))}1"

    # External transfers sheet
    ext_items = [i for i in flagged_items
                 if i.get("transfer_risk") in ("external-recipient", "external-share", "shared")]
    if ext_items:
        ws_ext = wb.create_sheet(title="External transfers")
        _write_sheet(ws_ext, ext_items, "E74C3C")
        ws_ext.sheet_properties.tabColor = "E74C3C"
        ws_sum.cell(row=sum_row, column=1, value="⚠ External transfers").font = Font(name="Arial", size=10, bold=True, color="E74C3C")
        ws_sum.cell(row=sum_row, column=2, value=len(ext_items)).font = Font(name="Arial", size=10, bold=True, color="E74C3C")
        ws_sum.cell(row=sum_row, column=3, value=sum(i.get("cpr_count", 0) for i in ext_items)).font = Font(name="Arial", size=10, bold=True, color="E74C3C")
        for ci in range(1, 4):
            ws_sum.cell(row=sum_row, column=ci).border = border
            ws_sum.cell(row=sum_row, column=ci).fill = _fill("FDE8E8")

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    fname = f"gdpr_scan_{_dt.datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    return buf.read(), fname


# ── Article 30 report ─────────────────────────────────────────────────────────

def _build_article30_docx() -> tuple[bytes, str]:
    """Generate a GDPR Article 30 Register of Processing Activities as .docx.
    Returns (bytes, filename). Strings are translated using the active LANG dict."""
    try:
        from docx import Document as _Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise ImportError("python-docx not installed — run: pip install python-docx")

    import datetime as _dt

    # Translate helper — falls back to English default if key missing
    def L(key: str, default: str = "") -> str:
        return LANG.get(key, default)

    # ── Data ─────────────────────────────────────────────────────────────────
    db    = _get_db() if DB_OK else None
    stats   = db.get_stats() if db else {}
    items   = db.get_flagged_items() if db else list(flagged_items)
    trend   = db.get_trend(10) if db else []
    overdue = db.get_overdue_items(5) if db else []

    # Build account_id → display_name map from the scan's stored user_ids
    # This lets us resolve GUIDs and "Microsoft Konto" placeholders that
    # were stored in account_name before _resolve_display_name was applied.
    _acct_map: dict[str, str] = {}
    if db:
        try:
            scan_id = stats.get("scan_id") or db.latest_scan_id()
            if scan_id:
                row = db._connect().execute(
                    "SELECT user_count, options FROM scans WHERE id=?", (scan_id,)
                ).fetchone()
                # user_ids are stored in the options JSON column
                opts_json = json.loads(row["options"] or "{}") if row else {}
                for u in opts_json.get("user_ids", []):
                    uid  = u.get("id", "")
                    name = u.get("displayName", "")
                    if uid and name:
                        _acct_map[uid] = name
        except Exception:
            pass
    # Also seed from in-memory flagged_items (catches current scan not yet in DB)
    for item in flagged_items:
        aid  = item.get("account_id", "")
        name = item.get("account_name", "")
        if aid and name and not _GUID_RE.match(name.strip()):
            _acct_map.setdefault(aid, name)

    def _acct_label(item: dict) -> str:
        """Return the best human-readable account label for an item."""
        aid  = item.get("account_id", "")
        name = item.get("account_name", "")
        # Try the lookup map first (most reliable — built from scan user_ids)
        if aid and aid in _acct_map:
            return _acct_map[aid]
        # Fall back to stored name, resolving GUIDs/placeholders against account_id
        return _resolve_display_name(name, aid)
    overdue_ids = {o["id"] for o in overdue}

    now_str   = _dt.datetime.now().strftime("%Y-%m-%d %H:%M")
    date_str  = _dt.datetime.now().strftime("%Y-%m-%d")
    fname     = f"article30_{date_str}.docx"

    # Aggregate by source
    by_source: dict = {}
    for item in items:
        st = item.get("source_type", "other")
        by_source.setdefault(st, []).append(item)

    SOURCE_LABELS = {
        "email":      "Exchange (Outlook)",
        "onedrive":   "OneDrive",
        "sharepoint": "SharePoint",
        "teams":      "Teams",
        "gmail":      "Gmail",
        "gdrive":     "Google Drive",
        "local":      "Lokal",
        "smb":        "Netværk (SMB)",
    }

    # ── Colour palette ────────────────────────────────────────────────────────
    DARK_BLUE  = RGBColor(0x1F, 0x38, 0x64)
    MID_BLUE   = RGBColor(0x00, 0x78, 0xD4)
    LIGHT_GREY = RGBColor(0xF2, 0xF2, 0xF2)
    RED        = RGBColor(0xC0, 0x39, 0x2B)
    ORANGE     = RGBColor(0xC5, 0x5A, 0x00)
    WHITE      = RGBColor(0xFF, 0xFF, 0xFF)

    def _hex(c: RGBColor) -> str:
        return f"{c[0]:02X}{c[1]:02X}{c[2]:02X}"

    # ── Document setup ────────────────────────────────────────────────────────
    doc = _Document()
    doc.core_properties.title   = "GDPR Article 30 — Register of Processing Activities"
    doc.core_properties.author  = "GDPRScanner"
    doc.core_properties.subject = "GDPR Compliance"

    # Page margins — A4 with 2.5 cm margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Helper: set cell background ──────────────────────────────────────────
    def _cell_bg(cell, hex_color: str):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        tcPr.append(shd)

    def _set_cell_border(cell, **kwargs):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for edge in ("top", "left", "bottom", "right"):
            cfg = kwargs.get(edge, {})
            el  = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"),   cfg.get("val",   "single"))
            el.set(qn("w:sz"),    cfg.get("sz",    "4"))
            el.set(qn("w:space"), cfg.get("space", "0"))
            el.set(qn("w:color"), cfg.get("color", "CCCCCC"))
            tcBorders.append(el)
        tcPr.append(tcBorders)

    def _para(text: str = "", bold=False, size=11, color=None,
              align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6) -> object:
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        if text:
            run = p.add_run(text)
            run.bold      = bold
            run.font.size = Pt(size)
            if color:
                run.font.color.rgb = color
        return p

    def _heading(text: str, level: int = 1):
        p  = doc.add_heading(text, level=level)
        r  = p.runs[0] if p.runs else p.add_run(text)
        r.font.color.rgb = DARK_BLUE
        r.font.size      = Pt(16 if level == 1 else 13)
        r.bold           = True
        p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
        p.paragraph_format.space_after  = Pt(4)
        return p

    def _kv(label: str, value: str, label_width=2.5, bold=False, highlight=False):
        """Two-column key-value paragraph using a 2-cell table row."""
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        w_label = int(label_width * 1440)
        w_value = int((16.0 - label_width) * 1440 * 0.6)  # approx content width
        c1, c2 = tbl.rows[0].cells
        _cell_bg(c1, "FFF3E0" if highlight else "F2F2F2")
        _cell_bg(c2, "FFF3E0" if highlight else "FFFFFF")
        c1.width = Inches(label_width)
        c2.width = Inches(16.0 - label_width)
        p1 = c1.paragraphs[0]; p1.clear()
        r1 = p1.add_run(label); r1.bold = True; r1.font.size = Pt(10)
        p2 = c2.paragraphs[0]; p2.clear()
        r2 = p2.add_run(value); r2.font.size = Pt(10); r2.bold = bold
        if highlight:
            r1.font.color.rgb = RGBColor(0x6B, 0x00, 0x6B)
            r2.font.color.rgb = RGBColor(0x6B, 0x00, 0x6B)
        for cell in (c1, c2):
            _set_cell_border(cell, top={"color": "E0E0E0"}, bottom={"color": "E0E0E0"},
                             left={"color": "E0E0E0"}, right={"color": "E0E0E0"})
        return tbl

    # ── Cover page ────────────────────────────────────────────────────────────
    _para()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(40)
    r = title_p.add_run(L("a30_title", "GDPR Article 30"))
    r.bold = True; r.font.size = Pt(28); r.font.color.rgb = DARK_BLUE

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub_p.add_run(L("a30_subtitle", "Register of Processing Activities"))
    r2.font.size = Pt(16); r2.font.color.rgb = MID_BLUE

    _para()
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = meta_p.add_run(f"{L('a30_generated','Generated')}: {now_str}  ·  GDPRScanner")
    r3.font.size = Pt(10); r3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Divider line
    _para()
    div = doc.add_paragraph()
    div_fmt = div.paragraph_format
    div_fmt.space_after = Pt(20)
    pPr = div._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6")
    bot.set(qn("w:color"), _hex(MID_BLUE))
    pBdr.append(bot); pPr.append(pBdr)

    doc.add_page_break()

    # ── Section 1: Summary ────────────────────────────────────────────────────
    _heading(L("a30_s1", "1. Summary"))

    total_items    = len(items)
    total_cpr      = sum(i.get("cpr_count", 0) for i in items)
    special_items  = [i for i in items if i.get("special_category") and
                      i["special_category"] not in ("[]", "", None, [])]
    photo_items    = [i for i in items if i.get("face_count", 0) > 0]
    gps_items      = [i for i in items if "gps_location" in (i.get("special_category") or [])]
    exif_pii_items = [i for i in items if "exif_pii" in (i.get("special_category") or [])]
    unique_subj    = stats.get("unique_subjects", 0)
    total_scanned  = stats.get("total_scanned", 0)
    scan_date      = _dt.datetime.fromtimestamp(
        stats.get("started_at", 0)).strftime("%Y-%m-%d %H:%M") if stats.get("started_at") else "—"
    special_items  = [i for i in items if i.get("special_category") and
                      i["special_category"] not in ("[]", "", None, [])]

    _kv(L("a30_scan_date",       "Scan date"),                scan_date)
    _kv(L("a30_items_scanned",   "Items scanned"),            str(total_scanned))
    _kv(L("a30_flagged",         "Flagged items"),            str(total_items))
    _kv(L("a30_cpr_hits",        "Total CPR hits"),           str(total_cpr))
    _kv(L("a30_data_subjects",   "Estimated data subjects"),  str(unique_subj))
    _kv(L("a30_overdue",         "Overdue items (>5 yrs)"),   str(len(overdue_ids)))
    if gps_items:
        _kv(L("a30_gps_items", "Items with GPS location data (Art. 4 — location = personal data)"),
            str(len(gps_items)))
    if exif_pii_items:
        _kv(L("a30_exif_pii_items", "Items with EXIF PII (author, description, keywords)"),
            str(len(exif_pii_items)))
    if photo_items:
        total_faces = sum(i.get("face_count", 0) for i in photo_items)
        _kv(L("a30_photo_items", "Photos with detected faces (Art. 9 biometric)"),
            f"{len(photo_items)} items / {total_faces} faces")
        _para(L("a30_photo_note",
                "Photographs of identifiable persons are biometric data under Art. 9 GDPR. "
                "Retention requires a documented legal basis under Art. 9(2). "
                "For school photographs of pupils under 15, parental consent is required "
                "(Databeskyttelsesloven §6). See Datatilsynet guidance on school photography."),
              size=9, space_after=4)
    if special_items:
        _kv(L("a30_special_cat", "Art. 9 special category items"),
            str(len(special_items)))
        _para(L("a30_special_cat_note",
                "These items contain health, criminal, biometric, religious, ethnic, "
                "trade union, political, or sexual orientation data. "
                "An explicit legal basis (Art. 9(2)) and possibly a DPIA (Art. 35) is required."),
              size=9, space_after=4)

    _para()

    # Per-source breakdown table
    _para(L("a30_by_source", "Breakdown by source"), bold=True, size=11, space_before=10)

    src_tbl = doc.add_table(rows=1, cols=5)
    src_tbl.style = "Table Grid"
    hdr_cells = src_tbl.rows[0].cells
    for cell, txt in zip(hdr_cells, [L("a30_col_source","Source"), L("a30_col_items","Items"),
                                     L("a30_col_cpr","CPR hits"), L("a30_col_overdue","Overdue"),
                                     L("a30_col_special","Art. 9")]):
        _cell_bg(cell, _hex(DARK_BLUE))
        p = cell.paragraphs[0]; p.clear()
        r = p.add_run(txt); r.bold = True
        r.font.size = Pt(10); r.font.color.rgb = WHITE

    for src_key in ("email", "onedrive", "sharepoint", "teams"):
        src_items = by_source.get(src_key, [])
        if not src_items:
            continue
        row   = src_tbl.add_row().cells
        n_ov   = sum(1 for i in src_items if i.get("id") in overdue_ids)
        n_cpr  = sum(i.get("cpr_count", 0) for i in src_items)
        n_spec = sum(1 for i in src_items if i.get("special_category") and
                     i["special_category"] not in ("[]", "", None, []))
        for cell, val in zip(row, [
            SOURCE_LABELS.get(src_key, src_key),
            str(len(src_items)), str(n_cpr), str(n_ov),
            str(n_spec) if n_spec else "—"
        ]):
            p = cell.paragraphs[0]; p.clear()
            r = p.add_run(val); r.font.size = Pt(10)
            if val != "0" and cell == row[3]:
                r.font.color.rgb = ORANGE
            if n_spec and cell == row[4]:
                r.font.color.rgb = RGBColor(0x7B, 0x00, 0x82)
                r.bold = True

    # ── Section 2: Data categories ────────────────────────────────────────────
    doc.add_page_break()
    _heading(L("a30_s2", "2. Personal Data Categories Identified"))

    _para(L("a30_s2_intro", "The following categories of personal data were detected during scanning."),
          size=10, space_after=8)

    # Aggregate PII from DB or from items
    pii_totals: dict = {}
    if db:
        rows = db._connect().execute(
            """SELECT pii_type, SUM(hit_count) FROM pii_hits
               WHERE scan_id=? GROUP BY pii_type""",
            (stats.get("scan_id") or db.latest_scan_id() or 0,)
        ).fetchall()
        for pii_type, count in rows:
            pii_totals[pii_type] = count

    PII_LABELS = {
        "PHONE":        L("a30_pii_phone",        "Phone numbers"),
        "EMAIL":        L("a30_pii_email",        "Email addresses"),
        "IBAN":         L("a30_pii_iban",         "IBAN bank numbers"),
        "BANK_ACCOUNT": L("a30_pii_bank",         "Bank account numbers"),
        "NAME":         L("a30_pii_name",         "Personal names (NER)"),
        "ADDRESS":      L("a30_pii_address",      "Addresses (NER)"),
        "ORG":          L("a30_pii_org",          "Organisations (NER)"),
    }

    pii_tbl = doc.add_table(rows=1, cols=3)
    pii_tbl.style = "Table Grid"
    for cell, txt in zip(pii_tbl.rows[0].cells,
                          [L("a30_col_category","Data category"), L("a30_col_count","Count"), L("a30_col_gdpr_class","GDPR classification")]):
        _cell_bg(cell, _hex(DARK_BLUE))
        p = cell.paragraphs[0]; p.clear()
        r = p.add_run(txt); r.bold = True
        r.font.size = Pt(10); r.font.color.rgb = WHITE

    # CPR row first — always
    cpr_row = pii_tbl.add_row().cells
    for cell, val in zip(cpr_row, [L("a30_cpr_label", "CPR numbers (Danish personal ID)"), str(total_cpr),
                                    L("a30_cpr_class", "Art. 9 — national identifier")]):
        p = cell.paragraphs[0]; p.clear()
        r = p.add_run(val); r.font.size = Pt(10)
        _cpr_class = L("a30_cpr_class", "Art. 9 — national identifier")
        if val == _cpr_class:
            r.font.color.rgb = RED; r.bold = True

    for pii_type, label in PII_LABELS.items():
        count = pii_totals.get(pii_type, 0)
        if not count:
            continue
        cls = L("a30_pii_class_9", "Art. 9 — health/sensitive") if pii_type in ("NAME", "ADDRESS") else L("a30_pii_class_4", "Art. 4 — personal data")
        row = pii_tbl.add_row().cells
        for cell, val in zip(row, [label, str(count), cls]):
            p = cell.paragraphs[0]; p.clear()
            r = p.add_run(val); r.font.size = Pt(10)

    # ── Section 3: Data inventory ─────────────────────────────────────────────
    doc.add_page_break()
    _heading(L("a30_s3", "3. Data Inventory"))

    _para(L("a30_s3_intro", "All flagged items are listed below with location, retention status, and compliance disposition."),
          size=10, space_after=8)

    # Split by user role for separate presentation
    student_items = [i for i in items if i.get("user_role") == "student"]
    staff_items   = [i for i in items if i.get("user_role") != "student"]

    _disp_map = {
        "unreviewed":       L("a30_disp_unreviewed",      "Unreviewed"),
        "retain-legal":     L("a30_disp_retain_legal",    "Retain — Legal obligation"),
        "retain-legitimate": L("a30_disp_retain_legit",   "Retain — Legitimate interest"),
        "retain-contract":  L("a30_disp_retain_contract", "Retain — Contract"),
        "delete-scheduled": L("a30_disp_delete_sched",    "Delete — Scheduled"),
        "deleted":          L("a30_disp_deleted",         "Deleted"),
    }

    def _inv_table(tbl_items: list):
        tbl = doc.add_table(rows=1, cols=6)
        tbl.style = "Table Grid"
        col_hdrs = [L("a30_col_name","Name / Subject"), L("a30_col_source","Source"),
                    L("a30_col_account","Account"), L("a30_col_modified","Modified"),
                    L("a30_col_cpr_short","CPR"), L("a30_col_disp","Disposition")]
        for cell, txt in zip(tbl.rows[0].cells, col_hdrs):
            _cell_bg(cell, _hex(DARK_BLUE))
            p = cell.paragraphs[0]; p.clear()
            r = p.add_run(txt); r.bold = True
            r.font.size = Pt(9); r.font.color.rgb = WHITE
        sorted_tbl = sorted(tbl_items,
            key=lambda x: (0 if x.get("id") in overdue_ids else 1, -x.get("cpr_count", 0)))
        for idx, item in enumerate(sorted_tbl[:500]):
            disp_rec = db.get_disposition(item["id"]) if db else None
            raw_disp = disp_rec.get("status", "unreviewed") if disp_rec else "unreviewed"
            disp_str = _disp_map.get(raw_disp, raw_disp.replace("-", " ").title())
            is_ov    = item.get("id") in overdue_ids
            row = tbl.add_row().cells
            vals = [
                (item.get("name", "")[:60] + ("…" if len(item.get("name", "")) > 60 else "")),
                SOURCE_LABELS.get(item.get("source_type", ""), item.get("source_type", "")),
                _acct_label(item),
                item.get("modified", ""),
                str(item.get("cpr_count", 0)),
                disp_str,
            ]
            bg = "FFF8F0" if is_ov else ("FFFFFF" if idx % 2 == 0 else "F8F8F8")
            for cell, val in zip(row, vals):
                _cell_bg(cell, bg)
                p = cell.paragraphs[0]; p.clear()
                r = p.add_run(val); r.font.size = Pt(8)
                if is_ov and cell == row[3]:
                    r.font.color.rgb = ORANGE
        if len(tbl_items) > 500:
            _para(f"… {len(tbl_items) - 500} {L('a30_more_items', 'additional items not shown.')}",
                  size=9, color=RGBColor(0x88, 0x88, 0x88), space_before=4)

    if staff_items:
        if student_items:
            _para(L("a30_inv_staff", "👔 Staff / Faculty"), bold=True, size=11, space_before=6, space_after=4)
        _inv_table(staff_items)

    if student_items:
        _para(L("a30_inv_students", "🎓 Students"), bold=True, size=11, space_before=14, space_after=2)
        _para(L("a30_student_consent_note",
                "Note: Student accounts in Danish folkeskole (pupils under age 15) require parental "
                "consent for processing of personal data under Databeskyttelsesloven §6. "
                "Items in student accounts must not be auto-deleted — any action requires "
                "review by school administration and, for pupils under 15, notification of parents "
                "or guardians as rights holders under GDPR Article 8."),
              size=9, color=RGBColor(0x88, 0x44, 0x00), space_after=6)
        _inv_table(student_items)

    # ── Section 4: Retention analysis ────────────────────────────────────────
    if overdue:
        doc.add_page_break()
        _heading(L("a30_s4", "4. Retention Analysis"))

        _para(L("a30_s4_intro", "The following items exceed the 5-year retention threshold and should be reviewed for deletion under GDPR Article 5(1)(e) — storage limitation."),
              size=10, space_after=8)

        ret_tbl = doc.add_table(rows=1, cols=5)
        ret_tbl.style = "Table Grid"
        for cell, txt in zip(ret_tbl.rows[0].cells,
                              [L("a30_col_name","Name"), L("a30_col_source","Source"), L("a30_col_account","Account"), L("a30_col_modified","Modified"), L("a30_col_cpr","CPR hits")]):
            _cell_bg(cell, _hex(ORANGE))
            p = cell.paragraphs[0]; p.clear()
            r = p.add_run(txt); r.bold = True
            r.font.size = Pt(9); r.font.color.rgb = WHITE

        for item in overdue[:200]:
            row = ret_tbl.add_row().cells
            for cell, val in zip(row, [
                item.get("name", "")[:55],
                SOURCE_LABELS.get(item.get("source_type", ""), ""),
                _acct_label(item),
                item.get("modified", ""),
                str(item.get("cpr_count", 0)),
            ]):
                p = cell.paragraphs[0]; p.clear()
                r = p.add_run(val); r.font.size = Pt(8)

    # ── Section 5: Scan history ───────────────────────────────────────────────
    if trend:
        sec_num = "5" if overdue else "4"
        doc.add_page_break()
        _heading(f"{sec_num}. {L('a30_s5','Compliance Trend').split('. ',1)[-1]}")

        _para(L("a30_s5_intro", "Flagged item counts over the last scans (most recent first)."),
              size=10, space_after=8)

        trend_tbl = doc.add_table(rows=1, cols=4)
        trend_tbl.style = "Table Grid"
        for cell, txt in zip(trend_tbl.rows[0].cells,
                              [L("a30_col_scan_date","Scan date"), L("a30_col_flagged","Flagged"), L("a30_col_overdue","Overdue"), L("a30_col_scan_type","Scan type")]):
            _cell_bg(cell, _hex(DARK_BLUE))
            p = cell.paragraphs[0]; p.clear()
            r = p.add_run(txt); r.bold = True
            r.font.size = Pt(9); r.font.color.rgb = WHITE

        for t in reversed(trend):
            row = trend_tbl.add_row().cells
            for cell, val in zip(row, [
                t.get("scan_date", ""),
                str(t.get("flagged_count", 0)),
                str(t.get("overdue_count", 0)),
                L("a30_scan_delta", "Delta") if t.get("delta") else L("a30_scan_full", "Full"),
            ]):
                p = cell.paragraphs[0]; p.clear()
                r = p.add_run(val); r.font.size = Pt(9)

    # ── Section: Deletion audit log ───────────────────────────────────────────
    del_log   = db.get_deletion_log(limit=500) if db else []
    del_stats = db.deletion_log_stats() if db else {}

    # Running section counter — starts at 3 (summary, categories, inventory always present)
    last_sec  = 3
    last_sec += 1 if overdue  else 0   # retention analysis
    last_sec += 1 if trend    else 0   # compliance trend

    if del_log:
        del_sec   = last_sec
        last_sec += 1
        doc.add_page_break()
        _heading(f"{del_sec}. {L('a30_s_dellog', 'Deletion Audit Log')}")

        _para(L("a30_dellog_intro",
                f"A total of {del_stats.get('total', len(del_log))} item(s) containing personal data "
                f"have been deleted via GDPRScanner. "
                f"CPR hits removed: {del_stats.get('cpr_hits_deleted', 0)}. "
                f"This log satisfies the accountability obligation under GDPR Article 5(2)."),
              size=10, space_after=8)

        # Summary by reason
        by_reason = del_stats.get("by_reason", {})
        if by_reason:
            _para(L("a30_dellog_by_reason", "Deletions by reason"), bold=True, size=10, space_before=4, space_after=4)
            reason_tbl = doc.add_table(rows=1, cols=2)
            reason_tbl.style = "Table Grid"
            for cell, txt in zip(reason_tbl.rows[0].cells,
                                  [L("a30_col_reason", "Reason"), L("a30_col_count", "Count")]):
                _cell_bg(cell, _hex(DARK_BLUE))
                p = cell.paragraphs[0]; p.clear()
                r = p.add_run(txt); r.bold = True
                r.font.size = Pt(9); r.font.color.rgb = WHITE
            REASON_LABELS = {
                "manual":               L("a30_reason_manual",    "Manual (individual card delete)"),
                "bulk":                 L("a30_reason_bulk",       "Bulk delete"),
                "retention":            L("a30_reason_retention",  "Retention policy enforcement"),
                "data-subject-request": L("a30_reason_dsr",        "Data subject erasure request (Art. 17)"),
            }
            for reason, count in sorted(by_reason.items()):
                row = reason_tbl.add_row().cells
                for cell, val in zip(row, [REASON_LABELS.get(reason, reason), str(count)]):
                    p = cell.paragraphs[0]; p.clear()
                    r = p.add_run(val); r.font.size = Pt(9)

        # Full log table
        _para(L("a30_dellog_records", "Deletion records"), bold=True, size=10, space_before=10, space_after=4)
        log_tbl = doc.add_table(rows=1, cols=7)
        log_tbl.style = "Table Grid"
        for cell, txt in zip(log_tbl.rows[0].cells, [
            L("a30_col_deleted_at",  "Deleted at"),
            L("a30_col_name",        "Name"),
            L("a30_col_source",      "Source"),
            L("a30_col_account",     "Account"),
            L("a30_col_cpr",         "CPR hits"),
            L("a30_col_reason",      "Reason"),
            L("a30_col_deleted_by",  "Deleted by"),
        ]):
            _cell_bg(cell, _hex(DARK_BLUE))
            p = cell.paragraphs[0]; p.clear()
            r = p.add_run(txt); r.bold = True
            r.font.size = Pt(8); r.font.color.rgb = WHITE

        for idx, entry in enumerate(del_log):
            ts  = _dt.datetime.fromtimestamp(entry.get("deleted_at", 0)).strftime("%Y-%m-%d %H:%M")
            bg  = "FFFFFF" if idx % 2 == 0 else "F8F8F8"
            row = log_tbl.add_row().cells
            for cell, val in zip(row, [
                ts,
                entry.get("item_name", "")[:40],
                SOURCE_LABELS.get(entry.get("source_type", ""), entry.get("source_type", "")),
                _acct_map.get(entry.get("account_id", "")) or _resolve_display_name(entry.get("account_name", ""), entry.get("account_id", "")),
                str(entry.get("cpr_count", 0)),
                REASON_LABELS.get(entry.get("reason", ""), entry.get("reason", "")),
                entry.get("deleted_by", "") or "—",
            ]):
                _cell_bg(cell, bg)
                p = cell.paragraphs[0]; p.clear()
                r = p.add_run(val); r.font.size = Pt(7)

    # ── Section: Article 9 special categories ────────────────────────────────
    if special_items:
        last_sec += 1
        doc.add_page_break()
        _heading(f"{last_sec}. {L('a30_s_special', 'Special Category Data (Article 9)')}")

        _para(L("a30_special_intro",
                f"{len(special_items)} item(s) were detected as containing special category "
                f"data under GDPR Article 9. These require an explicit legal basis beyond "
                f"Article 6, and processing should be covered by a Data Protection Impact "
                f"Assessment (DPIA) under Article 35."),
              size=10, space_after=8)

        # Category breakdown table
        from collections import Counter as _Counter
        cat_counts: dict = _Counter()
        for item in special_items:
            sc = item.get("special_category", [])
            if isinstance(sc, str):
                import json as _scjson
                try:
                    sc = _scjson.loads(sc)
                except Exception:
                    sc = []
            for c in sc:
                cat_counts[c] += 1

        if cat_counts:
            _para(L("a30_special_by_cat", "Detected categories"), bold=True, size=10,
                  space_before=4, space_after=4)
            cat_tbl = doc.add_table(rows=1, cols=2)
            cat_tbl.style = "Table Grid"
            for cell, txt in zip(cat_tbl.rows[0].cells,
                                  [L("a30_col_category", "Category"),
                                   L("a30_col_count", "Items")]):
                _cell_bg(cell, _hex(DARK_BLUE))
                p = cell.paragraphs[0]; p.clear()
                r = p.add_run(txt); r.bold = True
                r.font.size = Pt(9); r.font.color.rgb = WHITE
            CAT_LABELS = {
                "health":           L("a30_cat_health",    "Health data (Art. 9)"),
                "mental_health":    L("a30_cat_mental",    "Mental health (Art. 9)"),
                "criminal":         L("a30_cat_criminal",  "Criminal records (Art. 10)"),
                "trade_union":      L("a30_cat_union",     "Trade union membership (Art. 9)"),
                "religion":         L("a30_cat_religion",  "Religious beliefs (Art. 9)"),
                "ethnicity":        L("a30_cat_ethnicity", "Racial/ethnic origin (Art. 9)"),
                "political":        L("a30_cat_political", "Political opinions (Art. 9)"),
                "biometric":        L("a30_cat_biometric", "Biometric data (Art. 9)"),
                "sexual_orientation": L("a30_cat_sexual",  "Sexual orientation (Art. 9)"),
            }
            for cat, count in sorted(cat_counts.items(), key=lambda x: -x[1]):
                row = cat_tbl.add_row().cells
                for cell, val in zip(row, [CAT_LABELS.get(cat, cat), str(count)]):
                    p = cell.paragraphs[0]; p.clear()
                    r = p.add_run(val); r.font.size = Pt(9)

        # Item list (capped at 50)
        _para(L("a30_special_items", "Affected items (up to 50)"), bold=True, size=10,
              space_before=10, space_after=4)
        sc_tbl = doc.add_table(rows=1, cols=5)
        sc_tbl.style = "Table Grid"
        for cell, txt in zip(sc_tbl.rows[0].cells, [
            L("a30_col_name",     "Name"),
            L("a30_col_account",  "Account"),
            L("a30_col_source",   "Source"),
            L("a30_col_category", "Category"),
            L("a30_col_cpr",      "CPR hits"),
        ]):
            _cell_bg(cell, _hex(DARK_BLUE))
            p = cell.paragraphs[0]; p.clear()
            r = p.add_run(txt); r.bold = True
            r.font.size = Pt(8); r.font.color.rgb = WHITE

        for idx, item in enumerate(special_items[:50]):
            bg = "FFFFFF" if idx % 2 == 0 else "FFF0F8"
            sc = item.get("special_category", [])
            if isinstance(sc, str):
                try:
                    import json as _scj2; sc = _scj2.loads(sc)
                except Exception:
                    sc = []
            row = sc_tbl.add_row().cells
            for cell, val in zip(row, [
                item.get("name", "")[:35],
                _acct_map.get(item.get("account_id", "")) or item.get("account_name", ""),
                SOURCE_LABELS.get(item.get("source_type", ""), item.get("source_type", "")),
                ", ".join(CAT_LABELS.get(c, c) for c in sc)[:45],
                str(item.get("cpr_count", 0)),
            ]):
                _cell_bg(cell, bg)
                p = cell.paragraphs[0]; p.clear()
                r = p.add_run(val); r.font.size = Pt(7)

    # ── Section: Photographs / biometric data (#9) ───────────────────────────
    if photo_items:
        last_sec += 1
        doc.add_page_break()
        _heading(f"{last_sec}. {L('a30_s_photos', 'Photographs and Biometric Data (Article 9)')}")

        total_faces = sum(i.get("face_count", 0) for i in photo_items)
        _para(L("a30_photo_intro",
                f"{len(photo_items)} image file(s) containing {total_faces} detected face(s) "
                f"were found in the scan. Photographs of identifiable persons constitute "
                f"biometric data under GDPR Article 9 and are subject to the same "
                f"heightened protection as health or criminal records data."),
              size=10, space_after=8)

        _para(L("a30_photo_guidance", "Retention guidance"), bold=True, size=10,
              space_before=4, space_after=4)
        for line in [
            L("a30_photo_g1",
              "Photos may only be retained while the original purpose remains valid "
              "(Art. 5(1)(b) — purpose limitation)."),
            L("a30_photo_g2",
              "Pupils under 15 require parental consent (Databeskyttelsesloven §6). "
              "Consent must be freely given, specific, and documented."),
            L("a30_photo_g3",
              "Photos on public-facing websites must be removed promptly after a person "
              "leaves the organisation or withdraws consent (Art. 17 — right to erasure)."),
            L("a30_photo_g4",
              "Historical/archive use may justify longer retention under Art. 89 only "
              "with specific safeguards and case-by-case assessment."),
        ]:
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(line); r.font.size = Pt(9)

        # GPS items sub-section
        if gps_items:
            _para(L("a30_gps_title", "Items with GPS location data"), bold=True, size=10,
                  space_before=10, space_after=4)
            _para(L("a30_gps_intro",
                    "The following files contain GPS coordinates embedded in EXIF metadata. "
                    "Location data constitutes personal data under Art. 4 GDPR. For photos of children "
                    "or staff, GPS data may reveal sensitive patterns (home address, health institution, "
                    "religious site). Consider stripping EXIF before sharing or publishing."),
                  size=9, space_after=6)
            gps_tbl = doc.add_table(rows=1, cols=4)
            gps_tbl.style = "Table Grid"
            for cell, txt in zip(gps_tbl.rows[0].cells, [
                L("a30_col_name", "Name"),
                L("a30_gps_col_lat", "Latitude"),
                L("a30_gps_col_lon", "Longitude"),
                L("a30_col_date", "Modified"),
            ]):
                _cell_bg(cell, _hex(DARK_BLUE))
                p = cell.paragraphs[0]; p.clear()
                r = p.add_run(txt); r.bold = True
                r.font.size = Pt(8); r.font.color.rgb = WHITE
            for idx, item in enumerate(gps_items[:50]):
                bg = "FFFFFF" if idx % 2 == 0 else "E8F7FF"
                row = gps_tbl.add_row().cells
                exif = item.get("exif") or {}
                gps  = exif.get("gps") or {}
                for cell, val in zip(row, [
                    item.get("name", "")[:40],
                    str(gps.get("lat", ""))[:12],
                    str(gps.get("lon", ""))[:12],
                    item.get("modified", ""),
                ]):
                    _cell_bg(cell, bg)
                    p = cell.paragraphs[0]; p.clear()
                    r = p.add_run(val); r.font.size = Pt(7)

        # Photo item list (capped at 50)
        _para(L("a30_photo_items", "Detected photo items (up to 50)"), bold=True, size=10,
              space_before=10, space_after=4)
        ph_tbl = doc.add_table(rows=1, cols=6)
        ph_tbl.style = "Table Grid"
        for cell, txt in zip(ph_tbl.rows[0].cells, [
            L("a30_col_name",    "Name"),
            L("a30_col_account", "Account"),
            L("a30_col_source",  "Source"),
            L("a30_photo_col_faces", "Faces"),
            L("a30_gps_col",     "GPS"),
            L("a30_col_date",    "Modified"),
        ]):
            _cell_bg(cell, _hex(DARK_BLUE))
            p = cell.paragraphs[0]; p.clear()
            r = p.add_run(txt); r.bold = True
            r.font.size = Pt(8); r.font.color.rgb = WHITE

        for idx, item in enumerate(photo_items[:50]):
            bg = "FFFFFF" if idx % 2 == 0 else "E8F7FF"
            row = ph_tbl.add_row().cells
            for cell, val in zip(row, [
                item.get("name", "")[:40],
                _acct_map.get(item.get("account_id", "")) or item.get("account_name", ""),
                SOURCE_LABELS.get(item.get("source_type", ""), item.get("source_type", "")),
                str(item.get("face_count", 0)),
                "✔" if (item.get("exif") or {}).get("gps") else "",
                item.get("modified", ""),
            ]):
                _cell_bg(cell, bg)
                p = cell.paragraphs[0]; p.clear()
                r = p.add_run(val); r.font.size = Pt(7)

    # ── Section: Methodology ─────────────────────────────────────────────────
    # last_sec already reflects all optional sections that were added above
    doc.add_page_break()
    _heading(f"{last_sec}. {L('a30_s6_short', 'Methodology and Legal Basis')}")

    _para(L("a30_method_title", "Scanning methodology"), bold=True, size=11, space_before=6, space_after=4)
    for line in [
        L("a30_method_1", "CPR numbers are detected using pattern matching against the official Danish CPR format (DDMMYY-XXXX)."),
        L("a30_method_2", "Additional personal data (phone numbers, email addresses, IBANs, bank accounts, names, addresses, and organisations) is detected using regular expressions and spaCy NER."),
        L("a30_method_3", "CPR numbers stored in this document's database are SHA-256 hashed and never stored in plaintext."),
        L("a30_method_4", "Scanning covers Exchange mailboxes (all folders including Sent Items), OneDrive, SharePoint, and Microsoft Teams channel files via the Microsoft Graph API. When connected, Google Workspace scanning covers Gmail and Google Drive via a service account with domain-wide delegation."),
        L("a30_method_5", "When photo scanning is enabled, image files are analysed using OpenCV Haar cascade face detection to identify photographs of persons (Art. 9 biometric data)."),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(line); r.font.size = Pt(10)

    _para(L("a30_gdpr_title", "GDPR Articles referenced"), bold=True, size=11, space_before=10, space_after=4)
    for line in [
        L("a30_gdpr_1", "Article 5(1)(c) — Data minimisation: only necessary data should be retained"),
        L("a30_gdpr_2", "Article 5(1)(e) — Storage limitation: data must not be kept longer than necessary"),
        L("a30_gdpr_3", "Article 9 — Special categories: health, criminal, trade union, and similar data require explicit legal basis"),
        L("a30_gdpr_4", "Article 15 — Right of access: data subjects may request information about their data"),
        L("a30_gdpr_5", "Article 17 — Right to erasure: data subjects may request deletion"),
        L("a30_gdpr_6", "Article 30 — Records of processing activities: this document satisfies the obligation"),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(line); r.font.size = Pt(10)

    _para(f"{L('a30_generated','Generated')}: {now_str}  ·  GDPRScanner  ·  {L('a30_confidential','Confidential — GDPR compliance document')}",
          size=9, color=RGBColor(0x88, 0x88, 0x88), align=WD_ALIGN_PARAGRAPH.CENTER, space_before=20)

    # ── Serialise ─────────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read(), fname






@app.route("/api/local_ip")
def local_ip():
    """Return the machine's LAN IP so viewer links point to a routable address."""
    try:
        with socket.socket(socket.AF_INET, socket.SOCK_DGRAM) as _s:
            _s.connect(("8.8.8.8", 80))
            ip = _s.getsockname()[0]
    except Exception:
        ip = "127.0.0.1"
    return jsonify({"ip": ip})


@app.route("/api/scan/stream")
def scan_stream():
    q = queue.Queue(maxsize=512)
    _sse_queues.append(q)
    # Filter replay buffer: only include events from the current scan
    # (avoids replaying stale events from a previous scan)
    replay_scan_id = _sse_mod._current_scan_id
    buf = []
    if replay_scan_id:
        for msg in list(_sse_buffer):
            if f'"scan_id": "{replay_scan_id}"' in msg:
                buf.append(msg)
    else:
        buf = list(_sse_buffer)
    def generate():
        try:
            yield ": connected\n\n"
            if buf:
                yield f"event: sse_replay\ndata: {{\"count\": {len(buf)}}}\n\n"
            for msg in buf:
                yield msg
            if buf:
                yield "event: sse_replay_done\ndata: {}\n\n"
            logger.debug("[SSE] generator live, q_id=%d, replayed=%d", id(q), len(buf))
            while True:
                try:
                    msg = q.get(timeout=5)
                    yield msg
                except queue.Empty:
                    yield ": heartbeat\n\n"
        except GeneratorExit:
            pass
        finally:
            if q in _sse_queues:
                _sse_queues.remove(q)
    return Response(generate(), mimetype="text/event-stream",
                    headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})



# ── Blueprint registration ────────────────────────────────────────────────────
from routes.auth      import bp as auth_bp
from routes.users     import bp as users_bp
from routes.scan      import bp as scan_bp
from routes.sources   import bp as sources_bp
from routes.profiles  import bp as profiles_bp
from routes.email     import bp as email_bp, _send_report_email
from routes.database  import bp as database_bp
from routes.export    import bp as export_bp
from routes.app_routes import bp as app_routes_bp
from routes.scheduler import bp as scheduler_bp
from routes.google_auth import bp as google_auth_bp
from routes.google_scan import bp as google_scan_bp
from routes.viewer      import bp as viewer_bp

for _bp in [auth_bp, users_bp, scan_bp, sources_bp, profiles_bp,
            email_bp, database_bp, export_bp, app_routes_bp, scheduler_bp,
            google_auth_bp, google_scan_bp, viewer_bp]:
    app.register_blueprint(_bp)

# ── Entry point ───────────────────────────────────────────────────────────────
# ── Entry point ───────────────────────────────────────────────────────────────
if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="M365 CPR Scanner",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Headless (scheduled) usage:
  python gdpr_scanner.py --headless --output ~/Reports/

  Auth credentials (Application mode) must be provided via:
    environment variables:  M365_CLIENT_ID, M365_TENANT_ID, M365_CLIENT_SECRET
    or a settings JSON:     --settings /path/to/settings.json

  Scan options are loaded from ~/.gdpr_scanner_settings.json (saved automatically
  after any interactive scan), or overridden in the --settings file.

  SMTP config is loaded from ~/.gdpr_scanner_smtp.json (saved in the UI) or from
  an 'smtp' key in the --settings file.

Example cron (weekly, Mondays at 06:00):
  0 6 * * 1 M365_CLIENT_ID=... M365_TENANT_ID=... M365_CLIENT_SECRET=... \\
            python /path/to/gdpr_scanner.py --headless --output /reports/ \\
            --email-to compliance@company.com,ciso@company.com

Example Windows Task Scheduler (run batch file):
  m365_scan.bat:
    set M365_CLIENT_ID=<id>
    set M365_TENANT_ID=<tid>
    set M365_CLIENT_SECRET=<secret>
    python gdpr_scanner.py --headless --output C:\\Reports\\ --email-to compliance@company.com

Example --settings file with SMTP:
  {
    "client_id": "...", "tenant_id": "...", "client_secret": "...",
    "sources": ["email", "onedrive"],
    "options": {"older_than_days": 365, "delta": true},
    "smtp": {
      "host": "smtp.office365.com", "port": 587,
      "username": "scanner@company.com", "password": "...",
      "use_tls": true
    }
  }
""",
    )
    parser.add_argument("--port",     type=int, default=5100)
    parser.add_argument("--host",     default="0.0.0.0")
    parser.add_argument("--headless", action="store_true",
                        help="Run a non-interactive scan and export Excel, then exit")
    parser.add_argument("--output",   default=".",
                        help="Output directory for Excel export in headless mode (default: .)")
    parser.add_argument("--settings", default=None,
                        help="Path to a JSON settings file (overrides ~/.gdpr_scanner_settings.json)")
    parser.add_argument("--email-to", default=None,
                        help="Comma-separated recipient addresses — send Excel report by email (headless only)")
    parser.add_argument("--retention-years", type=int, default=None,
                        help="Auto-delete items older than N years after headless scan (requires --headless)")
    parser.add_argument("--fiscal-year-end", default=None,
                        help="Fiscal year end as MM-DD for retention cutoff (e.g. 12-31 for Bogforingsloven). Omit for rolling window.")
    parser.add_argument("--reset-db", action="store_true",
                        help="Reset the results database (~/.gdpr_scanner.db) — permanently deletes all scan history, "
                             "dispositions, and deletion log. Prompts for confirmation unless --yes is also passed.")
    parser.add_argument("--yes", action="store_true",
                        help="Skip confirmation prompts (use with --reset-db for scripted resets)")
    parser.add_argument("--purge", action="store_true",
                        help="Permanently delete all data files created by the scanner "
                             "(database, token cache, credentials, checkpoints, settings, OCR cache). "
                             "Use before decommissioning or moving to a new server. "
                             "Prompts for confirmation unless --yes is also passed.")
    parser.add_argument("--export-db", default=None, metavar="FILE",
                        help="Export the database to a ZIP archive (e.g. gdpr_export_2026.zip) and exit")
    parser.add_argument("--import-db", default=None, metavar="FILE",
                        help="Import a previously exported ZIP archive into the database and exit")
    parser.add_argument("--import-mode", default="merge", choices=["merge", "replace"],
                        help="Import mode: 'merge' (default) keeps existing data and adds dispositions/deletion log; "
                             "'replace' wipes the DB first then imports everything")
    parser.add_argument("--profile", default=None, metavar="NAME",
                        help="Name of the scan profile to use for headless mode")
    parser.add_argument("--list-profiles", action="store_true",
                        help="List all saved scan profiles and exit")
    parser.add_argument("--save-profile", default=None, metavar="NAME",
                        help="Save the provided CLI options as a named profile and exit")
    parser.add_argument("--delete-profile", default=None, metavar="NAME",
                        help="Delete a saved profile by name and exit")

    # ── File scanning CLI flags (#8) ──────────────────────────────────────────
    parser.add_argument("--scan-path", default=None, metavar="PATH",
                        help="Scan a local folder or SMB share for CPR numbers and PII. "
                             "Local: ~/Documents  SMB: //nas.school.dk/shares/staff")
    parser.add_argument("--smb-user", default=None, metavar="USER",
                        help="SMB username (e.g. DOMAIN\\username) for --scan-path on a network share")
    parser.add_argument("--smb-host", default=None, metavar="HOST",
                        help="SMB hostname (auto-detected from --scan-path if not specified)")
    parser.add_argument("--smb-domain", default=None, metavar="DOMAIN",
                        help="SMB/Windows domain (optional, included in --smb-user as DOMAIN\\user)")
    parser.add_argument("--smb-keychain-key", default=None, metavar="KEY",
                        help="Account name used to retrieve the SMB password from the OS keychain")
    parser.add_argument("--smb-store-creds", action="store_true",
                        help="Store SMB credentials in the OS keychain and exit. "
                             "Requires --smb-host and --smb-user. Prompts for password interactively.")
    parser.add_argument("--scan-label", default=None, metavar="LABEL",
                        help="Display label for --scan-path results (defaults to the path)")
    parser.add_argument("--scan-photos", action="store_true",
                        help="Enable face detection on image files during --scan-path scan (slower)")
    parser.add_argument("--max-file-mb", default=50, type=int, metavar="MB",
                        help="Maximum file size in MB to scan (default: 50). "
                             "Files larger than this are skipped.")

    args = parser.parse_args()

    # ── File scan CLI flags (#8) ─────────────────────────────────────────────
    if getattr(args, "smb_store_creds", False):
        if not FILE_SCANNER_OK:
            print("ERROR: file_scanner.py not found — cannot store credentials.")
            sys.exit(1)
        smb_host = getattr(args, "smb_host", None) or ""
        smb_user = getattr(args, "smb_user", None) or ""
        if not smb_user:
            print("ERROR: --smb-user required with --smb-store-creds")
            sys.exit(1)
        import getpass
        pw = getpass.getpass(f"SMB password for {smb_user}@{smb_host}: ")
        key = getattr(args, "smb_keychain_key", None) or smb_user
        ok = store_smb_password(smb_host, smb_user, pw, key)
        if ok:
            print(f"  [ok] Credentials stored in OS keychain (service=gdpr-scanner-nas, account={key})")
        else:
            print("  [warn] keyring not available — install: pip install keyring")
        sys.exit(0)

    if getattr(args, "scan_path", None):
        if not FILE_SCANNER_OK:
            print("ERROR: file_scanner.py not found — cannot scan file system.")
            sys.exit(1)
        source = {
            "path":         args.scan_path,
            "label":        getattr(args, "scan_label", None) or args.scan_path,
            "smb_host":     getattr(args, "smb_host", None) or "",
            "smb_user":     getattr(args, "smb_user", None) or "",
            "smb_domain":   getattr(args, "smb_domain", None) or "",
            "keychain_key": getattr(args, "smb_keychain_key", None) or "",
            "scan_photos":  bool(getattr(args, "scan_photos", False)),
            "max_file_mb":  int(getattr(args, "max_file_mb", 50)),
        }
        print(f"[file scan] {source['label']}")
        run_file_scan(source)
        # Write Excel report if output path provided
        if getattr(args, "output", None) and flagged_items:
            try:
                out_path = _write_excel_report(args.output)
                if out_path:
                    print(f"[file scan] report: {out_path}")
            except Exception as e:
                print(f"[file scan] report failed: {e}")
        sys.exit(0)

    # ── Profile management (15b) ──────────────────────────────────────────────
    if getattr(args, "list_profiles", False):
        import sys as _sys
        profiles = _profiles_load()
        if not profiles:
            print("  No profiles saved. Run a scan first, or use --save-profile to create one.")
        else:
            print(f"\n  {'#':<4} {'Name':<30} {'Sources':<30} {'Last run':<20} {'Scan ID'}")
            print(f"  {'-'*4} {'-'*30} {'-'*30} {'-'*20} {'-'*8}")
            for i, p in enumerate(profiles, 1):
                srcs    = ", ".join(p.get("sources", [])) or "—"
                last    = (p.get("last_run") or "never")[:19]
                scan_id = str(p.get("last_scan_id") or "—")
                print(f"  {i:<4} {p.get('name',''):<30} {srcs:<30} {last:<20} {scan_id}")
                if p.get("description"):
                    print(f"       {p['description']}")
        print()
        _sys.exit(0)

    if getattr(args, "save_profile", None):
        import sys as _sys
        name = args.save_profile
        # Build profile from CLI args
        sources = []
        if getattr(args, "sources", None):
            sources = [s.strip() for s in args.sources.split(",") if s.strip()]
        profile = _profile_from_settings({
            "sources":         sources,
            "user_ids":        [],
            "options":         {
                "email_body":    True,
                "attachments":   getattr(args, "attachments", False),
                "older_than_days": 0,
            },
            "retention_years": getattr(args, "retention_years", None),
            "fiscal_year_end": getattr(args, "fiscal_year_end", None),
            "email_to":        getattr(args, "email_to", "") or "",
        }, name=name)
        existing = _profile_get(name)
        if existing:
            profile["id"] = existing["id"]
        saved = _profile_save(profile)
        print(f"\n  ✔ Profile '{name}' saved (id: {saved['id']})")
        print(f"    Sources:   {', '.join(saved.get('sources', [])) or 'none'}")
        if saved.get("email_to"):
            print(f"    Email to:  {saved['email_to']}")
        if saved.get("retention_years"):
            print(f"    Retention: {saved['retention_years']} years")
        print()
        _sys.exit(0)

    if getattr(args, "delete_profile", None):
        import sys as _sys
        name = args.delete_profile
        ok   = _profile_delete(name)
        if ok:
            print(f"\n  ✔ Profile '{name}' deleted.\n")
        else:
            print(f"\n  ✖ Profile '{name}' not found.\n")
            print("  Available profiles:")
            for p in _profiles_load():
                print(f"    • {p.get('name')}")
            print()
        _sys.exit(0)

    # ── Resolve --profile for headless mode ───────────────────────────────────
    _active_profile_id: str | None = None
    if getattr(args, "profile", None) and args.headless:
        import sys as _sys
        p = _profile_get(args.profile)
        if not p:
            print(f"\n  ✖ Profile '{args.profile}' not found.\n")
            print("  Available profiles:")
            for pr in _profiles_load():
                print(f"    • {pr.get('name')}")
            print()
            _sys.exit(1)
        # Populate args from profile (profile overrides individual CLI flags)
        _active_profile_id = p["id"]
        if p.get("sources"):
            args.sources = ",".join(p["sources"])  # used by headless scan builder
        if p.get("retention_years") and not args.retention_years:
            args.retention_years = p["retention_years"]
        if p.get("fiscal_year_end") and not args.fiscal_year_end:
            args.fiscal_year_end = p["fiscal_year_end"]
        if p.get("email_to") and not args.email_to:
            args.email_to = p["email_to"]
        print(f"\n  Profile: '{p['name']}'")
        if p.get("description"):
            print(f"  {p['description']}")
        if p.get("last_run"):
            print(f"  Last run: {p['last_run'][:19]}")
        print()

    # ── Purge all scanner data files ─────────────────────────────────────────
    if getattr(args, "purge", False):
        import sys as _sys
        from gdpr_db import DB_PATH as _DB_PATH

        # All files created by either scanner
        PURGE_FILES = [
            # GDPRScanner
            (_DB_PATH,                                              "SQLite results database"),
            (_CONFIG_FILE,                                          "Azure app credentials"),
            (_SMTP_CONFIG_PATH,                                     "SMTP credentials"),
            (_SETTINGS_PATH,                                        "Headless scan settings"),
            (_ROLE_OVERRIDES_PATH,                                  "Manual role overrides"),
            (_FILE_SOURCES_PATH,                                    "File source definitions"),
            (_CHECKPOINT_PATH,                                      "Scan checkpoint (resume state)"),
            (_DELTA_PATH,                                           "Delta scan tokens"),
            (_LANG_OVERRIDE_FILE,                                   "Language preference"),
            (Path.home() / ".gdprscanner" / "schedule.json",           "Scheduler configuration"),
            # Document Scanner
            (Path.home() / ".document_scanner_ocr_cache.db",       "OCR cache"),
            (Path.home() / ".document_scanner_lang",               "Document Scanner language preference"),
            # MSAL token cache (created by msal library)
            (Path.home() / ".gdprscanner" / "msal_cache.bin",         "MSAL token cache"),
        ]

        print("\n  ── GDPR Scanner — Purge data files ──────────────────────────────")
        print("  This will permanently delete all data files created by the scanner.")
        print("  No scan results, credentials, or cached data will remain.\n")

        existing = [(p, desc) for p, desc in PURGE_FILES if p.exists()]
        if not existing:
            print("  No scanner data files found — nothing to delete.")
            _sys.exit(0)

        total_kb = sum(p.stat().st_size for p, _ in existing) / 1024
        print(f"  Files to delete ({len(existing)}, {total_kb:.0f} KB total):")
        for p, desc in existing:
            kb = p.stat().st_size / 1024
            print(f"    {desc:40s} {p.name}  ({kb:.0f} KB)")

        print()
        if not getattr(args, "yes", False):
            print("  ⚠  This cannot be undone. Export the database first if you need a record.")
            answer = input("  Type 'yes' to confirm: ").strip().lower()
            if answer != "yes":
                print("  Cancelled — no files deleted.")
                _sys.exit(0)

        deleted = 0
        failed  = 0
        for p, desc in existing:
            try:
                p.unlink()
                print(f"  ✔ Deleted: {p}")
                deleted += 1
            except Exception as e:
                print(f"  ✖ Failed:  {p} — {e}")
                failed += 1

        print(f"\n  Purge complete: {deleted} deleted, {failed} failed.")
        if failed == 0:
            print("  The scanner has left no data files on this machine.")
        _sys.exit(0)


    if args.reset_db:
        import sys as _sys
        from gdpr_db import DB_PATH as _DB_PATH
        db_path = _DB_PATH
        print(f"\n  Database reset requested: {db_path}")
        if db_path.exists():
            size_kb = round(db_path.stat().st_size / 1024, 1)
            print(f"  Current size: {size_kb} KB")
        else:
            print("  (database file does not exist yet — nothing to reset)")
            _sys.exit(0)

        if not args.yes:
            print("\n  ⚠  This will permanently delete:")
            print("       • All scan results and flagged items")
            print("       • CPR index and PII hit counts")
            print("       • All compliance dispositions")
            print("       • Deletion audit log")
            print("       • Scan history and trend data")
            print()
            answer = input("  Type 'yes' to confirm: ").strip().lower()
            if answer != "yes":
                print("  Cancelled — database not modified.")
                _sys.exit(0)

        if DB_OK:
            try:
                _get_db().reset()
                print(f"  ✔ Database reset complete: {db_path}")
            except Exception as e:
                print(f"  ✖ Reset failed: {e}")
                _sys.exit(1)
        else:
            print("  ✖ m365_db not available — cannot reset")
            _sys.exit(1)

        # Also clear the JSON checkpoint so the UI starts with no cached results
        _clear_checkpoint()
        if not _CHECKPOINT_PATH.exists():
            print(f"  ✔ Checkpoint cleared")

        # Clear delta tokens too — stale after a full DB reset
        if _DELTA_PATH.exists():
            _DELTA_PATH.unlink()
            print(f"  ✔ Delta tokens cleared")

        if not args.headless:
            _sys.exit(0)  # reset-only — done

    # ── Export database ───────────────────────────────────────────────────────
    if getattr(args, "export_db", None):
        import sys as _sys
        if not DB_OK:
            print("  ✖ m365_db not available — cannot export")
            _sys.exit(1)
        out = Path(args.export_db)
        print(f"\n  Exporting database to: {out}")
        try:
            meta = _get_db().export_db(out)
            print(f"  ✔ Export complete: {out}")
            print(f"  Exported at: {meta['exported_at']}")
            for table, count in meta["row_counts"].items():
                if count:
                    print(f"    {table:20s} {count} rows")
            print(f"  Size: {out.stat().st_size / 1024:.0f} KB")
        except Exception as e:
            print(f"  ✖ Export failed: {e}")
            _sys.exit(1)
        _sys.exit(0)

    # ── Import database ───────────────────────────────────────────────────────
    if getattr(args, "import_db", None):
        import sys as _sys
        if not DB_OK:
            print("  ✖ m365_db not available — cannot import")
            _sys.exit(1)
        src  = Path(args.import_db)
        mode = getattr(args, "import_mode", "merge")
        print(f"\n  Importing from: {src}")
        print(f"  Mode: {mode}")
        if mode == "replace":
            print("  ⚠  Replace mode will wipe the current database first.")
            if not getattr(args, "yes", False):
                answer = input("  Type 'yes' to confirm: ").strip().lower()
                if answer != "yes":
                    print("  Cancelled — database not modified.")
                    _sys.exit(0)
        try:
            result = _get_db().import_db(src, mode=mode)
            print(f"  ✔ Import complete ({mode} mode)")
            print(f"  Source export date: {result.get('exported_at', 'unknown')}")
            for table, count in result["imported"].items():
                if count:
                    print(f"    {table:20s} {count} rows imported")
        except Exception as e:
            print(f"  ✖ Import failed: {e}")
            _sys.exit(1)
        _sys.exit(0)

    if not MSAL_OK:
        print("⚠  msal not installed — run: pip install msal requests")
    if not SCANNER_OK:
        print("⚠  document_scanner not found — CPR scanning unavailable")

    if args.headless:
        # ── Headless / scheduled mode ─────────────────────────────────────────
        import sys as _sys
        print("\n  GDPRScanner — Headless mode")
        print("  ─────────────────────────────────────────")

        # Load settings from --settings file or saved defaults
        cfg: dict = {}
        if args.settings:
            try:
                cfg = json.loads(Path(args.settings).read_text(encoding="utf-8"))
                print(f"  Settings loaded from: {args.settings}")
            except Exception as e:
                print(f"  ✖ Cannot read settings file: {e}")
                _sys.exit(1)
        else:
            saved = _load_settings()
            if saved:
                cfg = saved
                print(f"  Settings loaded from: {_SETTINGS_PATH}")
            else:
                print(f"  ✖ No saved settings found. Run an interactive scan first, or provide --settings.")
                _sys.exit(1)

        # Auth credentials from environment or settings file
        client_id     = cfg.get("client_id")     or os.environ.get("M365_CLIENT_ID", "")
        tenant_id     = cfg.get("tenant_id")     or os.environ.get("M365_TENANT_ID", "")
        client_secret = cfg.get("client_secret") or os.environ.get("M365_CLIENT_SECRET", "")

        if not all([client_id, tenant_id, client_secret]):
            print("  ✖ App credentials required for headless mode.")
            print("    Set M365_CLIENT_ID, M365_TENANT_ID, M365_CLIENT_SECRET")
            print("    or include client_id / tenant_id / client_secret in --settings JSON.")
            _sys.exit(1)

        # Authenticate
        try:
            from m365_connector import M365Connector
            conn = M365Connector(client_id, tenant_id, client_secret=client_secret)
            conn.authenticate_app_mode()
            print("  ✔ Authenticated (Application / client credentials)")
        except Exception as e:
            print(f"  ✖ Authentication failed: {e}")
            _sys.exit(1)

        # Set connector in module globals (works whether running as __main__ or imported)
        _mod = _sys.modules[__name__]
        _mod._connector = conn

        # Build scan options from config
        sources  = cfg.get("sources", ["email", "onedrive"])
        user_ids = cfg.get("user_ids", [])
        opts     = cfg.get("options",  {})

        if not user_ids:
            # Default: scan all tenant users
            print("  No user_ids in settings — fetching all tenant users…")
            try:
                all_users = conn.list_users()
                user_ids  = [{"id": u["id"],
                              "displayName": _resolve_display_name(
                                  u.get("displayName", ""),
                                  u.get("mail") or u.get("userPrincipalName", ""))}
                             for u in all_users if u.get("id")]
                print(f"  Found {len(user_ids)} users")
            except Exception as e:
                print(f"  ✖ Could not list users: {e}")
                _sys.exit(1)

        scan_options = {
            "sources":  sources,
            "user_ids": user_ids,
            "options":  opts,
        }

        # Print scan summary
        print(f"  Sources: {', '.join(sources)}")
        print(f"  Users:   {len(user_ids)}")
        older = opts.get("older_than_days", 0)
        print(f"  Cutoff:  {'%d days' % older if older else 'All'}")
        print("  Scanning…\n")

        # Replace broadcast with a stdout logger for headless mode
        def _headless_broadcast(event: str, data: dict):
            if event == "scan_phase":
                print(f"  {data.get('phase', '')}", flush=True)
            elif event == "scan_start":
                resumed = data.get("resumed", 0)
                total   = data.get("total", 0)
                msg = f"  Items to scan: {total}"
                if resumed:
                    msg += f"  ({resumed} skipped — already scanned)"
                print(msg, flush=True)
            elif event == "scan_progress":
                pct     = data.get("pct", 0)
                name    = data.get("file", "")[:55]
                eta     = data.get("eta", "")
                bar     = "█" * (pct // 5) + "░" * (20 - pct // 5)
                eta_str = f"  {eta} left" if eta else ""
                print(f"\r  [{bar}] {pct:3d}%  {name:<55}{eta_str}", end="", flush=True)
            elif event == "scan_file_flagged":
                print(f"\n  ✔ {data.get('name', '')} — {data.get('cpr_count', 0)} CPR", flush=True)
            elif event == "scan_done":
                print(f"\n\n  Done — {data.get('flagged_count', 0)} flagged / {data.get('total_scanned', 0)} scanned", flush=True)
            elif event == "scan_error":
                print(f"\n  ✖ {data.get('file', '')}: {data.get('error', '')}", flush=True)
            elif event == "scan_cancelled":
                print(f"\n  Scan stopped after {data.get('completed', 0)} items.", flush=True)

        _orig_broadcast = _mod.broadcast
        _mod.broadcast  = _headless_broadcast

        try:
            run_scan(scan_options)
        except Exception as e:
            print(f"\n  ✖ Scan error: {e}")
            _sys.exit(1)
        finally:
            _mod.broadcast = _orig_broadcast

        if not flagged_items:
            print("  No flagged items — no Excel file written.")
            _sys.exit(0)

        # Export Excel
        out_dir = Path(args.output).expanduser()
        out_dir.mkdir(parents=True, exist_ok=True)
        import datetime as _dt
        fname    = f"m365_scan_{_dt.datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        out_path = out_dir / fname

        try:
            xl_bytes, fname = _build_excel_bytes()
            out_path = out_dir / fname
            out_path.write_bytes(xl_bytes)
            print(f"  Excel saved: {out_path}")
        except Exception as e:
            print(f"  ✖ Excel export failed: {e}")
            _sys.exit(1)

        # ── Email the report if --email-to was specified ──────────────────────
        email_to = getattr(args, "email_to", None)
        if email_to:
            recipients = [r.strip() for r in email_to.replace(";", ",").split(",") if r.strip()]
            # SMTP config: --settings file takes priority, then saved ~/.gdpr_scanner_smtp.json
            smtp_cfg = _load_smtp_config()
            if cfg.get("smtp"):
                smtp_cfg = {**smtp_cfg, **cfg["smtp"]}
            if not smtp_cfg.get("host"):
                print("  ✖ Cannot send email — no SMTP config found.")
                print("    Configure SMTP in the UI (✉ Email report panel) or add an 'smtp' key to --settings.")
            else:
                print(f"  Sending report to: {', '.join(recipients)}…")
                try:
                    _send_report_email(xl_bytes, fname, smtp_cfg, recipients)
                    print(f"  ✔ Report emailed to {', '.join(recipients)}")
                except Exception as e:
                    print(f"  ✖ Email send failed: {e}")
                    # Don't exit 1 — the Excel file was saved successfully

        # ── Retention auto-delete if --retention-years was specified ──────────
        retention_years   = getattr(args, "retention_years", None)
        fiscal_year_end   = getattr(args, "fiscal_year_end", None)
        if retention_years and DB_OK:
            try:
                from gdpr_db import overdue_cutoff
                cutoff = overdue_cutoff(retention_years, fiscal_year_end)
                overdue_items = _get_db().get_overdue_items(
                    retention_years, fiscal_year_end=fiscal_year_end
                )
                mode_str = f"fiscal year end {fiscal_year_end}" if fiscal_year_end else "rolling"
                print(f"\n  Retention policy: {retention_years} years ({mode_str})")
                print(f"  Cutoff date:      {cutoff}")
                print(f"  Overdue items:    {len(overdue_items)}")

                if not overdue_items:
                    print("  No overdue items to delete.")
                else:
                    # Confirm unless --yes / non-interactive
                    import sys as _sys2
                    if _sys2.stdin.isatty():
                        answer = input(f"\n  Delete {len(overdue_items)} overdue item(s)? [y/N] ").strip().lower()
                        if answer != "y":
                            print("  Skipped — no items deleted.")
                        else:
                            _do_retention_delete(overdue_items)
                    else:
                        # Non-interactive (cron) — delete automatically
                        print("  Non-interactive mode — deleting automatically…")
                        _do_retention_delete(overdue_items)
            except Exception as e:
                print(f"  ✖ Retention check failed: {e}")

        # ── Auto-delete items tagged delete-scheduled in disposition table ────
        if DB_OK:
            try:
                db = _get_db()
                if db:
                    # Find all flagged items whose disposition is delete-scheduled
                    scheduled = [
                        item for item in flagged_items
                        if item.get("id") and (
                            lambda d: d and d.get("status") == "delete-scheduled"
                        )(db.get_disposition(item.get("id", "")))
                    ]
                    if scheduled:
                        print(f"\n  Disposition auto-delete: {len(scheduled)} item(s) tagged 'delete-scheduled'")
                        import sys as _sys2
                        if _sys2.stdin.isatty():
                            answer = input(f"  Delete {len(scheduled)} scheduled item(s)? [y/N] ").strip().lower()
                            if answer != "y":
                                print("  Skipped.")
                                scheduled = []
                        else:
                            print("  Non-interactive mode — deleting automatically…")
                        if scheduled:
                            _do_retention_delete(scheduled)
            except Exception as e:
                print(f"  ✖ Disposition auto-delete failed: {e}")

        # Update profile last_run if a named profile was used
        if _active_profile_id:
            try:
                sid = _get_db().latest_scan_id() if DB_OK else None
                _profile_touch(_active_profile_id, sid)
            except Exception:
                pass

        print("\n  ✔ Headless scan complete.\n")
        _sys.exit(0)

    else:
        # ── Interactive web UI mode ───────────────────────────────────────────
        # Single-instance guard — prevent two servers sharing the same DB/settings.
        _lock_fh = None
        def _acquire_lock() -> bool:
            global _lock_fh
            from app_config import _DATA_DIR
            _DATA_DIR.mkdir(parents=True, exist_ok=True)
            try:
                _lock_fh = open(_DATA_DIR / "app.lock", "w")
                if sys.platform == "win32":
                    import msvcrt as _msvcrt
                    _msvcrt.locking(_lock_fh.fileno(), _msvcrt.LK_NBLCK, 1)
                else:
                    import fcntl as _fcntl
                    _fcntl.flock(_lock_fh, _fcntl.LOCK_EX | _fcntl.LOCK_NB)
                _lock_fh.write(str(_os.getpid()))
                _lock_fh.flush()
                return True
            except (IOError, OSError):
                if _lock_fh:
                    _lock_fh.close()
                return False

        if not _acquire_lock():
            print("GDPRScanner is already running. Stop the existing instance first.", file=sys.stderr)
            sys.exit(1)

        # Find a free port — auto-increment from the requested port if in use.
        import socket as _socket
        def _find_free_port(start: int, host: str) -> int:
            for p in range(start, start + 100):
                with _socket.socket(_socket.AF_INET, _socket.SOCK_STREAM) as s:
                    try:
                        s.bind((host, p))
                        return p
                    except OSError:
                        continue
            raise RuntimeError(f"No free port found in range {start}–{start + 99}")

        actual_port = _find_free_port(args.port, args.host)
        if actual_port != args.port:
            print(f"  [!] Port {args.port} in use — using {actual_port} instead")
        args.port = actual_port
        # Machine-readable port line — parseable by a parent process via stdout.
        print(f"GDPR_PORT={args.port}", flush=True)

        print(f"\n  GDPRScanner\n  ──────────────────────────────")
        print(f"  Open: http://{args.host}:{args.port}")

        # Start in-process scheduler (#19)
        try:
            import scan_scheduler as _sched_mod
            scan_scheduler = _sched_mod.scan_scheduler
            if scan_scheduler.start():
                _sched_cfg = _sched_mod.load_schedule_config()
                if _sched_cfg.get("enabled"):
                    _nxt = scan_scheduler.next_run_time() or "—"
                    print(f"  Scheduler: enabled (next run: {_nxt})")
                else:
                    print("  Scheduler: disabled (enable in Settings → Scheduler)")
            else:
                print("  Scheduler: unavailable (pip install apscheduler)")
        except Exception as _sched_err:
            print(f"  Scheduler: failed to start ({_sched_err})")

        print(f"  Press Ctrl+C to stop\n")
        app.run(host=args.host, port=args.port, debug=False, threaded=True)
