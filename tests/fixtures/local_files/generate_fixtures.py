"""
Generate binary fixture files for the local-file GDPR scan test suite.

Run from repo root:
    source venv/bin/activate
    python tests/fixtures/local_files/generate_fixtures.py

Fixtures produced
─────────────────
Document fixtures (require python-docx + openpyxl):
  09_cpr_in_docx.docx   — Word document with 2 CPR numbers          → Flag
  13_cpr_in_xlsx.xlsx   — Excel workbook with CPR numbers            → Flag

Audio fixtures (require mutagen):
  14_audio_artist_pii.mp3  — MP3 with artist/title tags (personal name)    → Flag
  15_audio_artist_pii.flac — FLAC with artist/title Vorbis comments        → Flag
  16_audio_no_pii.mp3      — MP3 with no metadata tags                     → No flag
  17_audio_no_pii.flac     — FLAC with no metadata                         → No flag

Video fixtures (require mutagen):
  18_video_gps.mp4      — MP4 with GPS coordinates + artist tag       → Flag
  19_video_no_pii.mp4   — MP4 with no metadata tags                   → No flag
"""
import struct
import tempfile
import os
from pathlib import Path
import sys

HERE = Path(__file__).parent

def _require(pkg):
    try:
        return __import__(pkg)
    except ImportError:
        print(f"Missing: {pkg}  →  pip install {pkg}", file=sys.stderr)
        sys.exit(1)

openpyxl = _require("openpyxl")
docx = _require("docx")
_require("mutagen")

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ── 09_cpr_in_docx.docx ───────────────────────────────────────────────────────
def make_docx():
    doc = Document()

    doc.add_heading("Elevjournal — Gudenaaskolen", level=1)

    p = doc.add_paragraph()
    p.add_run("Dette dokument indeholder personoplysninger og er fortroligt.")
    p.runs[0].italic = True

    doc.add_heading("Elevoplysninger", level=2)
    # Use labelled paragraphs so CPR values are always preceded by ": " —
    # avoids the _CPR_PREFIX_NOISE guard that fires when table-cell runs are
    # concatenated without a separator.
    fields = [
        ("Navn",       "Magnus Lund Eriksen"),
        ("CPR-nummer", "010172-1019"),
        ("Klasse",     "8B"),
        ("Adresse",    "Egevej 3, 8680 Ry"),
        ("Telefon",    "+45 40 12 34 56"),
        ("E-mail",     "magnus.eriksen@elev.gudenaaskolen.dk"),
    ]
    for label, value in fields:
        p = doc.add_paragraph()
        run_label = p.add_run(f"{label}: ")
        run_label.bold = True
        p.add_run(value + " ")

    doc.add_heading("Forældrekontakt", level=2)
    doc.add_paragraph(
        "Forældrene er orienteret om elevens situation den 15. marts 2026. "
        "Begge forældre deltog i mødet. Næste opfølgning er planlagt til "
        "maj 2026."
    )

    doc.add_heading("Anden elev — tabel", level=2)
    doc.add_paragraph(
        "Nedenstående tabel viser en anden elev, der deler klasse med Magnus."
    )
    for label, value in [
        ("Navn",         "Nora Bjerrum Nielsen"),
        ("Personnummer", "280490-0120"),
        ("Klasse",       "8B"),
    ]:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(value + " ")

    doc.add_heading("Sagsbehandlernote", level=2)
    doc.add_paragraph(
        "Sagsbehandler: M. Andersen\n"
        "Dato: 20. april 2026\n"
        "Der er ikke fundet grundlag for yderligere foranstaltninger."
    )

    out = HERE / "09_cpr_in_docx.docx"
    doc.save(str(out))
    print(f"Written: {out.name}")


# ── 13_cpr_in_xlsx.xlsx ───────────────────────────────────────────────────────
def make_xlsx():
    wb = Workbook()

    # Sheet 1: Elevliste
    ws1 = wb.active
    ws1.title = "Elevliste"

    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="2B5F9E")

    headers = ["Klasse", "Navn", "CPR-nummer", "Adresse", "Forælder tlf", "Bemærkninger"]
    for col, h in enumerate(headers, 1):
        cell = ws1.cell(row=1, column=col, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")

    students = [
        ("7A", "Magnus Lund Eriksen",   "010172-1019", "Egevej 3, 8680 Ry",        "+45 40 12 34 56", ""),
        ("7A", "Nora Bjerrum Nielsen",  "280490-0120", "Møllevej 11, 8680 Ry",     "+45 50 23 45 67", "Brillebærer"),
        ("7A", "Oliver Skov Madsen",    "250372-0100", "Kirkegade 2, 8660 Skanderborg", "+45 60 34 56 78", ""),
        ("7B", "Rasmus Dal Kristensen", "150365-1102", "Rosenvej 5, 8680 Ry",       "+45 21 56 78 90", ""),
        ("7B", "Sofie Holm Thomsen",    "111111-1010", "Birkevej 22, 8660 Skanderborg", "+45 31 67 89 01", "Allergi: nødder"),
        ("7B", "Emil Sand Jensen",      "010107-4102", "Hybenvej 7, 8680 Ry",       "+45 41 78 90 12", ""),
    ]
    for row_i, row_data in enumerate(students, 2):
        for col_i, val in enumerate(row_data, 1):
            ws1.cell(row=row_i, column=col_i, value=val)

    for col in ws1.columns:
        max_len = max(len(str(c.value or "")) for c in col)
        ws1.column_dimensions[col[0].column_letter].width = max_len + 4

    # Sheet 2: Medarbejdere
    ws2 = wb.create_sheet("Medarbejdere")
    emp_headers = ["ID", "Navn", "Personnummer", "Afdeling", "E-mail"]
    for col, h in enumerate(emp_headers, 1):
        cell = ws2.cell(row=1, column=col, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")

    employees = [
        ("EMP-001", "Christian Bøgh Hansen",  "150365-1102", "Ledelse",        "c.hansen@gudenaaskolen.dk"),
        ("EMP-002", "Mette Dahl Andersen",     "280490-0120", "Administration", "m.andersen@gudenaaskolen.dk"),
        ("EMP-003", "Søren Lykke Jakobsen",    "010172-1019", "Pædagogik",      "s.jakobsen@gudenaaskolen.dk"),
    ]
    for row_i, row_data in enumerate(employees, 2):
        for col_i, val in enumerate(row_data, 1):
            ws2.cell(row=row_i, column=col_i, value=val)

    for col in ws2.columns:
        max_len = max(len(str(c.value or "")) for c in col)
        ws2.column_dimensions[col[0].column_letter].width = max_len + 4

    out = HERE / "13_cpr_in_xlsx.xlsx"
    wb.save(str(out))
    print(f"Written: {out.name}")


# ── Audio / video helpers ─────────────────────────────────────────────────────

# Two silent MPEG1 Layer3 frames (128 kbps / 44100 Hz / mono).
# mutagen needs at least 2 consecutive frame headers to confirm sync.
# 4-byte header + 413 bytes frame body = 417 bytes × 2 = 834 bytes total.
_MPEG_FRAMES = (b'\xff\xfb\x90\x00' + b'\x00' * 413) * 2


def _flac_block_header(block_type: int, data_len: int, last: bool = False) -> bytes:
    first = (0x80 if last else 0x00) | block_type
    return bytes([first, (data_len >> 16) & 0xFF, (data_len >> 8) & 0xFF, data_len & 0xFF])


def _vorbis_comment_block(comments: dict) -> bytes:
    vendor = b'GDPRScanner fixture'
    data = struct.pack('<I', len(vendor)) + vendor
    data += struct.pack('<I', len(comments))
    for key, value in comments.items():
        entry = f'{key}={value}'.encode('utf-8')
        data += struct.pack('<I', len(entry)) + entry
    return data


def _minimal_flac(comments: dict) -> bytes:
    """Return bytes for a valid minimal FLAC file with Vorbis comments."""
    # STREAMINFO (34 bytes): 44100 Hz, mono, 16-bit, 0 samples, zero MD5.
    si = bytearray(34)
    si[0:2] = struct.pack('>H', 4096)   # min block size
    si[2:4] = struct.pack('>H', 4096)   # max block size
    # bytes 4-9: min/max frame sizes = 0 (unknown)
    # Bits 80-99: sample_rate=44100 (0xAC44 in 20-bit field)
    # Bits 100-102: channels-1 = 0 (mono)
    # Bits 103-107: bits_per_sample-1 = 15 (16-bit)
    # Bits 108-143: total_samples = 0; bytes 14-17 remain zero
    si[10] = 0x0A   # 0000_1010 — top 8 of 44100 in 20-bit field
    si[11] = 0xC4   # 1100_0100
    si[12] = 0x40   # bottom 4 of sample_rate | channels(000) | bps_msb(0)
    si[13] = 0xF0   # bps remaining 4 bits (1111) | top 4 of total_samples (0)

    vc = _vorbis_comment_block(comments)
    return (
        b'fLaC'
        + _flac_block_header(0, 34, last=not comments)  # STREAMINFO
        + bytes(si)
        + (_flac_block_header(4, len(vc), last=True) + vc if comments else b'')
    )


def _mp4_atom(name: bytes, data: bytes) -> bytes:
    return struct.pack('>I', 8 + len(data)) + name + data


def _minimal_mp4_base() -> bytes:
    """Return bytes for the smallest valid MPEG-4 container mutagen can tag."""
    # ftyp — identifies the file as M4A
    ftyp = _mp4_atom(
        b'ftyp',
        b'M4A ' + struct.pack('>I', 0) + b'M4A ' + b'mp42' + b'isom',
    )
    # mvhd version 0 — 100 bytes of content (ISO 14496-12 §8.2.2)
    mvhd = bytearray(100)
    mvhd[0:4] = b'\x00\x00\x00\x00'                          # version + flags
    struct.pack_into('>IIII', mvhd, 4, 0, 0, 1000, 0)        # creation, modification, timescale, duration
    struct.pack_into('>I', mvhd, 16, 0x00010000)              # rate = 1.0
    struct.pack_into('>H', mvhd, 20, 0x0100)                  # volume = 1.0
    # bytes 22-31: reserved (10 bytes, already zero)
    struct.pack_into('>9i', mvhd, 32,                         # unity matrix
        0x00010000, 0, 0, 0, 0x00010000, 0, 0, 0, 0x40000000)
    # bytes 68-91: pre-defined (24 bytes, already zero)
    struct.pack_into('>I', mvhd, 96, 0xFFFFFFFF)              # next_track_ID

    return ftyp + _mp4_atom(b'moov', _mp4_atom(b'mvhd', bytes(mvhd)))


def _mp4_with_tags(tags: dict) -> bytes:
    """Return bytes for a minimal MP4 with the given mutagen tag dict."""
    import mutagen.mp4
    tmp = tempfile.mktemp(suffix='.mp4')
    try:
        with open(tmp, 'wb') as fh:
            fh.write(_minimal_mp4_base())
        f = mutagen.mp4.MP4(tmp)
        f.add_tags()
        for key, value in tags.items():
            f.tags[key] = [value]
        f.save()
        with open(tmp, 'rb') as fh:
            return fh.read()
    finally:
        if os.path.exists(tmp):
            os.unlink(tmp)


# ── 14_audio_artist_pii.mp3 ───────────────────────────────────────────────────
def make_mp3_pii():
    from mutagen.easyid3 import EasyID3
    tmp = tempfile.mktemp(suffix='.mp3')
    try:
        t = EasyID3()
        t['artist'] = ['Emma Slot Henriksen']
        t['title']  = ['Fortrolig optagelse — personalemøde']
        t['date']   = ['2026-04-21']
        t.save(tmp)
        with open(tmp, 'rb') as fh:
            id3_bytes = fh.read()
    finally:
        if os.path.exists(tmp):
            os.unlink(tmp)

    out = HERE / '14_audio_artist_pii.mp3'
    out.write_bytes(id3_bytes + _MPEG_FRAMES)
    print(f"Written: {out.name}")


# ── 15_audio_artist_pii.flac ──────────────────────────────────────────────────
def make_flac_pii():
    out = HERE / '15_audio_artist_pii.flac'
    out.write_bytes(_minimal_flac({
        'ARTIST': 'Emma Slot Henriksen',
        'TITLE':  'Fortrolig optagelse — personalemøde',
        'DATE':   '2026-04-21',
    }))
    print(f"Written: {out.name}")


# ── 16_audio_no_pii.mp3 ───────────────────────────────────────────────────────
def make_mp3_no_pii():
    from mutagen.easyid3 import EasyID3
    tmp = tempfile.mktemp(suffix='.mp3')
    try:
        EasyID3().save(tmp)  # empty ID3 header, no tags
        with open(tmp, 'rb') as fh:
            id3_bytes = fh.read()
    finally:
        if os.path.exists(tmp):
            os.unlink(tmp)

    out = HERE / '16_audio_no_pii.mp3'
    out.write_bytes(id3_bytes + _MPEG_FRAMES)
    print(f"Written: {out.name}")


# ── 17_audio_no_pii.flac ──────────────────────────────────────────────────────
def make_flac_no_pii():
    out = HERE / '17_audio_no_pii.flac'
    out.write_bytes(_minimal_flac({}))   # no Vorbis comment block
    print(f"Written: {out.name}")


# ── 18_video_gps.mp4 ─────────────────────────────────────────────────────────
def make_mp4_gps():
    out = HERE / '18_video_gps.mp4'
    out.write_bytes(_mp4_with_tags({
        '©xyz': '+55.6761+012.5683+000.000/',   # Copenhagen
        '©ART': 'Emma Slot Henriksen',
        '©nam': 'Optagelse fra skolegården',
    }))
    print(f"Written: {out.name}")


# ── 19_video_no_pii.mp4 ──────────────────────────────────────────────────────
def make_mp4_no_pii():
    out = HERE / '19_video_no_pii.mp4'
    out.write_bytes(_minimal_mp4_base())   # no moov/udta/meta/ilst — no tags
    print(f"Written: {out.name}")


if __name__ == "__main__":
    make_docx()
    make_xlsx()
    make_mp3_pii()
    make_flac_pii()
    make_mp3_no_pii()
    make_flac_no_pii()
    make_mp4_gps()
    make_mp4_no_pii()
    print("Done.")
