"""
conftest.py — shared fixtures for GDPRScanner test suite.
"""
import sys
import tempfile
from pathlib import Path

import pytest

# Ensure the project root is on sys.path so all modules are importable
ROOT = Path(__file__).parent.parent
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))


# ── File fixtures ─────────────────────────────────────────────────────────────

@pytest.fixture()
def tmp_dir(tmp_path):
    return tmp_path


@pytest.fixture()
def docx_with_cpr(tmp_path):
    """Word document containing 3 CPR numbers in different positions."""
    from docx import Document
    doc = Document()
    doc.add_paragraph("Elev 1: CPR 290472-1234 er registreret i systemet.")
    doc.add_paragraph("Elev 2: personnummer 010185-4321.")
    tbl = doc.add_table(rows=2, cols=2)
    tbl.cell(0, 0).text = "Navn"
    tbl.cell(0, 1).text = "CPR"
    tbl.cell(1, 0).text = "Anne Hansen"
    tbl.cell(1, 1).text = "CPR: 150364-5678"
    p = tmp_path / "sample_with_cpr.docx"
    doc.save(p)
    return p


@pytest.fixture()
def docx_no_cpr(tmp_path):
    """Word document with no CPR numbers."""
    from docx import Document
    doc = Document()
    doc.add_paragraph("Ingen personoplysninger her.")
    doc.add_paragraph("Konto: 1234-5678  Telefon: 33 12 34 56")
    p = tmp_path / "sample_no_cpr.docx"
    doc.save(p)
    return p


@pytest.fixture()
def xlsx_with_cpr(tmp_path):
    """Excel workbook containing 1 CPR in a cell."""
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws["A1"] = "Navn"
    ws["B1"] = "CPR"
    ws["A2"] = "Test Person"
    ws["B2"] = "CPR: 290472-1234"
    p = tmp_path / "sample_with_cpr.xlsx"
    wb.save(p)
    return p


@pytest.fixture()
def xlsx_no_cpr(tmp_path):
    """Excel workbook with account numbers that look CPR-like."""
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws["A1"] = "Kontonummer"
    ws["B1"] = "Beløb"
    ws["A2"] = "12345678"      # 8-digit — too short
    ws["A3"] = "29047212345"   # 11-digit — too long
    ws["A4"] = "Reg: 2904"
    p = tmp_path / "sample_no_cpr.xlsx"
    wb.save(p)
    return p


@pytest.fixture()
def txt_with_art9(tmp_path):
    """Plain text with CPR adjacent to Article 9 health keywords."""
    content = (
        "Eleven CPR 290472-1234 har diagnosen diabetes og modtager behandling.\n"
        "Kontakt læge vedr. sygemelding."
    )
    p = tmp_path / "sample_art9.txt"
    p.write_text(content, encoding="utf-8")
    return p


@pytest.fixture()
def binary_garbage(tmp_path):
    """Binary file that must not crash the scanner."""
    p = tmp_path / "sample_binary.bin"
    p.write_bytes(bytes(range(256)) * 100)
    return p


@pytest.fixture()
def tmp_db(tmp_path):
    """Fresh in-memory-path SQLite DB for each test."""
    from gdpr_db import ScanDB
    db_path = tmp_path / "test.db"
    db = ScanDB(str(db_path))
    yield db
    try:
        db_path.unlink()
    except Exception:
        pass
